from flask import Flask, render_template, request, redirect, url_for, session, render_template_string, escape  
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
import os
import pandas as pd
from datetime import date, datetime as dt
from datetime import datetime
from itsdangerous import URLSafeTimedSerializer, BadSignature
from sqlalchemy import or_
from sqlalchemy import and_ 

import psycopg2
from datetime import date, datetime
from dateutil.relativedelta import relativedelta
import ast
import datetime
import os
import json

from flask_cors import CORS
import re
# import spacy
from flask_mail import Mail, Message
from flask import render_template, redirect, url_for, flash
from flask_mail import Mail, Message
from itsdangerous import URLSafeTimedSerializer
# from spacy.matcher import Matcher
from flask import Flask, request, render_template
from sqlalchemy import ForeignKey
from sqlalchemy.orm import relationship
from itsdangerous import URLSafeTimedSerializer
from flask import request, render_template, flash, redirect, url_for
import secrets
import secrets
from urllib.parse import quote_plus
from flask_migrate import Migrate
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from datetime import datetime
import pytz 
from sqlalchemy import case, desc
from sqlalchemy.orm import aliased
import hashlib
import random
import string
import base64
import PyPDF2
from PyPDF2 import PdfFileReader
import docx
import re
from collections import Counter
import math
import google.generativeai as genai
from details import config
import docx2txt


# from sqlalchemy.ext.declarative import declarative_base

# Base = declarative_base()


app = Flask(__name__)
cors = CORS(app)

app.config['MAIL_SERVER'] = config.microsoft_server
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = config.microsoft_mail_username
app.config['MAIL_PASSWORD'] = config.microsoft_mail_password


# app.config['MAIL_SERVER'] = config.gmail_server
# app.config['MAIL_PORT'] = 587
# app.config['MAIL_USE_TLS'] = True
# app.config['MAIL_USERNAME'] = config.gmail_username
# app.config['MAIL_PASSWORD'] = config.gmail_password


mail = Mail(app)

app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get("DATABASE_URL")

app.config['SECRET_KEY'] = secrets.token_hex(16)
engine = create_engine(app.config['SQLALCHEMY_DATABASE_URI'])
Session = sessionmaker(bind=engine)
app.config['SECRET_KEY'] = secrets.token_hex(16)
# Specify the folder where uploaded resumes will be stored
UPLOAD_FOLDER = 'C:/Users/Makonis/PycharmProjects/login/'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# cors = CORS(app)
# Specify the allowed resume file extensions
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}
db = SQLAlchemy(app)
migrate = Migrate(app, db)
from datetime import timedelta
#hello

# Specify the folder where uploaded resumes will be stored
# UPLOAD_FOLDER = 'static/'
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# cors = CORS(app)
# Specify the allowed resume file extensions
# ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}

# db = SQLAlchemy(app)
# migrate = Migrate(app, db)

def generate_verification_token(user_id):
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    return serializer.dumps(user_id)

class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True,nullable=False)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(300), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100), nullable=False)
    user_type = db.Column(db.String(20), nullable=False)
    client = db.Column(db.String(100))
    candidate = relationship("Candidate", back_populates="user", uselist=False)
    is_active = db.Column(db.Boolean, default=True)
    is_verified = db.Column(db.Boolean, default=False)
    created_by = db.Column(db.String(50))
    otp = db.Column(db.String(6), default=False)
    registration_completed = db.Column(db.String(50))
    filename = db.Column(db.String(100))
    # image_file = db.Column(db.String(1000))
    image_file=db.Column(db.LargeBinary)
    image_deleted=db.Column(db.Boolean, default=False)
    def serialize(self):
        return {
            'id': self.id,
            'username': self.username,
            'name': self.name,
            'email': self.email,
            'user_type': self.user_type,
            'client': self.client,
            'is_active': self.is_active,
            'is_verified': self.is_verified,
            'created_by': self.created_by,
            'otp': self.otp,
            'registration_completed': self.registration_completed
        }
        
class Candidate(db.Model):
    __tablename__ = 'candidates'
    id = db.Column(db.Integer, primary_key=True)
    # id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    job_id = db.Column(db.Integer)
    name = db.Column(db.String(100), nullable=False)
    mobile = db.Column(db.String(20), nullable=False)
    email = db.Column(db.String(100), nullable=False)
    client = db.Column(db.String(100), nullable=False)
    current_company = db.Column(db.String(100))
    position = db.Column(db.String(100))
    profile = db.Column(db.String(200))
    current_job_location = db.Column(db.String(100))
    preferred_job_location = db.Column(db.String(100))
    # resume = db.Column(db.String(1000))
    resume = db.Column(db.LargeBinary)
    skills = db.Column(db.String(500))
    qualifications = db.Column(db.String(200))
    experience = db.Column(db.String(200))
    relevant_experience = db.Column(db.String(200))
    current_ctc = db.Column(db.String(200))
    expected_ctc = db.Column(db.String(200))
    notice_period = db.Column(db.String(20))
    last_working_date = db.Column(db.Date)
    buyout = db.Column(db.Boolean, default=False)
    holding_offer = db.Column(db.String(20))
    total = db.Column(db.Integer)
    package_in_lpa = db.Column(db.Float)
    recruiter = db.Column(db.String(100))
    management = db.Column(db.String(100))
    status = db.Column(db.String(100))
    reason_for_job_change = db.Column(db.String(200))
    remarks = db.Column(db.String(200))
    # screening_done = db.Column(db.Boolean, default=False)
    # rejected_at_screening = db.Column(db.Boolean, default=False)
    # l1_cleared = db.Column(db.Boolean, default=False)
    # rejected_at_l1 = db.Column(db.Boolean, default=False)
    # dropped_after_clearing_l1 = db.Column(db.Boolean, default=False)
    # l2_cleared = db.Column(db.Boolean, default=False)
    # rejected_at_l2 = db.Column(db.Boolean, default=False)
    # dropped_after_clearing_l2 = db.Column(db.Boolean, default=False)
    # onboarded = db.Column(db.Boolean, default=False)
    # dropped_after_onboarding = db.Column(db.Boolean, default=False)
    date_created = db.Column(db.Date, default=date.today)
    time_created = db.Column(db.Time, default=datetime.now().time)
    comments = db.Column(db.String(1000))
    linkedin_url = db.Column(db.String(1000))
    user_id = db.Column(db.Integer, ForeignKey('users.id'))
    serving_notice_period = db.Column(db.String(200))
    period_of_notice = db.Column(db.String(1000))
    user = relationship("User", back_populates="candidate")
    reference = db.Column(db.String(200))
    reference_name = db.Column(db.String(200))
    reference_position = db.Column(db.String(200))
    reference_information = db.Column(db.String(200))
    data_updated_date = db.Column(db.Date)
    data_updated_time = db.Column(db.Time)
    resume_present = db.Column(db.Boolean, default=False)
    # resume_present = db.Column(db.Boolean, default=True)
    def serialize(self):
        return {
            'id': self.id,
            'job_id': self.job_id,
            'name': self.name,
            'mobile': self.mobile,
            'email': self.email,
            'client': self.client,
            'current_company': self.current_company,
            'position': self.position,
            'profile': self.profile,
            'current_job_location': self.current_job_location,
            'preferred_job_location': self.preferred_job_location,
            'resume': self.resume,
            'skills': self.skills,
            'qualifications': self.qualifications,
            'experience': self.experience,
            'relevant_experience': self.relevant_experience,
            'current_ctc': self.current_ctc,
            'expected_ctc': self.expected_ctc,
            'notice_period': self.notice_period,
            'last_working_date': self.last_working_date.strftime('%Y-%m-%d') if self.last_working_date else None,
            'buyout': self.buyout,
            'holding_offer': self.holding_offer,
            'total': self.total,
            'package_in_lpa': self.package_in_lpa,
            'recruiter': self.recruiter,
            'management': self.management,
            'status': self.status,
            'reason_for_job_change': self.reason_for_job_change,
            'remarks': self.remarks,
            # 'screening_done': self.screening_done,
            # 'rejected_at_screening': self.rejected_at_screening,
            # 'l1_cleared': self.l1_cleared,
            # 'rejected_at_l1': self.rejected_at_l1,
            # 'dropped_after_clearing_l1': self.dropped_after_clearing_l1,
            # 'l2_cleared': self.l2_cleared,
            # 'rejected_at_l2': self.rejected_at_l2,
            # 'dropped_after_clearing_l2': self.dropped_after_clearing_l2,
            # 'onboarded': self.onboarded,
            # 'dropped_after_onboarding': self.dropped_after_onboarding,
            'date_created': self.date_created.strftime('%Y-%m-%d'),
            'time_created': self.time_created.strftime('%H:%M:%S'),
            'comments': self.comments,
            'linkedin_url': self.linkedin_url,
            'user_id': self.user_id,
            'period_of_notice': self.period_of_notice,
            'reference': self.reference,
            'reference_name': self.reference_name,
            'reference_position': self.reference_position,
            'reference_information': self.reference_information,
            
        }
        
class Career_user(db.Model):
    __tablename__ = 'career_users'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True, nullable=False)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(50), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100), nullable=False)
    user_type = db.Column(db.String(50), default="career_visitor")


class Career_notification(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    recruiter_name = db.Column(db.String(100), nullable=False)
    notification_status = db.Column(db.Boolean, default=False)

    def __init__(self, recruiter_name, notification_status=False):
        self.recruiter_name = recruiter_name
        self.notification_status = notification_status


class JobPost(db.Model):
    __tablename__ = 'job_posts'

    id = db.Column(db.Integer, primary_key=True)
    client = db.Column(db.String(100))
    # experience_min = db.Column(db.Integer)
    # experience_max = db.Column(db.Integer)
    experience_min = db.Column(db.String(100))
    experience_max = db.Column(db.String(100))
    budget_min = db.Column(db.String(300))
    budget_max = db.Column(db.String(300))
    location = db.Column(db.String(100))
    shift_timings = db.Column(db.String(100))
    notice_period = db.Column(db.String(100))
    role = db.Column(db.String(100))
    detailed_jd = db.Column(db.Text)
    # jd_pdf =  db.Column(db.String(1000))
    jd_pdf =  db.Column(db.LargeBinary)
    mode = db.Column(db.String(100))
    recruiter = db.Column(db.String(1000))
    management = db.Column(db.String(100))
    date_created = db.Column(db.Date)
    time_created = db.Column(db.Time)
    job_status = db.Column(db.String(20))
    job_type = db.Column(db.String(100))
    contract_in_months = db.Column(db.String(100))
    # contract_in_months = db.Column(db.Integer, nullable=True)
    skills = db.Column(db.String(500))
    notification = db.Column(db.String(20))
    data_updated_date = db.Column(db.Date)
    data_updated_time = db.Column(db.Time)
    jd_pdf_present = db.Column(db.Boolean, default=False)
    # jd_pdf_present = db.Column(db.Boolean, default=True)
    no_of_positions = db.Column(db.String(100))
    def __init__(self, client, experience_min, experience_max, budget_min, budget_max, location, shift_timings, notice_period, role, detailed_jd, mode, recruiter, management, job_status, job_type, skills, jd_pdf, jd_pdf_present,contract_in_months, no_of_positions):
        self.client = client
        self.experience_min = experience_min
        self.experience_max = experience_max
        self.budget_min = budget_min
        self.budget_max = budget_max
        self.location = location
        self.shift_timings = shift_timings
        self.notice_period = notice_period
        self.role = role
        self.detailed_jd = detailed_jd
        self.mode = mode
        self.recruiter = recruiter
        self.management = management
        self.job_status = job_status
        self.job_type = job_type
        self.skills = skills
        self.jd_pdf = jd_pdf
        self.contract_in_months = contract_in_months
        self.jd_pdf_present = jd_pdf_present
        self.no_of_positions = no_of_positions

class Deletedcandidate(db.Model):
    _tablename_ = 'deletedcandidate'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), nullable=False)
    candidate_name = db.Column(db.String(100), nullable=False)
    candidate_email = db.Column(db.String(100), nullable=False)
    client = db.Column(db.String(100), nullable=False)
    profile = db.Column(db.String(100), nullable=False)
    status = db.Column(db.String(100), nullable=False)

# class Notification(db.Model):
#     id = db.Column(db.Integer, primary_key=True)
#     recruiter_name = db.Column(db.String(100), nullable=False)
#     notification_status = db.Column(db.Boolean, default=False)

#     def __init__(self, recruiter_name, notification_status=False):
#         self.recruiter_name = recruiter_name
#         self.notification_status = notification_status
class Notification(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    job_post_id = db.Column(db.Integer, db.ForeignKey('job_posts.id'))
    recruiter_name = db.Column(db.String(100), nullable=False)
    notification_status = db.Column(db.Boolean, default=False)
    num_notification = db.Column(db.Integer, default=0)  # New column added

    def __init__(self, job_post_id, recruiter_name, notification_status=False, num_notification=1):
        self.job_post_id = job_post_id
        self.recruiter_name = recruiter_name
        self.notification_status = notification_status
        self.num_notification = num_notification
        # self.num_notification = 0  # Default value for num_notification


class ScheduledMeeting(db.Model):
    __tablename__ = 'scheduled_meetings'
    
    id = db.Column(db.Integer, primary_key=True)
    event_id = db.Column(db.String(500), nullable=False)
    recruiter_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    subject = db.Column(db.String(255), nullable=False)
    start_date = db.Column(db.Date, nullable=False)
    start_time = db.Column(db.Time, nullable=False)
    end_date = db.Column(db.Date, nullable=False)
    end_time = db.Column(db.Time, nullable=False)
    attendees = db.Column(db.Text, nullable=False)
    cc_recipients = db.Column(db.Text, nullable=True)
    recruiter_email = db.Column(db.String(255), nullable=False)
    time_zone = db.Column(db.String(50), nullable=False)
    join_url = db.Column(db.String(512), nullable=True)  # Add this field
    
    created_at = db.Column(db.DateTime, default=db.func.current_timestamp(), nullable=False)


    # Relationship to User
    #recruiter = db.relationship("User", backref="scheduled_meetings")

    def __init__(self, event_id, subject, start_date, start_time,end_date,end_time, attendees, cc_recipients, recruiter_id, recruiter_email, time_zone, join_url):
        self.event_id = event_id
        self.subject = subject
        self.start_date = start_date
        self.start_time = start_time
        self.end_date = end_date
        self.end_time = end_time
        self.attendees = attendees
        self.cc_recipients = cc_recipients
        self.recruiter_id = recruiter_id
        self.recruiter_email = recruiter_email
        self.time_zone = time_zone
        self.join_url = join_url

###################################################################################################
from functools import lru_cache
from concurrent.futures import ThreadPoolExecutor
from sqlalchemy import case

@lru_cache(maxsize=100)
def generate_skills_for_job_role(job_role: str) -> dict:
    api_key = config.api_key
    if not api_key:
        raise ValueError("API_KEY environment variable not set")

    # Configure the genai API
    genai.configure(api_key=api_key)
    
    # Create a model instance
    model = genai.GenerativeModel('gemini-1.5-flash')

    # Construct the prompt for the Generative AI model
    prompt = f"""
    Given the following job role: {job_role}, provide 20 related technical topics for the given job role. Exclude soft skills (such as communication, fast learner). Present the output in the following format:

    sub_categories = {{
        '{job_role}': ['Core Java', 'J2EE', 'Servlets', 'Spring Boot', 'JSP', 'MySQL', 'PostgreSQL', 'Docker', 'Jenkins', ...]
    }}
    """

    # Generate content from the model
    response = model.generate_content(prompt)
    response_text = response.candidates[0].content.parts[0].text.strip()

    # Extract the generated sub-categories from the response
    match = re.search(r"sub_categories\s*=\s*({.*})", response_text, re.DOTALL)
    if match:
        response_dict_str = match.group(1).replace("'", "\"")
        try:
            response_dict = json.loads(response_dict_str)
        except json.JSONDecodeError as e:
            print(f"Error parsing response dictionary: {e}")
            return {}

        # Clean up and format the dictionary
        cleaned_dict = {key.strip(): [item.strip().lower() for item in value] for key, value in response_dict.items()}
        return cleaned_dict
    else:
        print("No valid sub-categories found in response.")
        return {}

def calculate_match_percentage(job_role: str, candidate_skills: str) -> tuple:
    if not job_role or not candidate_skills:
        print("Job role or candidate skills missing.")
        return 0.0, [], {}

    sub_skills_dict = generate_skills_for_job_role(job_role)
    
    if not sub_skills_dict:
        print("No sub-skills dictionary generated.")
        return 0.0, [], {}

    candidate_skills_set = set(candidate_skills.lower().split(", "))
    print(f"Candidate skills set: {candidate_skills_set}")

    matched_skills = []
    total_sub_skills = 0

    for sub_skills in sub_skills_dict.values():
        sub_skills_set = set(sub_skills)
        total_sub_skills += len(sub_skills_set)
        matched_skills.extend(list(sub_skills_set.intersection(candidate_skills_set)))
        print("\n")
        print(f"Sub-skills: {sub_skills_set}")
        print("\n")
        print(f"Matched skills: {matched_skills}")

    if total_sub_skills == 0:
        return 0.0, [], sub_skills_dict

    match_percentage = (len(matched_skills) / total_sub_skills) * 100
    print(f"Match percentage: {match_percentage}")
    return round(match_percentage, 2), matched_skills, sub_skills_dict


@app.route('/search_resumes', methods=['POST'])
def search_resumes():
    job_role = request.json.get('job_role')
    if not job_role:
        return jsonify({"error": "Job role is required"}), 400

    print(f"Received job role: {job_role}")

    # Generate skills for the job role
    sub_skills_dict = generate_skills_for_job_role(job_role)
    if not sub_skills_dict:
        return jsonify({"error": "No skills found for the provided job role"}), 500

    # Flatten the dictionary into a single list of skills
    gemini_skills = set(skill for skills in sub_skills_dict.values() for skill in skills)

    # Fetch necessary columns, including a check if the resume is present
    candidates = Candidate.query.with_entities(
        Candidate.id,
        Candidate.client,
        Candidate.name,
        Candidate.skills,
        Candidate.email,
        Candidate.mobile,
        Candidate.job_id,
        Candidate.profile,
        Candidate.date_created,
        Candidate.data_updated_date,
        Candidate.status,  # Include status
        case(
            (Candidate.resume.isnot(None) & (Candidate.resume != b"")),
            else_=False
        ).label('resume_present')
    ).all()

    def process_candidate(candidate):
        # Unpack the tuple
        candidate_id, client, name, skills, email, mobile, job_id, profile, date_created, data_updated_date, status, resume_present = candidate

        formatted_date_created = date_created.strftime('%Y-%m-%d') if date_created else None
        formatted_data_updated_date = data_updated_date.strftime('%Y-%m-%d') if data_updated_date else None

        if skills:
            candidate_skills_set = set(skills.lower().split(", "))
            gemini_skills_with_role = gemini_skills.union({job_role.lower()})
            matched_skills = list(gemini_skills_with_role.intersection(candidate_skills_set))
            if matched_skills:
                match_percentage = (len(matched_skills) / len(gemini_skills_with_role)) * 100
                return {
                    "candidate_id": candidate_id,
                    "candidate_name": name,
                    "match_percentage": round(match_percentage, 2),
                    "email": email,
                    "mobile": mobile,
                    "job_id": job_id,
                    "client": client,
                    "date_created": formatted_date_created,
                    "data_updated_date": formatted_data_updated_date,
                    "status": status,
                    "profile": profile,
                    "resume_present": resume_present,
                    "candidate_skills": skills.split(", "),  # Original candidate skills
                    "matched_skills": matched_skills,  # Skills that matched with Gemini-generated skills
                    "gemini_generated_skills": list(gemini_skills_with_role)  # Skills generated by Gemini for the job role plus job role itself
                }
        return None

    with ThreadPoolExecutor() as executor:
        results = executor.map(process_candidate, candidates)

    matching_resumes = [result for result in results if result]
    matching_resumes = sorted(matching_resumes, key=lambda x: x["match_percentage"], reverse=True)

    print(f"Matching resumes: {matching_resumes}")
    return jsonify(matching_resumes)



# @app.route('/search_resumes', methods=['POST'])
# def search_resumes():
#     job_role = request.json.get('job_role')
#     if not job_role:
#         return jsonify({"error": "Job role is required"}), 400

#     print(f"Received job role: {job_role}")

#     # Generate skills for the job role
#     sub_skills_dict = generate_skills_for_job_role(job_role)
#     if not sub_skills_dict:
#         return jsonify({"error": "No skills found for the provided job role"}), 500

#     # Flatten the dictionary into a single list of skills
#     gemini_skills = set(skill for skills in sub_skills_dict.values() for skill in skills)

#     # Fetch necessary columns, including a check if the resume is present
#     candidates = Candidate.query.with_entities(
#         Candidate.id,
#         Candidate.name,
#         Candidate.skills,
#         Candidate.email,
#         Candidate.mobile,
#         Candidate.job_id,
#         Candidate.client,
#         Candidate.profile,
#         case(
#             (Candidate.resume.isnot(None) & (Candidate.resume != b"")),  # If resume is binary
#             else_=False
#         ).label('resume_present')
#     ).all()

#     def process_candidate(candidate):
#         if candidate.skills:
#             candidate_skills_set = set(candidate.skills.lower().split(", "))
#             gemini_skills_with_role = gemini_skills.union({job_role.lower()})
#             matched_skills = list(gemini_skills_with_role.intersection(candidate_skills_set))
#             if matched_skills:
#                 match_percentage = (len(matched_skills) / len(gemini_skills_with_role)) * 100
#                 return {
#                     "candidate_id": candidate.id,
#                     "candidate_name": candidate.name,
#                     "match_percentage": round(match_percentage, 2),
#                     "email": candidate.email,
#                     "mobile": candidate.mobile,
#                     "job_id": candidate.job_id,
#                     "client":candidate.client,
#                     "profile": candidate.profile,
#                     "resume_present": candidate.resume_present,  # Use the value from the query
#                     "candidate_skills": candidate.skills.split(", "),  # Original candidate skills
#                     "matched_skills": matched_skills,  # Skills that matched with Gemini-generated skills
#                     "gemini_generated_skills": list(gemini_skills_with_role)  # Skills generated by Gemini for the job role plus job role itself
#                 }
#         return None

#     with ThreadPoolExecutor(max_workers=1000) as executor:  # Adjust the number of workers to optimize performance
#         results = executor.map(process_candidate, candidates)

#     matching_resumes = [result for result in results if result]
#     matching_resumes = sorted(matching_resumes, key=lambda x: x["match_percentage"], reverse=True)

#     print(f"Matching resumes: {matching_resumes}")
#     return jsonify(matching_resumes)



# @app.route('/search_resumes', methods=['POST'])
# def search_resumes():
#     job_role = request.json.get('job_role')
#     if not job_role:
#         return jsonify({"error": "Job role is required"}), 400

#     print(f"Received job role: {job_role}")

#     # Generate skills for the job role
#     sub_skills_dict = generate_skills_for_job_role(job_role)
#     if not sub_skills_dict:
#         return jsonify({"error": "No skills found for the provided job role"}), 500

#     # Flatten the dictionary into a single list of skills
#     gemini_skills = set(skill for skills in sub_skills_dict.values() for skill in skills)

#     # Fetch only necessary columns to reduce overhead
#     candidates = Candidate.query.with_entities(
#         Candidate.id, Candidate.name, Candidate.skills, Candidate.email, Candidate.mobile, Candidate.job_id, Candidate.profile
#     ).all()
#     matching_resumes = []

#     def process_candidate(candidate):
#         if candidate.skills:
#             candidate_skills_set = set(candidate.skills.lower().split(", "))
#             gemini_skills_with_role = gemini_skills.union({job_role.lower()})
#             matched_skills = list(gemini_skills_with_role.intersection(candidate_skills_set))
#             if matched_skills:
#                 match_percentage = (len(matched_skills) / len(gemini_skills_with_role)) * 100
#                 return {
#                     "candidate_id": candidate.id,
#                     "candidate_name": candidate.name,
#                     "match_percentage": round(match_percentage, 2),
#                     "email": candidate.email,
#                     "mobile": candidate.mobile,
#                     "job_id": candidate.job_id,
#                     "profile": candidate.profile,
#                     "candidate_skills": candidate.skills.split(", "),  # Original candidate skills
#                     "matched_skills": matched_skills,  # Skills that matched with Gemini-generated skills
#                     "gemini_generated_skills": list(gemini_skills_with_role)  # Skills generated by Gemini for the job role plus job role itself
#                 }
#         return None

#     with ThreadPoolExecutor(max_workers=1000) as executor:  # Adjust the number of workers based on your server capabilities
#         results = executor.map(process_candidate, candidates)
    
#     matching_resumes = [result for result in results if result]
#     matching_resumes = sorted(matching_resumes, key=lambda x: x["match_percentage"], reverse=True)

#     print(f"Matching resumes: {matching_resumes}")
#     return jsonify(matching_resumes)

####################################################################################################

from flask import Flask, request, jsonify
import requests
import json
from msal import ConfidentialClientApplication



TENANT_ID = '8a7c6498-6635-4dbc-8a5e-f38efccfef3e'
CLIENT_ID = '7ba39e41-0ec7-411d-8649-6607574db5f9'
CLIENT_SECRET = '2ET8Q~TkfssfstnbCWmFP2U24phkklo_w080uc7E'
AUTHORITY = f'https://login.microsoftonline.com/{TENANT_ID}'
SCOPES = ['https://graph.microsoft.com/.default']
GRAPH_ENDPOINT = 'https://graph.microsoft.com/v1.0'

# Initialize MSAL ConfidentialClientApplication
msal_app = ConfidentialClientApplication(
    CLIENT_ID,
    authority=AUTHORITY,
    client_credential=CLIENT_SECRET,
)

# Function to acquire access token
def get_access_token():
    result = msal_app.acquire_token_for_client(scopes=SCOPES)
    if "access_token" in result:
        return result["access_token"]
    else:
        print("Failed to obtain access token:", result.get("error_description"))
        return None




def strip_fractional_seconds(date_time_str):
    if '.' in date_time_str:
        date_time_str = date_time_str.split('.')[0]
    return date_time_str

def convert_to_ist(date_time_str, from_timezone):
    from_tz = pytz.timezone(from_timezone)
    
    # Strip fractional seconds if present
    date_time_str = strip_fractional_seconds(date_time_str)
    
    # Parse the datetime string and localize it to the given timezone
    utc_time = datetime.strptime(date_time_str, "%Y-%m-%dT%H:%M:%S")
    utc_time = from_tz.localize(utc_time)

    # Convert to IST timezone
    ist_timezone = pytz.timezone('Asia/Kolkata')
    ist_time = utc_time.astimezone(ist_timezone)
    
    return ist_time

def format_time_for_graph_api(time_str):
    parts = time_str.split(":")
    hour = parts[0].zfill(2)  # Add leading zero if necessary
    minutes = parts[1] if len(parts) > 1 else "00"
    return f"{hour}:{minutes}:00"  # Ensure seconds are included



# Function to create an event with attendees and CC recipients
def create_event(subject, start_date, start_time, end_date, end_time, attendees, cc_recipients, recruiter_email, time_zone):
    access_token = get_access_token()

    if not access_token:
        return None, "Access token not available"

    headers = {
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/json'
    }

    # Format start and end times to include seconds
    start_time = format_time_for_graph_api(start_time)
    end_time = format_time_for_graph_api(end_time)

    # Combine date and time into ISO 8601 format
    start_date_time = f"{start_date}T{start_time}"
    end_date_time = f"{end_date}T{end_time}"

    # Prepare attendees, including those in CC
    all_attendees = [
        {
            'emailAddress': {
                'address': attendee,
                'name': ''
            },
            'type': 'required'
        } for attendee in attendees
    ] + [
        {
            'emailAddress': {
                'address': cc,
                'name': ''
            },
            'type': 'optional'  # CC recipients are marked as 'optional'
        } for cc in cc_recipients
    ]

    event = {
        'subject': subject,
        'start': {
            'dateTime': start_date_time,
            'timeZone': time_zone
        },
        'end': {
            'dateTime': end_date_time,
            'timeZone': time_zone
        },
        'attendees': all_attendees,
        'isOnlineMeeting': True,
        'onlineMeetingProvider': 'teamsForBusiness'
    }

    response = requests.post(
        f'https://graph.microsoft.com/v1.0/users/{recruiter_email}/events',
        headers=headers,
        data=json.dumps(event)
    )

    if response.status_code != 201:
        return None, f"Error creating event: {response.status_code} - {response.text}"
    
    return response.json(), None



# Function to create an event with attendees and CC recipients
# def create_event(subject, start_date, start_time, end_date, end_time, attendees, cc_recipients, recruiter_email, time_zone):
#     access_token = get_access_token()

#     if not access_token:
#         return None, "Access token not available"

#     headers = {
#         'Authorization': f'Bearer {access_token}',
#         'Content-Type': 'application/json'
#     }

#     # Combine date and time into ISO 8601 format without seconds
#     start_date_time = f"{start_date}T{start_time[:5]}"  # Removing seconds
#     end_date_time = f"{end_date}T{end_time[:5]}"        # Removing seconds

#     # Prepare attendees, including those in CC
#     all_attendees = [
#         {
#             'emailAddress': {
#                 'address': attendee,
#                 'name': ''
#             },
#             'type': 'required'
#         } for attendee in attendees
#     ] + [
#         {
#             'emailAddress': {
#                 'address': cc,
#                 'name': ''
#             },
#             'type': 'optional'  # CC recipients are marked as 'optional'
#         } for cc in cc_recipients
#     ]

#     event = {
#         'subject': subject,
#         'start': {
#             'dateTime': start_date_time,
#             'timeZone': time_zone
#         },
#         'end': {
#             'dateTime': end_date_time,
#             'timeZone': time_zone
#         },
#         'attendees': all_attendees,
#         'isOnlineMeeting': True,
#         'onlineMeetingProvider': 'teamsForBusiness'
#     }

#     response = requests.post(
#         f'https://graph.microsoft.com/v1.0/users/{recruiter_email}/events',
#         headers=headers,
#         data=json.dumps(event)
#     )

#     if response.status_code != 201:
#         return None, f"Error creating event: {response.status_code} - {response.text}"
    
#     return response.json(), None

# Function to update an event
def update_event(event_id, subject, start_date, start_time, end_date, end_time, attendees, cc_recipients, recruiter_email, time_zone):
    access_token = get_access_token()

    if not access_token:
        return None, "Access token not available"

    headers = {
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/json'
    }

    # Combine date and time into ISO 8601 format without seconds
    start_date_time = f"{start_date}T{start_time[:5]}"  # Removing seconds
    end_date_time = f"{end_date}T{end_time[:5]}"        # Removing seconds

    # Prepare attendees, including those in CC
    all_attendees = [
        {
            'emailAddress': {
                'address': attendee,
                'name': ''
            },
            'type': 'required'
        } for attendee in attendees
    ] + [
        {
            'emailAddress': {
                'address': cc,
                'name': ''
            },
            'type': 'optional'
        } for cc in cc_recipients
    ]

    event = {
        'subject': subject,
        'start': {
            'dateTime': start_date_time,
            'timeZone': time_zone
        },
        'end': {
            'dateTime': end_date_time,
            'timeZone': time_zone
        },
        'attendees': all_attendees,
        'isOnlineMeeting': True,
        'onlineMeetingProvider': 'teamsForBusiness'
    }

    response = requests.patch(
        f'https://graph.microsoft.com/v1.0/users/{recruiter_email}/events/{event_id}',
        headers=headers,
        data=json.dumps(event)
    )

    if response.status_code != 200:
        return None, f"Error updating event: {response.status_code} - {response.text}"
    
    return response.json(), None


# Function to delete an event
def delete_event(event_id, recruiter_email):
    access_token = get_access_token()

    if not access_token:
        return None, "Access token not available"

    headers = {
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/json'
    }

    response = requests.delete(
        f'https://graph.microsoft.com/v1.0/users/{recruiter_email}/events/{event_id}',
        headers=headers
    )

    if response.status_code != 204:
        return None, f"Error deleting event: {response.status_code} - {response.text}"
    
    return True, None



@app.route('/create_event', methods=['POST'])
def handle_create_event():
    data = request.json
    
    if not data:
        return jsonify({'error': 'Invalid request, no JSON body provided'}), 400
    
    subject = data.get('subject')
    start_date = data.get('start_date')
    start_time = data.get('start_time')
    end_date = data.get('end_date')
    end_time = data.get('end_time')
    attendees = data.get('attendees')
    cc_recipients = data.get('cc_recipients', [])  # Defaults to an empty list if not provided
    recruiter_email = data.get('recruiter_email')
    time_zone = data.get('time_zone')  # Default to UTC if not provided
    
    if not all([subject, start_date, start_time, end_date, end_time, attendees, recruiter_email]):
        return jsonify({'error': 'Missing required fields'}), 400
    
    # Combine date and time into ISO 8601 format
    start_date_time_str = f"{start_date}T{start_time}:00"
    end_date_time_str = f"{end_date}T{end_time}:00"

    # Convert start and end times to IST
    start_date_time_ist = convert_to_ist(start_date_time_str, time_zone)
    end_date_time_ist = convert_to_ist(end_date_time_str, time_zone)

    # Extract date and time portions for storage
    start_date_ist = start_date_time_ist.strftime("%Y-%m-%d")
    start_time_ist = start_date_time_ist.strftime("%H:%M:%S")
    end_date_ist = end_date_time_ist.strftime("%Y-%m-%d")
    end_time_ist = end_date_time_ist.strftime("%H:%M:%S")
    
    event_response, error = create_event(
        subject=subject,
        start_date=start_date_ist,
        start_time=start_time_ist,
        end_date=end_date_ist,
        end_time=end_time_ist,
        attendees=attendees,
        cc_recipients=cc_recipients,
        recruiter_email=recruiter_email,
        time_zone="Asia/Kolkata"  # Use IST timezone
    )

    if event_response:
        event_id = event_response.get('id')
        
        # Correctly extract joinUrl from the onlineMeeting field
        join_url = event_response.get('onlineMeeting', {}).get('joinUrl', '')
        print("join_url:", join_url)

        # Query the recruiter based on their email
        recruiter = db.session.query(User).filter_by(email=recruiter_email).first()
        if recruiter:
            recruiter_id = recruiter.id
        else:
            return jsonify({"error": "Recruiter not found"}), 404

        # Insert the event details into the scheduled_meeting table
        new_meeting = ScheduledMeeting(
            event_id=event_id,
            recruiter_id=recruiter_id,
            subject=subject,
            start_date=start_date_ist,
            start_time=start_time_ist,
            end_date=end_date_ist,
            end_time=end_time_ist,
            attendees=','.join(attendees),
            cc_recipients=','.join(cc_recipients),
            recruiter_email=recruiter_email,
            time_zone="Asia/Kolkata",  # Use IST timezone
            join_url=join_url  # Save joinUrl in the database
        )

        db.session.add(new_meeting)
        db.session.commit()
        
        return jsonify({
            'message': 'Event created and saved successfully.',
            'event': event_response,
            'joinUrl': join_url
        }), 200
    else:
        return jsonify({'error': error}), 500


# Route to create an event
# @app.route('/create_event', methods=['POST'])
# def handle_create_event():
#     data = request.json
    
#     if not data:
#         return jsonify({'error': 'Invalid request, no JSON body provided'}), 400
    
#     subject = data.get('subject')
#     start_date = data.get('start_date')
#     start_time = data.get('start_time')
#     end_date = data.get('end_date')
#     end_time = data.get('end_time')
#     attendees = data.get('attendees')
#     cc_recipients = data.get('cc_recipients', [])  # Defaults to an empty list if not provided
#     recruiter_email = data.get('recruiter_email')
#     time_zone = data.get('time_zone')  # Default to UTC if not provided
    
#     if not all([subject, start_date, start_time, end_date, end_time, attendees, recruiter_email]):
#         return jsonify({'error': 'Missing required fields'}), 400
    
#     # Fetch recruiter_id based on recruiter_email
#     recruiter = db.session.query(User).filter_by(email=recruiter_email).first()
#     if recruiter:
#         recruiter_id = recruiter.id
#     else:
#         return jsonify({"error": "Recruiter not found"}), 404
    
#     event_response, error = create_event(
#         subject=subject,
#         start_date=start_date,
#         start_time=start_time,
#         end_date=end_date,
#         end_time=end_time,
#         attendees=attendees,
#         cc_recipients=cc_recipients,
#         recruiter_email=recruiter_email,
#         time_zone=time_zone
#     )
#     if event_response:
#         event_id = event_response.get('id')
        
#         # Correctly extract joinUrl from the onlineMeeting field
#         join_url = event_response.get('onlineMeeting', {}).get('joinUrl', '')
#         print("join_url:", join_url)

#         # Insert the event details into the scheduled_meeting table
#         new_meeting = ScheduledMeeting(
#             event_id=event_id,
#             recruiter_id=recruiter_id,
#             subject=subject,
#             start_date=start_date,
#             start_time=start_time,
#             end_date=end_date,
#             end_time=end_time,
#             attendees=','.join(attendees),
#             cc_recipients=','.join(cc_recipients),
#             recruiter_email=recruiter_email,
#             time_zone=time_zone,
#             join_url=join_url  # Save joinUrl in the database
#         )

#         db.session.add(new_meeting)
#         db.session.commit()
        
#         return jsonify({
#             'message': 'Event created and saved successfully.',
#             'event': event_response,
#             'joinUrl': join_url
#         }), 200
#     else:
#         return jsonify({'error': error}), 500


# Route to update an event
@app.route('/update_event', methods=['POST'])
def handle_update_event():
    data = request.json
    meeting_id = data.get('meeting_id')
    
    meetings = ScheduledMeeting.query.filter_by(id=meeting_id).first()
    event_id = meetings.event_id

    if not data:
        return jsonify({'error': 'Invalid request, no JSON body provided'}), 400
    
    subject = data.get('subject')
    start_date = data.get('start_date')
    start_time = data.get('start_time')
    end_date = data.get('end_date')
    end_time = data.get('end_time')
    attendees = data.get('attendees')
    cc_recipients = data.get('cc_recipients', [])  # Defaults to an empty list if not provided
    recruiter_email = data.get('recruiter_email')
    time_zone = data.get('time_zone')  # Default to UTC if not provided
    
    if not all([subject, start_date, start_time, end_date, end_time, attendees, recruiter_email,meeting_id]):
        return jsonify({'error': 'Missing required fields'}), 400
    
    event_response, error = update_event(
        event_id=event_id,
        subject=subject,
        start_date=start_date,
        start_time=start_time,
        end_date=end_date,
        end_time=end_time,
        attendees=attendees,
        cc_recipients=cc_recipients,
        recruiter_email=recruiter_email,
        time_zone=time_zone
    )
    
    if event_response:
        # Fetch the existing meeting record
        scheduled_meeting = db.session.query(ScheduledMeeting).filter_by(event_id=event_id).first()
        if scheduled_meeting:
            # Update the existing meeting record
            scheduled_meeting.subject = subject
            scheduled_meeting.start_date = start_date
            scheduled_meeting.start_time = start_time
            scheduled_meeting.end_date = end_date
            scheduled_meeting.end_time = end_time
            scheduled_meeting.attendees = ','.join(attendees)
            scheduled_meeting.cc_recipients = ','.join(cc_recipients)
            scheduled_meeting.time_zone = time_zone
            db.session.commit()
            return jsonify({'message': 'Event updated successfully.', 'event': event_response}), 200
        else:
            return jsonify({'error': 'Event not found'}), 404
    else:
        return jsonify({'error': error}), 500


@app.route('/delete_event', methods=['POST'])
def handle_delete_event():
    data = request.json
    
    if not data:
        return jsonify({'error': 'Invalid request, no JSON body provided'}), 400

    meeting_id = data.get('meeting_id')
    
    if not meeting_id:
        return jsonify({'error': 'Missing meeting_id field'}), 400

    # Fetch the meeting details
    meeting = ScheduledMeeting.query.filter_by(id=meeting_id).first()

    if not meeting:
        return jsonify({'error': 'Meeting not found'}), 404

    event_id = meeting.event_id
    recruiter_email = meeting.recruiter_email

    # Call the delete_event function
    success, error = delete_event(event_id, recruiter_email)
    
    if success:
        # Remove the meeting record from the database
        db.session.delete(meeting)
        db.session.commit()
        return jsonify({'message': 'Event deleted successfully.'}), 200
    else:
        return jsonify({'error': error}), 500


@app.route('/get_all_meetings', methods=['POST'])
def get_all_meetings():
    data = request.json
    recruiter_id = data.get('recruiter_id')
    
    if not recruiter_id:
        return jsonify({'error': 'Missing recruiter_id field'}), 400
    
    try:
        # Get all users associated with the recruiter_id
        users = User.query.filter_by(id=recruiter_id).all()
        user_emails = {user.email for user in users}

        # Query the ScheduledMeeting table for all records
        all_meetings = ScheduledMeeting.query.all()

        # Filter meetings based on valid emails
        meetings_data = []
        for meeting in all_meetings:
            attendees = meeting.attendees.split(',') if meeting.attendees else []
            cc_recipients = meeting.cc_recipients.split(',') if meeting.cc_recipients else []

            # Check if any attendees or CC recipients are in the list of user emails
            valid_attendees = [email for email in attendees if email in user_emails]
            valid_cc_recipients = [email for email in cc_recipients if email in user_emails]

            # Check if recruiter_email is a valid user email
            valid_recruiter_email = meeting.recruiter_email in user_emails

            # Include the meeting if recruiter_email is valid, or if there are valid attendees or cc_recipients
            if valid_recruiter_email or valid_attendees or valid_cc_recipients:
                meeting_dict = {
                    'meeting_id': meeting.id,
                    'event_id': meeting.event_id,
                    'recruiter_id': meeting.recruiter_id,
                    'subject': meeting.subject,
                    'start_date': meeting.start_date.strftime('%Y-%m-%d'),
                    'start_time': meeting.start_time.strftime('%H:%M:%S'),
                    'end_date': meeting.end_date.strftime('%Y-%m-%d'),
                    'end_time': meeting.end_time.strftime('%H:%M:%S'),
                    'email': meeting.recruiter_email,
                    'attendees': meeting.attendees,  # Already stored as a comma-separated string
                   'cc_recipients': meeting.cc_recipients,  # Already stored as a comma-separated string
                    'time_zone': meeting.time_zone,
                    'join_url': meeting.join_url
                }
                meetings_data.append(meeting_dict)
        
        if not meetings_data:
            return jsonify({'message': 'No relevant meetings found for this recruiter'}), 404
        
        return jsonify({'meetings': meetings_data}), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# @app.route('/get_all_meetings', methods=['POST'])
# def get_all_meetings():
#     data = request.json
#     recruiter_id = data.get('recruiter_id')
    
#     if not recruiter_id:
#         return jsonify({'error': 'Missing recruiter_id field'}), 400
    
#     try:
#         # Get all users associated with the recruiter_id # L1 - SELECTED					

#         users = User.query.filter_by(id=recruiter_id).all()
#         user_emails = {user.email for user in users}

#         # Query the ScheduledMeeting table for all records
#         all_meetings = ScheduledMeeting.query.all()

#         # Filter meetings based on valid emails
#         meetings_data = []
#         for meeting in all_meetings:
#             attendees = meeting.attendees.split(',') if meeting.attendees else []
#             cc_recipients = meeting.cc_recipients.split(',') if meeting.cc_recipients else []

#             # Check if any attendees or CC recipients are in the list of user emails
#             valid_attendees = [email for email in attendees if email in user_emails]
#             valid_cc_recipients = [email for email in cc_recipients if email in user_emails]

#             # Check if recruiter_email is a valid user email
#             valid_recruiter_email = meeting.recruiter_email in user_emails

#             # Include the meeting if recruiter_email is valid, or if there are valid attendees or cc_recipients
#             if valid_recruiter_email or valid_attendees or valid_cc_recipients:
#                 meeting_dict = {
#                     'meeting_id': meeting.id,
#                     'event_id': meeting.event_id,
#                     'recruiter_id': meeting.recruiter_id,
#                     'subject': meeting.subject,
#                     'start_date': meeting.start_date.strftime('%Y-%m-%d'),
#                     'start_time': meeting.start_time.strftime('%H:%M:%S'),
#                     'end_date': meeting.end_date.strftime('%Y-%m-%d'),
#                     'end_time': meeting.end_time.strftime('%H:%M:%S'),
#                     'attendees': attendees,
#                     'cc_recipients': cc_recipients,
#                     'recruiter_email': meeting.recruiter_email,
#                     'time_zone': meeting.time_zone,
#                     'join_url': meeting.join_url
#                 }
#                 meetings_data.append(meeting_dict)
        
#         if not meetings_data:
#             return jsonify({'message': 'No relevant meetings found for this recruiter'}), 404
        
#         return jsonify({'meetings': meetings_data}), 200
    
#     except Exception as e:
#         return jsonify({'error': str(e)}), 500





# @app.route('/get_all_meetings', methods=['POST'])
# def get_all_meetings():
#     data = request.json
#     recruiter_id = data.get('recruiter_id')
    
#     if not recruiter_id:
#         return jsonify({'error': 'Missing recruiter_id field'}), 400
    
#     try:
#         # Get all users associated with the recruiter_id
#         users = User.query.filter_by(id=recruiter_id).all()
#         user_emails = {user.email for user in users}

#         # Query the ScheduledMeeting table for all records
#         all_meetings = ScheduledMeeting.query.all()

#         # Filter meetings based on valid emails
#         meetings_data = []
#         for meeting in all_meetings:
#             attendees = meeting.attendees.split(',') if meeting.attendees else []
#             cc_recipients = meeting.cc_recipients.split(',') if meeting.cc_recipients else []

#             # Check if any attendees or CC recipients are in the list of user emails
#             valid_attendees = [email for email in attendees if email in user_emails]
#             valid_cc_recipients = [email for email in cc_recipients if email in user_emails]

#             # Check if recruiter_email is a valid user email
#             valid_recruiter_email = meeting.recruiter_email in user_emails

#             # Include the meeting if recruiter_email is valid, or if there are valid attendees or cc_recipients
#             if valid_recruiter_email or valid_attendees or valid_cc_recipients:
#                 meeting_dict = {
#                     'meeting_id': meeting.id,
#                     'event_id': meeting.event_id,
#                     'recruiter_id': meeting.recruiter_id,
#                     'subject': meeting.subject,
#                     'start_date': meeting.start_date.strftime('%Y-%m-%d'),
#                     'start_time': meeting.start_time.strftime('%H:%M:%S'),
#                     'end_date': meeting.end_date.strftime('%Y-%m-%d'),
#                     'end_time': meeting.end_time.strftime('%H:%M:%S'),
#                     'attendees': valid_attendees,
#                     'cc_recipients': valid_cc_recipients,
#                     'recruiter_email': meeting.recruiter_email,
#                     'time_zone': meeting.time_zone,
#                     'join_url': meeting.join_url
#                 }
#                 meetings_data.append(meeting_dict)
        
#         if not meetings_data:
#             return jsonify({'message': 'No relevant meetings found for this recruiter'}), 404
        
#         return jsonify({'meetings': meetings_data}), 200
    
#     except Exception as e:
#         return jsonify({'error': str(e)}), 500



# @app.route('/get_all_meetings', methods=['POST'])
# def get_all_meetings():
#     data = request.json
#     recruiter_id = data.get('recruiter_id')
    
#     if not recruiter_id:
#         return jsonify({'error': 'Missing recruiter_id field'}), 400
    
#     try:
#         # Query the ScheduledMeeting table for all records matching the recruiter_id
#         meetings = ScheduledMeeting.query.filter_by(recruiter_id=recruiter_id).all()
        
#         if not meetings:
#             return jsonify({'message': 'No meetings found for this recruiter'})
        
#         # Convert the result to a list of dictionaries
#         meetings_data = []
#         for meeting in meetings:
#             meeting_dict = {
#                 'meeting_id': meeting.id,
#                 'event_id': meeting.event_id,
#                 'recruiter_id': meeting.recruiter_id,
#                 'subject': meeting.subject,
#                 'start_date': meeting.start_date.strftime('%Y-%m-%d'),  # Format date
#                 'start_time': meeting.start_time.strftime('%H:%M:%S'),  # Format time

#                 'end_date': meeting.end_date.strftime('%Y-%m-%d'), 
#                 'end_time': meeting.end_time.strftime('%H:%M:%S'),
                
#                 'attendees': meeting.attendees,  # Already stored as a comma-separated string
#                 'cc_recipients': meeting.cc_recipients,  # Already stored as a comma-separated string
#                 'recruiter_email': meeting.recruiter_email,
#                 'time_zone': meeting.time_zone,
#                 'join_url': meeting.join_url
#             }
#             meetings_data.append(meeting_dict)
        
#         return jsonify({'meetings': meetings_data}), 200
    
#     except Exception as e:
#         return jsonify({'error': str(e)}), 500


def get_calendar_events(access_token, recruiter_email):
    headers = {'Authorization': f'Bearer {access_token}'}
    response = requests.get(
        f'https://graph.microsoft.com/v1.0/users/{recruiter_email}/events',
        headers=headers
    )

    if response.status_code == 200:
        events = response.json().get('value', [])
        # Extract the 'id' (event_id) from each event
        event_ids = [event['id'] for event in events]
        return event_ids
    else:
        return []

def delete_events_not_in_graph(recruiter_email, graph_event_ids):
    # Query to get all event IDs from the database for the given recruiter_email
    db_event_ids = db.session.query(ScheduledMeeting.event_id).filter_by(recruiter_email=recruiter_email).all()
    db_event_ids = [id[0] for id in db_event_ids]  # Convert list of tuples to list of ids

    # Identify events that are in the database but not in the Graph API response
    events_to_delete = set(db_event_ids) - set(graph_event_ids)

    # Delete these events from the database
    if events_to_delete:
        db.session.query(ScheduledMeeting).filter(ScheduledMeeting.event_id.in_(events_to_delete)).delete(synchronize_session=False)
        db.session.commit()

@app.route('/sync_events', methods=['POST'])
def sync_events():
    recruiter_email = request.json.get('recruiter_email')
    
    if not recruiter_email:
        return jsonify({"error": "Missing recruiter_email in the request"}), 400
    
    access_token = get_access_token()
    graph_event_ids = get_calendar_events(access_token, recruiter_email)
    delete_events_not_in_graph(recruiter_email, graph_event_ids)
    
    return jsonify({"message": "Events synced successfully"}), 200







# Your app credentials
# Your app credentials
# TENANT_ID = '8a7c6498-6635-4dbc-8a5e-f38efccfef3e'
# CLIENT_ID = '7ba39e41-0ec7-411d-8649-6607574db5f9'
# CLIENT_SECRET = '2ET8Q~TkfssfstnbCWmFP2U24phkklo_w080uc7E'
# AUTHORITY = f'https://login.microsoftonline.com/{TENANT_ID}'
# SCOPES = ['https://graph.microsoft.com/.default']
# GRAPH_ENDPOINT = 'https://graph.microsoft.com/v1.0'

# # Initialize MSAL ConfidentialClientApplication
# msal_app = ConfidentialClientApplication(
#     CLIENT_ID,
#     authority=AUTHORITY,
#     client_credential=CLIENT_SECRET,
# )

# # Function to acquire access token
# def get_access_token():
#     result = msal_app.acquire_token_for_client(scopes=SCOPES)
#     if "access_token" in result:
#         return result["access_token"]
#     else:
#         print("Failed to obtain access token:", result.get("error_description"))
#         return None

# def create_event(subject, start_date, start_time, end_date, end_time, attendees, recruiter_email, time_zone):
#     access_token = get_access_token()
#     # print("access_token : ",access_token)
#     #access_token = "eyJ0eXAiOiJKV1QiLCJub25jZSI6Ik84c2dMcENrUERqQW5yNHNkUGZCeHZDSmF6MURwWncxZWJPSTEtUUgtMG8iLCJhbGciOiJSUzI1NiIsIng1dCI6IktRMnRBY3JFN2xCYVZWR0JtYzVGb2JnZEpvNCIsImtpZCI6IktRMnRBY3JFN2xCYVZWR0JtYzVGb2JnZEpvNCJ9.eyJhdWQiOiJodHRwczovL2dyYXBoLm1pY3Jvc29mdC5jb20iLCJpc3MiOiJodHRwczovL3N0cy53aW5kb3dzLm5ldC84YTdjNjQ5OC02NjM1LTRkYmMtOGE1ZS1mMzhlZmNjZmVmM2UvIiwiaWF0IjoxNzI0MDQ3Mzg5LCJuYmYiOjE3MjQwNDczODksImV4cCI6MTcyNDA1MTI4OSwiYWlvIjoiRTJkZ1lQRHVWVnFrdEQ2Rk1kSEZwL2pIREYweEFBPT0iLCJhcHBfZGlzcGxheW5hbWUiOiJBVFMiLCJhcHBpZCI6IjdiYTM5ZTQxLTBlYzctNDExZC04NjQ5LTY2MDc1NzRkYjVmOSIsImFwcGlkYWNyIjoiMSIsImlkcCI6Imh0dHBzOi8vc3RzLndpbmRvd3MubmV0LzhhN2M2NDk4LTY2MzUtNGRiYy04YTVlLWYzOGVmY2NmZWYzZS8iLCJpZHR5cCI6ImFwcCIsIm9pZCI6IjE3YTE2ZGJhLWEwMjMtNDE0Ny04ZWU1LTRiNGU5ZTZjMGRhNiIsInJoIjoiMC5BU3NBbUdSOGlqVm12RTJLWHZPT19NX3ZQZ01BQUFBQUFBQUF3QUFBQUFBQUFBQXJBQUEuIiwic3ViIjoiMTdhMTZkYmEtYTAyMy00MTQ3LThlZTUtNGI0ZTllNmMwZGE2IiwidGVuYW50X3JlZ2lvbl9zY29wZSI6IkFTIiwidGlkIjoiOGE3YzY0OTgtNjYzNS00ZGJjLThhNWUtZjM4ZWZjY2ZlZjNlIiwidXRpIjoiRXY3OFdJdXVvMENCSnEzNUNoMVFBQSIsInZlciI6IjEuMCIsIndpZHMiOlsiMDk5N2ExZDAtMGQxZC00YWNiLWI0MDgtZDVjYTczMTIxZTkwIl0sInhtc19pZHJlbCI6IjI0IDciLCJ4bXNfdGNkdCI6MTUzODcxMzQ1Nn0.WRGMVDNuoPoh8ozfZQF-7gI1euO_1mysBr-66yFEejsyuYTGDT_AUAgBJfPRzH-obcvLK4g0UCPt9xtPhVEuaYwiTKCDkJyx4iJ5gku8IjA5K_JruYUWa5n9rMv_o5sJ21H3-kW1pepJxF3OhTHEdPSNjbI3pd9ItIgyw5sr3H055mBbH_7v8GNrFU18zQ8ZJiAodgqyWz-JSU4vdw8OCu3VDRrmcGiHfe65rFjyKr25uAE4Qug_kLdhGoODAXGbUco5ASyUiH8Rr7MflNpMkXSiRBq31mLb7Na-M0lDFLyxfiRnKTJT2QA0EyloQz4TSOOez9hOBJaA-KFH02-wkg"

#     if not access_token:
#         return None

#     headers = {
#         'Authorization': f'Bearer {access_token}',
#         'Content-Type': 'application/json'
#     }

#     # Combine date and time into ISO 8601 format without seconds
#     start_date_time = f"{start_date}T{start_time[:5]}"  # Removing seconds
#     end_date_time = f"{end_date}T{end_time[:5]}"        # Removing seconds

#     event = {
#         'subject': subject,
#         'start': {
#             'dateTime': start_date_time,
#             'timeZone': time_zone
#         },
#         'end': {
#             'dateTime': end_date_time,
#             'timeZone': time_zone
#         },
#         'attendees': [
#             {
#                 'emailAddress': {
#                     'address': attendee,
#                     'name': ''
#                 },
#                 'type': 'required'
#             } for attendee in attendees
#         ],
#         'isOnlineMeeting': True,
#         'onlineMeetingProvider': 'teamsForBusiness'
#     }

#     response = requests.post(
#         f'https://graph.microsoft.com/v1.0/users/{recruiter_email}/events',
#         headers=headers,
#         data=json.dumps(event)
#     )

#     if response.status_code != 201:
#         return None, f"Error creating event: {response.status_code} - {response.text}"
    
#     return response.json(), None

# @app.route('/create_event', methods=['POST'])
# def handle_create_event():
#     data = request.json
    
#     if not data:
#         return jsonify({'error': 'Invalid request, no JSON body provided'}), 500
    
#     subject = data.get('subject')
#     start_date = data.get('start_date')
#     start_time = data.get('start_time')
#     end_date = data.get('end_date')
#     end_time = data.get('end_time')
#     attendees = data.get('attendees')
#     recruiter_email = data.get('recruiter_email')
#     time_zone = data.get('time_zone')  # Default to UTC if not provided
    
#     if not all([subject, start_date, start_time, end_date, end_time, attendees, recruiter_email]):
#         return jsonify({'error': 'Missing required fields'}), 500
    
#     event_response, error = create_event(
#         subject=subject,
#         start_date=start_date,
#         start_time=start_time,
#         end_date=end_date,
#         end_time=end_time,
#         attendees=attendees,
#         recruiter_email=recruiter_email,
#         time_zone=time_zone
#     )
    
#     if event_response:
#         return jsonify({'message': 'Event created successfully.', 'event': event_response}), 200
#     else:
#         return jsonify({'error': error}), 500



#############################################################################################################################

# from msal import ConfidentialClientApplication
# import requests
# from flask import Flask, request, jsonify


# # Constants
# CLIENT_ID = '7ba39e41-0ec7-411d-8649-6607574db5f9'
# TENANT_ID = '8a7c6498-6635-4dbc-8a5e-f38efccfef3e'
# CLIENT_SECRET = '2ET8Q~TkfssfstnbCWmFP2U24phkklo_w080uc7E'
# GRAPH_API_ENDPOINT = 'https://graph.microsoft.com/v1.0'

# # Function to generate access token
# def get_access_token():
#     token_response = msal_app.acquire_token_for_client(scopes=SCOPES)
#     if 'access_token' in token_response:
#         return token_response['access_token']
#     else:
#         print(f"Failed to get access token: {token_response.get('error_description')}")
#         return None


# # Function to construct event details
# def construct_event_detail(subject, start_date, end_date, from_time, end_time, attendees_list):
#     return {
#         "subject": subject,
#         "start": {
#             "dateTime": f"{start_date}T{from_time}",
#             "timeZone": "Asia/Kolkata"  # Use the appropriate timezone
#         },
#         "end": {
#             "dateTime": f"{end_date}T{end_time}",
#             "timeZone": "Asia/Kolkata"  # Use the appropriate timezone
#         },
#         "attendees": attendees_list
#     }

# @app.route('/schedule-event', methods=['POST'])
# def schedule_event():
#     data = request.json
#     subject = data.get('subject')
#     start_date = data.get('startDate')
#     end_date = data.get('endDate')
#     from_time = data.get('from_time')
#     end_time = data.get('end_time')
#     attendees_ids = data.get('attendees_id')  # List of attendee IDs or a single ID
#     user_email = data.get('user_email')  # Email of the user to schedule the event for

#     # Ensure attendees_ids is always a list
#     if not isinstance(attendees_ids, list):
#         attendees_ids = [attendees_ids]

#     # Fetch details of attendees based on the list of IDs (replace this with actual ORM queries)
#     attendees_details = Candidate.query.filter(Candidate.id.in_(attendees_ids)).all()

#     # Create the Attendees list
#     attendees_list = [
#         {
#             "emailAddress": {
#                 "address": attendee.email,
#                 "name": attendee.name
#             },
#             "type": "required"
#         } for attendee in attendees_details
#     ]

#     # Generate access token
#     access_token = generate_access_token()
#     headers = {
#         'Authorization': f'Bearer {access_token}',
#         'Content-Type': 'application/json'
#     }

#     # Construct event details
#     event_payload = construct_event_detail(subject, start_date, end_date, from_time, end_time, attendees_list)

#     # Schedule the event in the user's calendar
#     response = requests.post(
#         f'{GRAPH_API_ENDPOINT}/users/{user_email}/events',
#         headers=headers,
#         json=event_payload
#     )

#     if response.status_code == 201:
#         return jsonify({"message": "Event scheduled successfully", "event_id": response.json()['id']}), 201
#     else:
#         return jsonify({"error": response.status_code, "message": response.text}), response.status_code



# class GraphClient:
#     def __init__(self, access_token):
#         self.access_token = access_token
#         self.base_url = 'https://graph.microsoft.com/v1.0'
    
#     def post(self, endpoint, **kwargs):
#         url = f'{self.base_url}{endpoint}'
#         headers = {
#             'Authorization': f'Bearer {self.access_token}',
#             'Content-Type': 'application/json'
#         }
#         response = requests.post(url, headers=headers, **kwargs)
#         return response


# CLIENT_ID = '7ba39e41-0ec7-411d-8649-6607574db5f9'
# TENANT_ID = '8a7c6498-6635-4dbc-8a5e-f38efccfef3e'
# CLIENT_SECRET = '2ET8Q~TkfssfstnbCWmFP2U24phkklo_w080uc7E'
# SCOPE = ['https://graph.microsoft.com/.default']

# def get_access_token():
#     app = ConfidentialClientApplication(
#         CLIENT_ID,
#         authority=f'https://login.microsoftonline.com/{TENANT_ID}',
#         client_credential=CLIENT_SECRET,
#     )

#     token_response = app.acquire_token_for_client(scopes=SCOPE)
#     access_token = token_response.get('access_token')

#     if not access_token:
#         print("Token response:", token_response)
#         raise Exception("Could not obtain access token")
    
#     return access_token

# @app.route('/schedule-event', methods=['POST'])
# def schedule_event():
#     data = request.json
#     subject = data.get('subject')
#     start_date = data.get('startDate')
#     end_date = data.get('endDate')
#     from_time = data.get('from_time')
#     end_time = data.get('end_time')
#     attendees_ids = data.get('attendees_id')  # List of attendee IDs or a single ID
#     user_email = data.get('user_email')  # Email of the user to schedule the event for

#     # Ensure attendees_ids is always a list
#     if not isinstance(attendees_ids, list):
#         attendees_ids = [attendees_ids]

#     # Fetch details of attendees based on the list of IDs
#     attendees_details = Candidate.query.filter(Candidate.id.in_(attendees_ids)).all()

#     token = get_access_token()
#     client = GraphClient(token)

#     # Create Attendees list
#     attendee_list = [
#         {
#             "emailAddress": {
#                 "address": attendee.email,
#                 "name": attendee.name
#             },
#             "type": "required"
#         } for attendee in attendees_details
#     ]

#     # Create the event request body
#     request_body = {
#         "subject": subject,
#         "start": {
#             "dateTime": f"{start_date}T{from_time}",
#             "timeZone": "Asia/Kolkata"  # IST time zone
#         },
#         "end": {
#             "dateTime": f"{end_date}T{end_time}",
#             "timeZone": "Asia/Kolkata"  # IST time zone
#         },
#         "attendees": attendee_list
#     }

#     # Create the event in the user's calendar
#     response = client.post(f'/users/{user_email}/events', json=request_body)

#     if response.status_code == 201:
#         return jsonify(response.json())
#     else:
#         return jsonify({
#             "error": response.status_code,
#             "message": response.text
#         })



#####################################################################################################

@app.route('/update_resume_present', methods=['POST'])
def update_resume_present():
    try:
        # Update all rows in the Candidate table to set resume_present to True
        Candidate.query.update({Candidate.resume_present: True})
        db.session.commit()
        return jsonify({"message": "All candidates' resume_present set to True"}), 200
    except Exception as e:
        db.session.rollback()  # Rollback in case of any error
        return jsonify({"error": str(e)}), 500




@app.route('/update_jd_present', methods=['POST'])
def update_jd_present():
    try:
        # Update jd_pdf_present based on the presence of jd_pdf for all jobs in a single query
        JobPost.query.update({JobPost.jd_pdf_present: JobPost.jd_pdf.isnot(None)})
        db.session.commit()
        return jsonify({"message": "All jobs' jd_pdf_present updated based on jd_pdf presence"}), 200
    except Exception as e:
        db.session.rollback()  # Rollback in case of any error
        return jsonify({"error": str(e)}), 500


###########################################################################################
################## Questions Generation from Gemini ##############################
################################################################
def extract_text_from_pdf(pdf_bytes):
    """Extracts text from a PDF file."""
    text = ""
    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        text += page.get_text()
    return text

def extract_text_from_docx(pdf_bytes):
    """Extracts text from a DOCX file."""
    doc = docx.Document(BytesIO(pdf_bytes))
    text = ''
    for para in doc.paragraphs:
        text += para.text
    return text

@app.route('/generate_questions', methods=['POST'])
def generate_questions():
    data = request.json
    user_id = data.get('user_id')
    recruiter_prompt = data.get('recruiter_prompt')
    pdf_base64 = data.get('resume')

    if not pdf_base64:
        return jsonify({"error": "Resume not provided or invalid"}), 400

    try:
        # Decode the base64 PDF file
        pdf_bytes = base64.b64decode(pdf_base64)
        doc_bytes = extract_text_from_docx(pdf_base64)
    except Exception as e:
        return jsonify({"error": "Failed to decode base64 PDF"}), 400

    try:
        # Extract text from the PDF
        pdf_text = extract_text_from_pdf(pdf_bytes)
        doc_bytes = extract_text_from_docx(doc_bytes)
        print("pdf_text:", pdf_text)
    except Exception as e:
        return jsonify({"error": "Failed to extract text from PDF"}), 500

    # Combine the extracted text with the prompt
    prompt = f"{recruiter_prompt}\n\n{pdf_text}"

    # Configure and use Generative AI (assuming `genai` library setup)
    
    api_key = config.api_key
    if api_key is None:
        raise ValueError("API_KEY environment variable not set")

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')

    try:
        # Use the extracted text to generate a response
        response = model.generate_content(prompt)
    except Exception as e:
        return jsonify({"error": "Failed to generate content using Generative AI"}), 500

    # Ensure response content is available
    response_text = getattr(response, 'text', '')
    print("response_text:", response_text)
    if not response_text:
        return jsonify({"error": "Empty response from Generative AI"}), 500

    # Split response into lines and structure according to headings
    response_lines = response_text.split('\n')

    structured_response = []
    current_section = None
    questions_dict = {}
    
    for line in response_lines:
        line = line.strip()
        if line.endswith(":"):
            # Start of a new section
            if current_section:
                structured_response.append({
                    "heading": current_section.strip("##").strip(),
                    "questions": questions_dict
                })
            heading = line[:-1]  # Remove the trailing colon
            current_section = heading
            questions_dict = {}
        elif current_section is not None and line:
            # Check if the line is a question or sub-heading
            if line.startswith("**") and line.endswith("**"):
                # print("line :",line)
                sub_heading = line[2:-2].strip()  # Remove the leading and trailing '**'
                questions_dict[sub_heading] = []
            elif questions_dict:
                # Add questions to the last sub-heading
                last_sub_heading = list(questions_dict.keys())[-1]
                questions_dict[last_sub_heading].append(line.strip())
    
    # Add the last section if it exists
    if current_section:
        structured_response.append({
            "heading": current_section.strip("##").strip(),
            "questions": questions_dict
        })

    # Prepare the response in JSON format
    response_data = {
        'user_id': user_id,
        'response': structured_response
    }

    return jsonify(response_data)


###############################################################################################

# Extract text from PDF
def extract_text_from_pdf(file_binary):
    text = ''
    pdf_reader = PdfFileReader(BytesIO(file_binary))
    for page_num in range(pdf_reader.numPages):
        text += pdf_reader.getPage(page_num).extract_text()
    return text

# Extract text from DOCX
def extract_text_from_docx(file_binary):
    doc = docx.Document(BytesIO(file_binary))
    text = ''
    for para in doc.paragraphs:
        text += para.text
    return text

# Extract phone number
def extract_phone_number(text):
    phone_regex = r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]'
    phone_matches = re.findall(phone_regex, text)
    return phone_matches[-1].strip() if phone_matches else "No phone number found"

# Extract name
def extract_name(text):
    lines = text.split('\n')
    name_words = []

    for line in lines[:5]:
        if re.search(r'\b(phone|email)\b', line, re.IGNORECASE):
            continue
        
        # Extract potential name words
        words = re.findall(r'\b[A-Za-z]+', line)
        name_words.extend(words)

        if len(name_words) >= 2:
            return ' '.join(name_words[:3]).rstrip('.,')

    return "No name found"

# Extract email
def extract_email(text):
    email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    email_matches = re.findall(email_regex, text)
    return email_matches[-1].rstrip('.,') if email_matches else "No email found"

# Extract text from resume
def extract_text_from_resume(file_binary):
    try:
        return extract_text_from_pdf(file_binary)
    except:
        pass
    try:
        return extract_text_from_docx(file_binary)
    except Exception:
        raise ValueError("Unsupported file format")

# Preprocess text
# def preprocess_text(text):
#     text = re.sub(r'\W+', ' ', text).lower().strip()
#     return text

# Get job details including Gemini sub-skills
def get_job_details(job_id):
    job_post = JobPost.query.filter_by(id=job_id).first()
    if not job_post:
        return None

    job_details = {
        'client': job_post.client,
        'detailed_jd': job_post.detailed_jd if job_post.detailed_jd else "",
        'skills': job_post.skills.split(', ') if job_post.skills else [],
        'experience_min': job_post.experience_min,
        'experience_max': job_post.experience_max
    }

    return job_details

# Generate sub-skills from Gemini
# def generate_sub_skills_from_gemini(skill):

#     api_key = config.api_key
#     genai.configure(api_key=api_key)
#     model = genai.GenerativeModel('gemini-1.5-flash')

#     # prompt = f"""
#     # Given the following skills: {skill}, provide 10 related  technical topics for each technical skill, do not include soft skills (such as communication, fast learner). Present the output in the following format:

#     # sub_categories = {{
#     #     'Skill1': ['Topic1', 'Topic2', 'Topic3', ...],
#     #     'Skill2': ['Topic1', 'Topic2', 'Topic3', ...],
#     #     ...
#     # }}
#     # """

#     # prompt = f"""
#     # With the following skills: {skill}, outline the prerequisite technical skills and related branches. Ensure that soft skills (such as communication, fast learner) are not included. Structure the output as specified:

#     # sub_categories = {{
#     #     'Skill1': ['Topic1', 'Topic2', 'Topic3', ...],
#     #     'Skill2': ['Topic1', 'Topic2', 'Topic3', ...],
#     #     ...
#     # }}
#     # """

#     prompt = f"""
#     With the following skills: {skill}, outline 30 prerequisite technical skills and related branches for each skill. Ensure that soft skills (such as communication, fast learner) are not included. Structure the output as specified:

#     sub_categories = {{
#         'Skill1': ['Topic1', 'Topic2', 'Topic3', ..., 'Topic30'],
#         'Skill2': ['Topic1', 'Topic2', 'Topic3', ..., 'Topic30'],
#         ...
#     }}
#     """


#     response = model.generate_content(prompt)
#     response_text = response.candidates[0].content.parts[0].text.strip()

#     match = re.search(r"sub_categories\s*=\s*({.*})", response_text, re.DOTALL)
#     if match:
#         response_dict_str = match.group(1).replace("'", "\"")
#         try:
#             response_dict = json.loads(response_dict_str)
#         except json.JSONDecodeError as e:
#             print(f"Error parsing response dictionary: {e}")
#             return {}

#         cleaned_dict = {key.strip(): [item.strip() for item in value] for key, value in response_dict.items()}
#         print("cleaned_dict : ",cleaned_dict)
#         return cleaned_dict
#     else:
#         return {}

def generate_sub_skills_from_gemini(skill):
    # Configure the model
    api_key = config.api_key
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')

    # Construct the prompt
    prompt = f"""
    With the following skill: {skill}, outline 30 prerequisite technical skills and related branches. Ensure that soft skills (such as communication, fast learner) are not included. Structure the output as specified:

    sub_categories = {{
        '{skill}': ['Topic1', 'Topic2', 'Topic3', ..., 'Topic30']
    }}
    """

    try:
        # Generate content from the model
        response = model.generate_content(prompt)
        response_text = response.candidates[0].content.parts[0].text.strip()

        # Improved regex to match and parse JSON-like structures
        match = re.search(r"sub_categories\s*=\s*({.*})", response_text, re.DOTALL)
        if match:
            response_dict_str = match.group(1).replace("'", "\"")
            response_dict = json.loads(response_dict_str)
            cleaned_dict = {key.strip(): [item.strip() for item in value] for key, value in response_dict.items()}
            return cleaned_dict

    except Exception as e:
        print(f"Error generating sub-skills for {skill}: {e}")

    return {}


# Extract skills from resume text
# def extract_skills_from_resume(resume_text, gemini_sub_skills):
#     resume_skills = set()
#     resume_text = preprocess_text(resume_text)
#     for skill, sub_skills in gemini_sub_skills.items():
#         if skill.lower() in resume_text:
#             resume_skills.add(skill.lower())
#         for sub_skill in sub_skills:
#             if sub_skill.lower() in resume_text:
#                 resume_skills.add(skill.lower())
#                 break
#     return list(resume_skills)


def preprocess_text(text):
    # Convert to lowercase
    text = text.lower()
    # # Remove special characters and numbers
    # text = re.sub(r'\W+', ' ', text)
    # # Remove extra whitespace
    # text = re.sub(r'\s+', ' ', text).strip()
    return text




def extract_skills_from_gemini_resume(resume_text, gemini_sub_skills):
    resume_skills = set()
    resume_text = preprocess_text(resume_text)
    
    for skill, sub_skills in gemini_sub_skills.items():
        # Preprocess each subskill
        subskills_list = [preprocess_text(subskill) for subskill in sub_skills]
        
        # Create a pattern that matches any of the subskills within a broader description
        combined_pattern = r'\b(?:' + '|'.join(re.escape(subskill) for subskill in subskills_list) + r')\b'
        
        if re.search(combined_pattern, resume_text):
            resume_skills.add(skill)

    return list(resume_skills)



# Calculate skill match percentage
def calculate_skill_match_percentage(matched_skills, skills):
    if not skills:
        return 0
    return (len(matched_skills) / len(skills)) * 100



def extract_experience_from_resume(resume_text):
    experience_patterns = [
        r'(\d+(\.\d+)?)\s*(?:year|yr|years|yrs)?\s*(\d+)?\s*(?:month|months|mo|mos)?',
        r'(\d+)\s*(?:year|yr|years|yrs)?\s*(\d+)?\s*months?',
        r'(\d+(\.\d+)?)\s*(?:year|years|yr|yrs)?',
        r'(\d+(\.\d+)?)\s*(?:-year)',
        r'Demonstrated\s+(\d+(\.\d+)?)\s*(?:year|years|yr|yrs)',
        r'(\d+(\.\d+)?)\s*(?:years|yr|yrs)?\s*of\s*experience',
        r'(\d+(\.\d+)?)\s*years?\s*of\s*(?:IT|technical)?\s*experience',
        r'(\d+)\s*-\s*(\d+)\s*(?:years|yrs|yr|months|mos|mo)?\s*experience',
        r'experience\s*of\s*(\d+(\.\d+)?)\s*(?:years|yrs|yr|months|mos|mo)',
        r'over\s*(\d+(\.\d+)?)\s*(?:years|yrs|yr|months|mos|mo)?\s*experience',
        r'(?:years|yrs|yr|months|mos|mo)\s*of\s*experience\s*(\d+(\.\d+)?)',
        r'(\d+)\s*(?:years|yrs|yr|months|mos|mo)\s*\(to\s*date\)',
        r'total\s*experience\s*of\s*(\d+(\.\d+)?)\s*(?:years|yrs|yr|months|mos|mo)',
        r'experience\s*(?:from|since)\s*\d{4}\s*to\s*\d{4}',
        r'around\s*(\d+(\.\d+)?)\s*(?:years|yrs|yr|months|mos|mo)?\s*of\s*experience'
    ]

    internship_pattern = r'\binternship\b|\bintern\b'
    phone_pattern = re.compile(r'\b(\+?\d[\d\-\.\s]+)?\d{10}\b')

    internship_match = re.search(internship_pattern, resume_text, re.IGNORECASE)
    if internship_match:
        return 0

    resume_text_no_phone = phone_pattern.sub('', resume_text)
    for pattern in experience_patterns:
        match = re.search(pattern, resume_text_no_phone, re.IGNORECASE)
        if match:
            years = float(match.group(1))
            months = int(match.group(3) or 0) if len(match.groups()) > 2 else 0
            total_months = int(years * 12) + months
            if total_months < 600:
                return total_months

    return 0

# Merge periods
def merge_periods(periods):
    periods.sort(key=lambda x: x[0])
    merged = []
    for current in periods:
        if not merged:
            merged.append(current)
        else:
            last = merged[-1]
            if current[0] <= last[1]:
                merged[-1] = (last[0], max(last[1], current[1]))
            else:
                merged.append(current)
    return merged

@app.route('/check_resume_match', methods=['POST'])
def check_resume_match():
    data = request.json
    job_id = data.get('job_id')
    user_id = data.get('user_id')
    candidate_experience_str = data.get('candidate_experence')
    print("candidate_experience_str  :",candidate_experience_str)
    
    if 'resume' not in data:
        return jsonify({'error': 'No resume provided in the request'}), 400

    try:
        resume_binary = base64.b64decode(data['resume'])
    except Exception as e:
        return jsonify({'error': f'Error decoding base64 resume: {str(e)}'}), 400

    try:
        resume_text = extract_text_from_resume(resume_binary)
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    candidate_name = extract_name(resume_text)
    candidate_phone = extract_phone_number(resume_text)
    candidate_email = extract_email(resume_text)

    job_details = get_job_details(job_id)
    if not job_details:
        return jsonify({'error': 'Job details not found'}), 404

    gemini_sub_skills = {skill: generate_sub_skills_from_gemini(skill) for skill in job_details['skills']}
    matched_skills = extract_skills_from_gemini_resume(resume_text, gemini_sub_skills)
    
    # print("Resume Text:", resume_text)  # Debug
    # print("Gemini Sub Skills:", gemini_sub_skills)  # Debug
    # print("Matched Skills:", matched_skills)  # Debug

    skill_match_percentage = calculate_skill_match_percentage(matched_skills, job_details['skills'])

    

    # Convert candidate_experience_str to float or int if possible, or fallback to None
    if candidate_experience_str is not None and candidate_experience_str.strip():  # Check if not None and not empty
        try:
            candidate_experience = float(candidate_experience_str)
        except ValueError:
            candidate_experience = None
    else:
        candidate_experience = None
    
    # Calculate candidate experience only if candidate_experience is None or not provided
    if candidate_experience is None or candidate_experience == 0.0:
        print("resume_text :",resume_text)
        candidate_experience_months = extract_experience_from_resume(resume_text)
        candidate_experience_years = candidate_experience_months / 12
        candidate_experience_formatted = f"{math.floor(candidate_experience_years)}.{candidate_experience_months % 12}"
    else:
        candidate_experience_months = int(float(candidate_experience) * 12)  # Convert years to months
        candidate_experience_years = candidate_experience_months / 12
        candidate_experience_formatted = f"{math.floor(candidate_experience_years)}.{candidate_experience_months % 12}"

    experience_min_months = int(float(job_details['experience_min']) * 12)
    experience_max_months = int(float(job_details['experience_max']) * 12)
    if candidate_experience_months >= experience_min_months and candidate_experience_months <= experience_max_months:
        experience_match_percentage = 100
        experience_unmatch_percentage = 0
    elif candidate_experience_months > experience_max_months:
        experience_match_percentage = 100
        experience_unmatch_percentage = 0
    else:
        if candidate_experience_months < experience_min_months:
            experience_match_percentage = (candidate_experience_months / experience_min_months) * 100
            experience_unmatch_percentage = 100 - experience_match_percentage
        elif candidate_experience_months > experience_max_months:
            experience_match_percentage = (experience_max_months / candidate_experience_months) * 100
            experience_unmatch_percentage = 100 - experience_match_percentage

    overall_match_percentage = (skill_match_percentage + experience_match_percentage) / 2

    response_data = {
        'client': job_details['client'],
        'detailed_jd': job_details['detailed_jd'],
        'experience_min': job_details['experience_min'],
        'experience_max': job_details['experience_max'],
        'skills': job_details['skills'],
        'skill_match_percentage': skill_match_percentage,
        'experience_match_percentage': experience_match_percentage,
        'experience_unmatch_percentage': experience_unmatch_percentage,
        'overall_match_percentage': overall_match_percentage,
        'matched_skills': matched_skills,
        'resume_skills': matched_skills,
        'gemini_sub_skills': gemini_sub_skills,
        'user_id': user_id,
        'job_id': job_id,
        'candidate_experience': candidate_experience_formatted,
        'candidate_name': candidate_name,
        'candidate_phone': candidate_phone,
        'candidate_email': candidate_email
    }

    return jsonify(response_data), 200
############################################################################################################

# # Extract text from PDF
# def extract_text_from_pdf(file_binary):
#     text = ''
#     pdf_reader = PdfFileReader(BytesIO(file_binary))
#     for page_num in range(pdf_reader.numPages):
#         text += pdf_reader.getPage(page_num).extract_text()
#     return text

# # Extract text from DOCX
# def extract_text_from_docx(file_binary):
#     doc = docx.Document(BytesIO(file_binary))
#     text = ''
#     for para in doc.paragraphs:
#         text += para.text
#     return text

# # Extract phone number
# def extract_phone_number(text):
#     phone_regex = r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]'
#     phone_matches = re.findall(phone_regex, text)
#     return phone_matches[-1].strip() if phone_matches else "No phone number found"

# # Extract name
# def extract_name(text):
#     lines = text.split('\n')
#     name_words = []

#     for line in lines[:5]:
#         if re.search(r'\b(phone|email)\b', line, re.IGNORECASE):
#             continue
        
#         # Extract potential name words
#         words = re.findall(r'\b[A-Za-z]+', line)
#         name_words.extend(words)

#         if len(name_words) >= 2:
#             return ' '.join(name_words[:3]).rstrip('.,')

#     return "No name found"

# # Extract email
# def extract_email(text):
#     email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
#     email_matches = re.findall(email_regex, text)
#     return email_matches[-1].rstrip('.,') if email_matches else "No email found"

# # Extract text from resume
# def extract_text_from_resume(file_binary):
#     try:
#         return extract_text_from_pdf(file_binary)
#     except:
#         pass
#     try:
#         return extract_text_from_docx(file_binary)
#     except Exception:
#         raise ValueError("Unsupported file format")

# # Preprocess text
# # def preprocess_text(text):
# #     text = re.sub(r'\W+', ' ', text).lower().strip()
# #     return text

# # Get job details including Gemini sub-skills
# def get_job_details(job_id):
#     job_post = JobPost.query.filter_by(id=job_id).first()
#     if not job_post:
#         return None

#     job_details = {
#         'client': job_post.client,
#         'detailed_jd': job_post.detailed_jd if job_post.detailed_jd else "",
#         'skills': job_post.skills.split(', ') if job_post.skills else [],
#         'experience_min': job_post.experience_min,
#         'experience_max': job_post.experience_max
#     }

#     return job_details

# # Generate sub-skills from Gemini
# def generate_sub_skills_from_gemini(skill):
#     api_key = config.api_key
#     genai.configure(api_key=api_key)
#     model = genai.GenerativeModel('gemini-1.5-flash')

#     # prompt = f"""
#     # Given the following skills: {skill}, provide 10 related  technical topics for each technical skill, do not include soft skills (such as communication, fast learner). Present the output in the following format:

#     # sub_categories = {{
#     #     'Skill1': ['Topic1', 'Topic2', 'Topic3', ...],
#     #     'Skill2': ['Topic1', 'Topic2', 'Topic3', ...],
#     #     ...
#     # }}
#     # """

#     prompt = f"""
#     With the following skills: {skill}, outline the prerequisite technical skills and related branches. Ensure that soft skills (such as communication, fast learner) are not included. Structure the output as specified:

#     sub_categories = {{
#         'Skill1': ['Topic1', 'Topic2', 'Topic3', ...],
#         'Skill2': ['Topic1', 'Topic2', 'Topic3', ...],
#         ...
#     }}
#     """

#     response = model.generate_content(prompt)
#     response_text = response.candidates[0].content.parts[0].text.strip()

#     match = re.search(r"sub_categories\s*=\s*({.*})", response_text, re.DOTALL)
#     if match:
#         response_dict_str = match.group(1).replace("'", "\"")
#         try:
#             response_dict = json.loads(response_dict_str)
#         except json.JSONDecodeError as e:
#             print(f"Error parsing response dictionary: {e}")
#             return {}

#         cleaned_dict = {key.strip(): [item.strip() for item in value] for key, value in response_dict.items()}
#         print("cleaned_dict : ",cleaned_dict)
#         return cleaned_dict
#     else:
#         return {}

# # Extract skills from resume text
# # def extract_skills_from_resume(resume_text, gemini_sub_skills):
# #     resume_skills = set()
# #     resume_text = preprocess_text(resume_text)
# #     for skill, sub_skills in gemini_sub_skills.items():
# #         if skill.lower() in resume_text:
# #             resume_skills.add(skill.lower())
# #         for sub_skill in sub_skills:
# #             if sub_skill.lower() in resume_text:
# #                 resume_skills.add(skill.lower())
# #                 break
# #     return list(resume_skills)


# def preprocess_text(text):
#     # Convert to lowercase
#     text = text.lower()
#     # # Remove special characters and numbers
#     # text = re.sub(r'\W+', ' ', text)
#     # # Remove extra whitespace
#     # text = re.sub(r'\s+', ' ', text).strip()
#     return text




# def extract_skills_from_gemini_resume(resume_text, gemini_sub_skills):
#     resume_skills = set()
#     resume_text = preprocess_text(resume_text)
    
#     for skill, sub_skills in gemini_sub_skills.items():
#         # Preprocess each subskill
#         subskills_list = [preprocess_text(subskill) for subskill in sub_skills]
        
#         # Create a pattern that matches any of the subskills within a broader description
#         combined_pattern = r'\b(?:' + '|'.join(re.escape(subskill) for subskill in subskills_list) + r')\b'
        
#         if re.search(combined_pattern, resume_text):
#             resume_skills.add(skill)

#     return list(resume_skills)



# # Calculate skill match percentage
# def calculate_skill_match_percentage(matched_skills, skills):
#     if not skills:
#         return 0
#     return (len(matched_skills) / len(skills)) * 100



# def extract_experience_from_resume(resume_text):
#     experience_patterns = [
#         r'(\d+(\.\d+)?)\s*(?:year|yr|years|yrs)?\s*(\d+)?\s*(?:month|months|mo|mos)?',
#         r'(\d+)\s*(?:year|yr|years|yrs)?\s*(\d+)?\s*months?',
#         r'(\d+(\.\d+)?)\s*(?:year|years|yr|yrs)?',
#         r'(\d+(\.\d+)?)\s*(?:-year)',
#         r'Demonstrated\s+(\d+(\.\d+)?)\s*(?:year|years|yr|yrs)',
#         r'(\d+(\.\d+)?)\s*(?:years|yr|yrs)?\s*of\s*experience',
#         r'(\d+(\.\d+)?)\s*years?\s*of\s*(?:IT|technical)?\s*experience',
#         r'(\d+)\s*-\s*(\d+)\s*(?:years|yrs|yr|months|mos|mo)?\s*experience',
#         r'experience\s*of\s*(\d+(\.\d+)?)\s*(?:years|yrs|yr|months|mos|mo)',
#         r'over\s*(\d+(\.\d+)?)\s*(?:years|yrs|yr|months|mos|mo)?\s*experience',
#         r'(?:years|yrs|yr|months|mos|mo)\s*of\s*experience\s*(\d+(\.\d+)?)',
#         r'(\d+)\s*(?:years|yrs|yr|months|mos|mo)\s*\(to\s*date\)',
#         r'total\s*experience\s*of\s*(\d+(\.\d+)?)\s*(?:years|yrs|yr|months|mos|mo)',
#         r'experience\s*(?:from|since)\s*\d{4}\s*to\s*\d{4}',
#         r'around\s*(\d+(\.\d+)?)\s*(?:years|yrs|yr|months|mos|mo)?\s*of\s*experience'
#     ]

#     internship_pattern = r'\binternship\b|\bintern\b'
#     phone_pattern = re.compile(r'\b(\+?\d[\d\-\.\s]+)?\d{10}\b')

#     internship_match = re.search(internship_pattern, resume_text, re.IGNORECASE)
#     if internship_match:
#         return 0

#     resume_text_no_phone = phone_pattern.sub('', resume_text)
#     for pattern in experience_patterns:
#         match = re.search(pattern, resume_text_no_phone, re.IGNORECASE)
#         if match:
#             years = float(match.group(1))
#             months = int(match.group(3) or 0) if len(match.groups()) > 2 else 0
#             total_months = int(years * 12) + months
#             if total_months < 600:
#                 return total_months

#     return 0

# # Merge periods
# def merge_periods(periods):
#     periods.sort(key=lambda x: x[0])
#     merged = []
#     for current in periods:
#         if not merged:
#             merged.append(current)
#         else:
#             last = merged[-1]
#             if current[0] <= last[1]:
#                 merged[-1] = (last[0], max(last[1], current[1]))
#             else:
#                 merged.append(current)
#     return merged

# @app.route('/check_resume_match', methods=['POST'])
# def check_resume_match():
#     data = request.json
#     job_id = data.get('job_id')
#     user_id = data.get('user_id')
#     candidate_experience_str = data.get('candidate_experence')
#     print("candidate_experience_str  :",candidate_experience_str)
    
#     if 'resume' not in data:
#         return jsonify({'error': 'No resume provided in the request'}), 400

#     try:
#         resume_binary = base64.b64decode(data['resume'])
#     except Exception as e:
#         return jsonify({'error': f'Error decoding base64 resume: {str(e)}'}), 400

#     try:
#         resume_text = extract_text_from_resume(resume_binary)
#     except ValueError as e:
#         return jsonify({'error': str(e)}), 400

#     candidate_name = extract_name(resume_text)
#     candidate_phone = extract_phone_number(resume_text)
#     candidate_email = extract_email(resume_text)

#     job_details = get_job_details(job_id)
#     if not job_details:
#         return jsonify({'error': 'Job details not found'}), 404

#     gemini_sub_skills = {skill: generate_sub_skills_from_gemini(skill) for skill in job_details['skills']}
#     matched_skills = extract_skills_from_gemini_resume(resume_text, gemini_sub_skills)
    
#     # print("Resume Text:", resume_text)  # Debug
#     # print("Gemini Sub Skills:", gemini_sub_skills)  # Debug
#     # print("Matched Skills:", matched_skills)  # Debug

#     skill_match_percentage = calculate_skill_match_percentage(matched_skills, job_details['skills'])

    

#     # Convert candidate_experience_str to float or int if possible, or fallback to None
#     if candidate_experience_str is not None and candidate_experience_str.strip():  # Check if not None and not empty
#         try:
#             candidate_experience = float(candidate_experience_str)
#         except ValueError:
#             candidate_experience = None
#     else:
#         candidate_experience = None
    
#     # Calculate candidate experience only if candidate_experience is None or not provided
#     if candidate_experience is None or candidate_experience == 0.0:
#         print("resume_text :",resume_text)
#         candidate_experience_months = extract_experience_from_resume(resume_text)
#         candidate_experience_years = candidate_experience_months / 12
#         candidate_experience_formatted = f"{math.floor(candidate_experience_years)}.{candidate_experience_months % 12}"
#     else:
#         candidate_experience_months = int(float(candidate_experience) * 12)  # Convert years to months
#         candidate_experience_years = candidate_experience_months / 12
#         candidate_experience_formatted = f"{math.floor(candidate_experience_years)}.{candidate_experience_months % 12}"

#     experience_min_months = int(float(job_details['experience_min']) * 12)
#     experience_max_months = int(float(job_details['experience_max']) * 12)
#     if candidate_experience_months >= experience_min_months and candidate_experience_months <= experience_max_months:
#         experience_match_percentage = 100
#         experience_unmatch_percentage = 0
#     elif candidate_experience_months > experience_max_months:
#         experience_match_percentage = 100
#         experience_unmatch_percentage = 0
#     else:
#         if candidate_experience_months < experience_min_months:
#             experience_match_percentage = (candidate_experience_months / experience_min_months) * 100
#             experience_unmatch_percentage = 100 - experience_match_percentage
#         elif candidate_experience_months > experience_max_months:
#             experience_match_percentage = (experience_max_months / candidate_experience_months) * 100
#             experience_unmatch_percentage = 100 - experience_match_percentage

#     overall_match_percentage = (skill_match_percentage + experience_match_percentage) / 2

#     response_data = {
#         'client': job_details['client'],
#         'detailed_jd': job_details['detailed_jd'],
#         'experience_min': job_details['experience_min'],
#         'experience_max': job_details['experience_max'],
#         'skills': job_details['skills'],
#         'skill_match_percentage': skill_match_percentage,
#         'experience_match_percentage': experience_match_percentage,
#         'experience_unmatch_percentage': experience_unmatch_percentage,
#         'overall_match_percentage': overall_match_percentage,
#         'matched_skills': matched_skills,
#         'resume_skills': matched_skills,
#         'gemini_sub_skills': gemini_sub_skills,
#         'user_id': user_id,
#         'job_id': job_id,
#         'candidate_experience': candidate_experience_formatted,
#         'candidate_name': candidate_name,
#         'candidate_phone': candidate_phone,
#         'candidate_email': candidate_email
#     }

#     return jsonify(response_data), 200


#################################################################################################


# sub_categories = {
#     # Existing categories...
#     'Python': ['numpy', 'pandas', 'matplotlib', 'scikit-learn', 'tensorflow', 'keras'],
#     'Java': ['Spring Boot', 'J2EE', 'Hibernate', 'Servlets', 'JSP', 'Microservices'],
#     'C++': ['STL', 'OOP', 'Templates', 'Multi-threading', 'Boost'],
#     'AI': ['Machine Learning', 'Deep Learning', 'Natural Language Processing', 'Computer Vision'],
#     'JavaScript': ['React', 'Node.js', 'Angular', 'Vue.js', 'TypeScript', 'ES6'],
#     'HTML': ['HTML5', 'XML', 'JSON', 'Bootstrap'],
#     'CSS': ['CSS3', 'SASS', 'LESS'],
#     'Database Management': ['SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Oracle'],
#     'Web Development': ['RESTful APIs', 'GraphQL', 'WebSockets', 'Django', 'Flask'],
#     'DevOps': ['Docker', 'Kubernetes', 'Jenkins', 'Ansible', 'Terraform'],
#     'Cloud Computing': ['AWS', 'Azure', 'Google Cloud', 'Firebase'],
#     'Mobile Development': ['iOS Development', 'Android Development', 'React Native', 'Flutter'],
#     'Testing': ['Unit Testing', 'Integration Testing', 'Selenium', 'JUnit', 'Cypress'],
#     'Version Control': ['Git', 'SVN', 'Mercurial'],
#     'Agile Methodology': ['Scrum', 'Kanban', 'Agile Project Management'],
#     'Big Data': ['Hadoop', 'Spark', 'Hive', 'MapReduce', 'BigQuery'],
#     'Security': ['SSL/TLS', 'OWASP Top 10', 'Encryption', 'Firewalls', 'Penetration Testing'],
#     'UI/UX Design': ['Adobe XD', 'Sketch', 'Figma', 'User Interface Design', 'User Experience Design'],
#     'PHP': ['Laravel', 'Symfony', 'CodeIgniter', 'WordPress', 'Magento'],
#     'Ruby': ['Ruby on Rails', 'Sinatra', 'Rack', 'RSpec', 'Cucumber'],
#     'Go': ['Goroutines', 'Channels', 'RESTful APIs', 'Concurrency', 'Google App Engine'],
#     'Swift': ['iOS SDK', 'Cocoa Touch', 'SwiftUI', 'Combine', 'Realm'],
#     'Kotlin': ['Android SDK', 'Coroutines', 'Jetpack', 'Room', 'Ktor'],
#     'R': ['Data Visualization', 'Data Manipulation', 'Statistical Analysis', 'Shiny', 'ggplot2'],
#     'Perl': ['CGI', 'Mojolicious', 'DBI', 'Moose', 'Dancer'],
#     'Scala': ['Akka', 'Play Framework', 'Spark Streaming', 'Slick', 'ScalaTest'],
#     'TypeScript': ['Angular', 'React', 'Vue.js', 'Node.js', 'TypeORM'],
#     'Lua': ['Corona SDK', 'LÖVE', 'LuaJIT', 'Torch', 'MoonScript'],
#     'Shell Scripting': ['Bash', 'Shell Commands', 'Shell Variables', 'Script Automation', 'SED and AWK'],
#     'Assembly Language': ['x86 Assembly', 'ARM Assembly', 'MIPS Assembly', 'NASM', 'Keil'],
#     'MATLAB': ['Simulink', 'Image Processing Toolbox', 'Signal Processing Toolbox', 'Machine Learning Toolbox', 'Deep Learning Toolbox'],
#     'VB.NET': ['ASP.NET', 'Windows Forms', 'LINQ', 'ADO.NET', 'Entity Framework'],
#     'C#': ['.NET Framework', '.NET Core', 'ASP.NET Core', 'Unity', 'LINQ'],
#     'Dart': ['Flutter', 'Dart VM', 'AngularDart', 'Aqueduct', 'DartPad'],
#     'Julia': ['DataFrames', 'Plots', 'JuMP', 'Parallel Computing', 'Metaprogramming'],
#     'Haskell': ['GHC', 'Functional Programming', 'Monads', 'QuickCheck', 'Parsec'],
#     'Groovy': ['Grails', 'Spock Framework', 'Gradle', 'Geb', 'Ratpack'],
#     'Elixir': ['Phoenix Framework', 'Ecto', 'OTP', 'GenServer', 'Mix'],
#     'F#': ['.NET Framework', 'Fable', 'Functional Programming', 'Type Providers', 'Fable'],
#     'Rust': ['Ownership', 'Borrowing', 'Concurrency', 'Actix', 'Rocket'],
#     'Objective-C': ['Cocoa Touch', 'iOS Development', 'Swift', 'Core Data', 'UIKit'],
#     'COBOL': ['Mainframe', 'JCL', 'CICS', 'DB2', 'COBOL-IT'],
#     'Fortran': ['Numerical Computing', 'Parallel Programming', 'DO loops', 'Array Operations', 'GFortran'],
#     'Ada': ['Concurrency', 'Safety-Critical Systems', 'GNAT', 'SPARK', 'Real-Time Systems'],
#     'Scheme': ['Functional Programming', 'LISP', 'MIT/GNU Scheme', 'Racket', 'Guile'],
#     'Prolog': ['Logic Programming', 'Datalog', 'SWI-Prolog', 'GNU Prolog', 'CLP(FD)'],
#     'LISP': ['Common Lisp', 'Clojure', 'Scheme', 'Emacs Lisp', 'Racket'],
#     'Smalltalk': ['Object-Oriented Programming', 'Squeak', 'Pharo', 'Seaside', 'Morphic'],
#     'Erlang': ['OTP', 'Concurrency', 'Distributed Systems', 'Elixir', 'Riak'],
#     'D': ['Systems Programming', 'Garbage Collection', 'DUB', 'Vibe.d', 'GTKD'],
#     'PowerShell': ['Windows Administration', 'Scripting', 'Azure PowerShell', 'PowerShell Core', 'Active Directory'],
#     'Clojure': ['Functional Programming', 'Concurrency', 'LISP', 'ClojureScript', 'Datomic'],
#     'VHDL': ['FPGA', 'Digital Design', 'RTL Design', 'ModelSim', 'Vivado'],
#     'Verilog': ['ASIC', 'FPGA', 'SystemVerilog', 'Simulation', 'Synthesis'],

#     # Additional categories...
#     'Ruby on Rails': ['ActiveRecord', 'ActionView', 'ActionMailer', 'RSpec', 'Capistrano'],
#     'AngularJS': ['Controllers', 'Directives', 'Services', 'Filters', 'Routing'],
#     'Ember.js': ['Ember Data', 'Templates', 'Components', 'Controllers', 'Routing'],
#     'Backbone.js': ['Models', 'Views', 'Collections', 'Routers', 'Underscore.js'],
#     'Meteor.js': ['Meteor Methods', 'Blaze', 'Tracker', 'Accounts', 'Reactivity'],
#     'Polymer': ['Custom Elements', 'Templates', 'Data Binding', 'Events', 'Shadow DOM'],
#     'Golang': ['Concurrency', 'Channels', 'Interfaces', 'Benchmarking', 'WebAssembly'],
#     'Lua': ['Coroutines', 'Metatables', 'LuaSocket', 'Lapis', 'MoonScript'],
#     'Django': ['Models', 'Views', 'Forms', 'Templates', 'Admin'],
#     'Flask': ['Routes', 'Templates', 'Blueprints', 'SQLAlchemy', 'RESTful APIs'],
#     'Vue.js': ['Components', 'Directives', 'Vue Router', 'Vuex', 'Vue CLI'],
#     'Spring Framework': ['Dependency Injection', 'Spring MVC', 'Spring Security', 'Spring Data', 'Spring Boot'],
#     'Hibernate': ['Mapping', 'Transactions', 'Caching', 'Query Language', 'Entity Manager'],
#     'Node.js': ['Express.js', 'NPM', 'Event-driven Programming', 'RESTful APIs', 'Socket.io'],
#     'React Native': ['Components', 'Navigation', 'State Management', 'Hooks', 'Redux'],
#     'Flutter': ['Widgets', 'State Management', 'Material Design', 'Firebase Integration', 'Plugins'],
#     'ASP.NET': ['MVC', 'Web Forms', 'Entity Framework', 'Identity', 'Core'],
#     'Symfony': ['Bundles', 'Doctrine', 'Twig', 'Console', 'Security'],
#     'Laravel': ['Eloquent ORM', 'Blade Templates', 'Migrations', 'Events', 'Queues'],
#     'CodeIgniter': ['Controllers', 'Models', 'Views', 'Helpers', 'Form Validation'],
#     'Yii': ['Models', 'Controllers', 'Views', 'Widgets', 'Gii'],
#     'MVC Frameworks': ['Model', 'View', 'Controller', 'Routing', 'Templating'],
#     'Serverless Architecture': ['AWS Lambda', 'Azure Functions', 'Google Cloud Functions', 'Event-driven', 'NoOps'],
#     'Microservices Architecture': ['Service Discovery', 'API Gateway', 'Circuit Breaker', 'Event Sourcing', 'Service Mesh'],
#     'GraphQL': ['Schemas', 'Resolvers', 'Mutations', 'Subscriptions', 'Apollo Client'],
#     'RESTful APIs': ['HTTP Methods', 'Authentication', 'JSON', 'OpenAPI', 'Swagger'],
#     'SOAP': ['WSDL', 'UDDI', 'XML', 'WS-Security', 'Web Services'],
#     'Blockchain Development': ['Smart Contracts', 'Decentralized Apps', 'Consensus Algorithms', 'Cryptocurrencies', 'Hyperledger'],
#     'AR/VR Development': ['Unity3D', 'Unreal Engine', 'ARKit', 'ARCore', 'Vuforia'],
#     'Quantum Computing': ['Qubits', 'Quantum Gates', 'Quantum Algorithms', 'Quantum Teleportation', 'Superposition'],
#     'Low-Code Development': ['Visual Programming', 'Drag-and-drop', 'Workflow Automation', 'App Deployment', 'No-code Platforms'],
#     'PWA': ['Service Workers', 'App Shell', 'Push Notifications', 'Offline Support', 'Add to Home Screen'],
#     'NoSQL Databases': ['Document Store', 'Key-Value Store', 'Column-Family Store', 'Graph Store', 'Time Series Store'],
#     'Blockchain': ['Decentralized Ledger', 'Mining', 'Cryptocurrency', 'Smart Contracts', 'Distributed Consensus'],
#     'RPA (Robotic Process Automation)': ['Automation Anywhere', 'UiPath', 'Blue Prism', 'Process Mining', 'RPA Bots'],
#     'Cybersecurity': ['Ethical Hacking', 'Network Security', 'Cryptography', 'Incident Response', 'Security Operations'],
#     'Artificial Intelligence': ['Machine Learning', 'Neural Networks', 'Computer Vision', 'Natural Language Processing', 'AI Ethics'],
#     '5G Networks': ['Network Slicing', 'Edge Computing', 'Millimeter Wave', 'Virtual RAN', 'Beamforming'],
#     'Edge Computing': ['Fog Computing', 'Cloudlet', 'Mobile Edge Computing', 'Edge Analytics', 'Edge Security'],
#     'Quantum Cryptography': ['Quantum Key Distribution', 'Quantum Networks', 'Post-Quantum Cryptography', 'Quantum Hacking', 'Quantum Encryption'],
#     'Augmented Reality': ['Marker-based AR', 'Markerless AR', 'SLAM', 'AR Cloud', 'AR Glasses'],
#     'Virtual Reality': ['Immersive VR', '360-degree Video', 'VR Gaming', 'VR Training', 'VR Therapy'],
#     'Ethical Hacking': ['Footprinting', 'Scanning', 'Enumeration', 'Sniffing', 'Social Engineering'],
#     'Penetration Testing': ['Black Box Testing', 'White Box Testing', 'Gray Box Testing', 'Vulnerability Assessment', 'Exploitation'],
#     'Incident Response': ['Threat Detection', 'Forensics Analysis', 'Incident Containment', 'Root Cause Analysis', 'Remediation'],
#     'IoT Security': ['Device Authentication', 'Data Encryption', 'Firmware Updates', 'Access Control', 'Privacy Protection'],
#     'Cloud Security': ['Identity and Access Management', 'Data Encryption', 'Security Compliance', 'Threat Intelligence', 'Security Orchestration'],
#     'Mobile Security': ['App Hardening', 'Secure APIs', 'Jailbreaking Detection', 'Rooting Detection', 'Code Obfuscation'],
#     'AI Ethics': ['Fairness', 'Accountability', 'Transparency', 'Privacy', 'Bias'],
#     'Data Science': ['Statistical Analysis', 'Data Mining', 'Predictive Modeling', 'Data Visualization', 'Big Data'],
#     'Bioinformatics': ['Genomics', 'Proteomics', 'Computational Biology', 'Biostatistics', 'Bioinformatics Tools'],
#     'Health Informatics': ['Electronic Health Records', 'Health Data Analytics', 'Medical Imaging', 'Clinical Decision Support', 'Telemedicine'],
#     'Geographic Information Systems (GIS)': ['Spatial Analysis', 'Remote Sensing', 'GIS Software', 'Cartography', 'Geocoding'],
#     'Business Intelligence (BI)': ['Data Warehousing', 'Data Integration', 'OLAP', 'ETL Processes', 'Data Mining'],
#     'Financial Technology (FinTech)': ['Payments', 'Blockchain', 'Robo-advisors', 'InsurTech', 'RegTech'],
#     'Legal Tech': ['Legal Research', 'Contract Management', 'eDiscovery', 'Document Automation', 'Case Management'],
#     'EdTech': ['Learning Management Systems', 'Online Education', 'eLearning Platforms', 'Educational Technology Tools', 'MOOCs'],
#     'GovTech': ['Digital Transformation', 'Smart Cities', 'Civic Tech', 'Open Data', 'Government Analytics'],
#     'AgTech': ['Precision Agriculture', 'Farm Management Software', 'AgriTech Startups', 'Vertical Farming', 'AgriTech Solutions'],
#     'Space Tech': ['Satellite Technology', 'Space Exploration', 'Rocket Science', 'Spacecraft Design', 'Astronautics'],
#     'Green Tech': ['Renewable Energy', 'Sustainability', 'Environmental Monitoring', 'CleanTech', 'Green Building'],
#     'MarTech': ['Marketing Automation', 'Digital Marketing Analytics', 'Customer Relationship Management (CRM)', 'Content Management Systems', 'SEO'],
#     'InsurTech': ['Digital Insurance Platforms', 'Claims Processing Automation', 'Underwriting Automation', 'Insurance Analytics', 'InsurTech Startups'],
#     'Retail Tech': ['E-commerce Platforms', 'Point-of-Sale Systems', 'Supply Chain Management', 'Retail Analytics', 'Omni-channel Retailing'],
#     'HR Tech': ['Human Resource Information Systems (HRIS)', 'Talent Acquisition Software', 'Employee Engagement Platforms', 'Performance Management Systems', 'HR Analytics'],
#     'Legal Tech': ['Legal Research', 'Contract Management', 'eDiscovery', 'Document Automation', 'Case Management'],
#     'AdTech': ['Programmatic Advertising', 'Ad Exchange Platforms', 'Ad Fraud Detection', 'Ad Targeting', 'Data Management Platforms'],
#     'Logistics Tech': ['Supply Chain Optimization', 'Fleet Management Systems', 'Last-Mile Delivery Solutions', 'Warehouse Management Systems', 'Freight Forwarding Software'],
#     'Food Tech': ['Food Delivery Platforms', 'Restaurant Management Software', 'Food Safety Technology', 'AgriTech Solutions', 'Food Analytics'],
#     'Sports Tech': ['Sports Analytics', 'Fitness Tracking Apps', 'Athlete Performance Management', 'Sports Biomechanics', 'eSports Platforms'],
#     'Travel Tech': ['Online Travel Agencies', 'Travel Management Software', 'Booking Platforms', 'Travel Expense Management', 'Hospitality Tech'],
#     'Telecom Tech': ['5G Networks', 'Fiber Optics', 'Mobile Infrastructure', 'Telecom Software', 'Network Security'],
#     'Entertainment Tech': ['Streaming Platforms', 'Gaming Consoles', 'Music Streaming Services', 'Video Production Tools', 'Virtual Reality (VR)'],
#     'Media Tech': ['Digital Publishing Platforms', 'Content Management Systems', 'Media Analytics', 'Broadcasting Technology', 'Media Distribution'],
#     'Fashion Tech': ['Virtual Try-On', 'Fashion E-commerce Platforms', 'AR Fashion Apps', 'Fashion Data Analytics', 'Sustainable Fashion Tech'],
#     'Real Estate Tech': ['Property Management Software', 'Real Estate Marketplaces', 'VR Property Tours', 'Real Estate CRM', 'Automated Valuation Models'],
#     'Health Tech': ['Telemedicine Platforms', 'Health Data Analytics', 'Medical Imaging AI', 'Patient Management Software', 'Health Wearables'],
#     'Construction Tech': ['Building Information Modeling (BIM)', 'Construction Management Software', 'Drones in Construction', 'IoT in Construction', 'Prefabrication Technology'],
#     'Legal Tech': ['Legal Research', 'Contract Management', 'eDiscovery', 'Document Automation', 'Case Management'],
#     'Energy Tech': ['Smart Grids', 'Renewable Energy Technology', 'Energy Storage Solutions', 'Energy Management Systems', 'Microgrid Technology'],
#     'Agritech': ['Precision Agriculture', 'Farm Management Software', 'AgriTech Startups', 'Vertical Farming', 'AgriTech Solutions'],
#     'Govtech': ['Digital Transformation', 'Smart Cities', 'Civic Tech', 'Open Data', 'Government Analytics'],
#     'Retail Tech': ['E-commerce Platforms', 'Point-of-Sale Systems', 'Supply Chain Management', 'Retail Analytics', 'Omni-channel Retailing'],
#     'Legal Tech': ['Legal Research', 'Contract Management', 'eDiscovery', 'Document Automation', 'Case Management'],
#     'Travel Tech': ['Online Travel Agencies', 'Travel Management Software', 'Booking Platforms', 'Travel Expense Management', 'Hospitality Tech'],
#     'Telecom Tech': ['5G Networks', 'Fiber Optics', 'Mobile Infrastructure', 'Telecom Software', 'Network Security'],
#     'EdTech': ['Learning Management Systems', 'Online Education', 'eLearning Platforms', 'Educational Technology Tools', 'MOOCs'],
#     'Construction Tech': ['Building Information Modeling (BIM)', 'Construction Management Software', 'Drones in Construction', 'IoT in Construction', 'Prefabrication Technology'],
#     'Health Tech': ['Telemedicine Platforms', 'Health Data Analytics', 'Medical Imaging AI', 'Patient Management Software', 'Health Wearables'],
#     'FinTech': ['Payments', 'Blockchain', 'Robo-advisors', 'InsurTech', 'RegTech'],
#     'Fashion Tech': ['Virtual Try-On', 'Fashion E-commerce Platforms', 'AR Fashion Apps', 'Fashion Data Analytics', 'Sustainable Fashion Tech'],
#     'Energy Tech': ['Smart Grids', 'Renewable Energy Technology', 'Energy Storage Solutions', 'Energy Management Systems', 'Microgrid Technology'],
#     'Media Tech': ['Digital Publishing Platforms', 'Content Management Systems', 'Media Analytics', 'Broadcasting Technology', 'Media Distribution'],
#     'Food Tech': ['Food Delivery Platforms', 'Restaurant Management Software', 'Food Safety Technology', 'AgriTech Solutions', 'Food Analytics'],
#     'Automotive Tech': ['Autonomous Vehicles', 'Connected Cars', 'Electric Vehicle Technology', 'Vehicle Telematics', 'Automotive Cybersecurity'],
#     'Blockchain': ['Decentralized Ledger', 'Mining', 'Cryptocurrency', 'Smart Contracts', 'Distributed Consensus'],
#     'Cybersecurity': ['Ethical Hacking', 'Network Security', 'Cryptography', 'Incident Response', 'Security Operations'],
#     '5G Networks': ['Network Slicing', 'Edge Computing', 'Millimeter Wave', 'Virtual RAN', 'Beamforming'],
#     'Edge Computing': ['Fog Computing', 'Cloudlet', 'Mobile Edge Computing', 'Edge Analytics', 'Edge Security'],
#     'Quantum Cryptography': ['Quantum Key Distribution', 'Quantum Networks', 'Post-Quantum Cryptography', 'Quantum Hacking', 'Quantum Encryption'],
#     'Augmented Reality': ['Marker-based AR', 'Markerless AR', 'SLAM', 'AR Cloud', 'AR Glasses'],
#     'Virtual Reality': ['Immersive VR', '360-degree Video', 'VR Gaming', 'VR Training', 'VR Therapy'],
#     'Ethical Hacking': ['Footprinting', 'Scanning', 'Enumeration', 'Sniffing', 'Social Engineering'],
#     'Penetration Testing': ['Black Box Testing', 'White Box Testing', 'Gray Box Testing', 'Vulnerability Assessment', 'Exploitation'],
#     'Incident Response': ['Threat Detection', 'Forensics Analysis', 'Incident Containment', 'Root Cause Analysis', 'Remediation'],
#     'IoT Security': ['Device Authentication', 'Data Encryption', 'Firmware Updates', 'Access Control', 'Privacy Protection'],
#     'Cloud Security': ['Identity and Access Management', 'Data Encryption', 'Security Compliance', 'Threat Intelligence', 'Security Orchestration'],
#     'Mobile Security': ['App Hardening', 'Secure APIs', 'Jailbreaking Detection', 'Rooting Detection', 'Code Obfuscation'],
#     'AI Ethics': ['Fairness', 'Accountability', 'Transparency', 'Privacy', 'Bias'],
#     'Data Science': ['Statistical Analysis', 'Data Mining', 'Predictive Modeling', 'Data Visualization', 'Big Data'],
#     'Bioinformatics': ['Genomics', 'Proteomics', 'Computational Biology', 'Biostatistics', 'Bioinformatics Tools'],
#     'Health Informatics': ['Electronic Health Records', 'Health Data Analytics', 'Medical Imaging', 'Clinical Decision Support', 'Telemedicine'],
#     'Geographic Information Systems (GIS)': ['Spatial Analysis', 'Remote Sensing', 'GIS Software', 'Cartography', 'Geocoding'],
#     'Business Intelligence (BI)': ['Data Warehousing', 'Data Integration', 'OLAP', 'ETL Processes', 'Data Mining'],
#     'Financial Technology (FinTech)': ['Payments', 'Blockchain', 'Robo-advisors', 'InsurTech', 'RegTech'],
#     'Legal Tech': ['Legal Research', 'Contract Management', 'eDiscovery', 'Document Automation', 'Case Management'],
#     'EdTech': ['Learning Management Systems', 'Online Education', 'eLearning Platforms', 'Educational Technology Tools', 'MOOCs'],
#     'GovTech': ['Digital Transformation', 'Smart Cities', 'Civic Tech', 'Open Data', 'Government Analytics'],
#     'AgTech': ['Precision Agriculture', 'Farm Management Software', 'AgriTech Startups', 'Vertical Farming', 'AgriTech Solutions'],
#     'Space Tech': ['Satellite Technology', 'Space Exploration', 'Rocket Science', 'Spacecraft Design', 'Astronautics'],
#     'Green Tech': ['Renewable Energy', 'Sustainability', 'Environmental Monitoring', 'CleanTech', 'Green Building'],
#     'MarTech': ['Marketing Automation', 'Digital Marketing Analytics', 'Customer Relationship Management (CRM)', 'Content Management Systems', 'SEO'],
#     'InsurTech': ['Digital Insurance Platforms', 'Claims Processing Automation', 'Underwriting Automation', 'Insurance Analytics', 'InsurTech Startups'],
#     'Retail Tech': ['E-commerce Platforms', 'Point-of-Sale Systems', 'Supply Chain Management', 'Retail Analytics', 'Omni-channel Retailing'],
#     'HR Tech': ['Human Resource Information Systems (HRIS)', 'Talent Acquisition Software', 'Employee Engagement Platforms', 'Performance Management Systems', 'HR Analytics'],
#     'Legal Tech': ['Legal Research', 'Contract Management', 'eDiscovery', 'Document Automation', 'Case Management'],
#     'AdTech': ['Programmatic Advertising', 'Ad Exchange Platforms', 'Ad Fraud Detection', 'Ad Targeting', 'Data Management Platforms'],
#     'Logistics Tech': ['Supply Chain Optimization', 'Fleet Management Systems', 'Last-Mile Delivery Solutions', 'Warehouse Management Systems', 'Freight Forwarding Software'],
#     'Food Tech': ['Food Delivery Platforms', 'Restaurant Management Software', 'Food Safety Technology', 'AgriTech Solutions', 'Food Analytics'],
#     'Sports Tech': ['Sports Analytics', 'Fitness Tracking Apps', 'Athlete Performance Management', 'Sports Biomechanics', 'eSports Platforms'],
#     'Travel Tech': ['Online Travel Agencies', 'Travel Management Software', 'Booking Platforms', 'Travel Expense Management', 'Hospitality Tech'],
#     'Telecom Tech': ['5G Networks', 'Fiber Optics', 'Mobile Infrastructure', 'Telecom Software', 'Network Security'],
#     'Entertainment Tech': ['Streaming Platforms', 'Gaming Consoles', 'Music Streaming Services', 'Video Production Tools', 'Virtual Reality (VR)'],
#     'Media Tech': ['Digital Publishing Platforms', 'Content Management Systems', 'Media Analytics', 'Broadcasting Technology', 'Media Distribution'],
#     'Fashion Tech': ['Virtual Try-On', 'Fashion E-commerce Platforms', 'AR Fashion Apps', 'Fashion Data Analytics', 'Sustainable Fashion Tech'],
#     'Real Estate Tech': ['Property Management Software', 'Real Estate Marketplaces', 'VR Property Tours', 'Real Estate CRM', 'Automated Valuation Models'],
#     'Health Tech': ['Telemedicine Platforms', 'Health Data Analytics', 'Medical Imaging AI', 'Patient Management Software', 'Health Wearables'],
#     'Construction Tech': ['Building Information Modeling (BIM)', 'Construction Management Software', 'Drones in Construction', 'IoT in Construction', 'Prefabrication Technology'],
#     'Energy Tech': ['Smart Grids', 'Renewable Energy Technology', 'Energy Storage Solutions', 'Energy Management Systems', 'Microgrid Technology'],
#     'Agritech': ['Precision Agriculture', 'Farm Management Software', 'AgriTech Startups', 'Vertical Farming', 'AgriTech Solutions'],
#     'Govtech': ['Digital Transformation', 'Smart Cities', 'Civic Tech', 'Open Data', 'Government Analytics'],
#     'Retail Tech': ['E-commerce Platforms', 'Point-of-Sale Systems', 'Supply Chain Management', 'Retail Analytics', 'Omni-channel Retailing'],
#     'Legal Tech': ['Legal Research', 'Contract Management', 'eDiscovery', 'Document Automation', 'Case Management'],
#     'Travel Tech': ['Online Travel Agencies', 'Travel Management Software', 'Booking Platforms', 'Travel Expense Management', 'Hospitality Tech'],
#     'Telecom Tech': ['5G Networks', 'Fiber Optics', 'Mobile Infrastructure', 'Telecom Software', 'Network Security'],
#     'EdTech': ['Learning Management Systems', 'Online Education', 'eLearning Platforms', 'Educational Technology Tools', 'MOOCs'],
#     'Construction Tech': ['Building Information Modeling (BIM)', 'Construction Management Software', 'Drones in Construction', 'IoT in Construction', 'Prefabrication Technology'],
#     'Health Tech': ['Telemedicine Platforms', 'Health Data Analytics', 'Medical Imaging AI', 'Patient Management Software', 'Health Wearables'],
#     'FinTech': ['Payments', 'Blockchain', 'Robo-advisors', 'InsurTech', 'RegTech'],
#     'Fashion Tech': ['Virtual Try-On', 'Fashion E-commerce Platforms', 'AR Fashion Apps', 'Fashion Data Analytics', 'Sustainable Fashion Tech'],
#     'Energy Tech': ['Smart Grids', 'Renewable Energy Technology', 'Energy Storage Solutions', 'Energy Management Systems', 'Microgrid Technology'],
#     'Media Tech': ['Digital Publishing Platforms', 'Content Management Systems', 'Media Analytics', 'Broadcasting Technology', 'Media Distribution'],
#     'Food Tech': ['Food Delivery Platforms', 'Restaurant Management Software', 'Food Safety Technology', 'AgriTech Solutions', 'Food Analytics'],
#     'Automotive Tech': ['Autonomous Vehicles', 'Connected Cars', 'Electric Vehicle Technology', 'Vehicle Telematics', 'Automotive Cybersecurity'],
#     'Legal Tech': ['Legal Research', 'Contract Management', 'eDiscovery', 'Document Automation', 'Case Management'],
#     'Blockchain': ['Decentralized Ledger', 'Mining', 'Cryptocurrency', 'Smart Contracts', 'Distributed Consensus'],
#     'Cybersecurity': ['Ethical Hacking', 'Network Security', 'Cryptography', 'Incident Response', 'Security Operations'],
#     '5G Networks': ['Network Slicing', 'Edge Computing', 'Millimeter Wave', 'Virtual RAN', 'Beamforming'],
#     'Edge Computing': ['Fog Computing', 'Cloudlet', 'Mobile Edge Computing', 'Edge Analytics', 'Edge Security'],
#     'Quantum Cryptography': ['Quantum Key Distribution', 'Quantum Networks', 'Post-Quantum Cryptography', 'Quantum Hacking', 'Quantum Encryption'],
#     'Augmented Reality': ['Marker-based AR', 'Markerless AR', 'SLAM', 'AR Cloud', 'AR Glasses'],
#     'Virtual Reality': ['Immersive VR', '360-degree Video', 'VR Gaming', 'VR Training', 'VR Therapy'],
#     'Ethical Hacking': ['Footprinting', 'Scanning', 'Enumeration', 'Sniffing', 'Social Engineering'],
#     'Penetration Testing': ['Black Box Testing', 'White Box Testing', 'Gray Box Testing', 'Vulnerability Assessment', 'Exploitation'],
#     'Incident Response': ['Threat Detection', 'Forensics Analysis', 'Incident Containment', 'Root Cause Analysis', 'Remediation'],
#     'IoT Security': ['Device Authentication', 'Data Encryption', 'Firmware Updates', 'Access Control', 'Privacy Protection'],
#     'Cloud Security': ['Identity and Access Management', 'Data Encryption', 'Security Compliance', 'Threat Intelligence', 'Security Orchestration'],
#     'Mobile Security': ['App Hardening', 'Secure APIs', 'Jailbreaking Detection', 'Rooting Detection', 'Code Obfuscation'],
#     'AI Ethics': ['Fairness', 'Accountability', 'Transparency', 'Privacy', 'Bias'],
#     'Data Science': ['Statistical Analysis', 'Data Mining', 'Predictive Modeling', 'Data Visualization', 'Big Data'],
#     'Bioinformatics': ['Genomics', 'Proteomics', 'Computational Biology', 'Biostatistics', 'Bioinformatics Tools'],
#     'Health Informatics': ['Electronic Health Records', 'Health Data Analytics', 'Medical Imaging', 'Clinical Decision Support', 'Telemedicine'],
#     'Geographic Information Systems (GIS)': ['Spatial Analysis', 'Remote Sensing', 'GIS Software', 'Cartography', 'Geocoding'],
#     'Business Intelligence (BI)': ['Data Warehousing', 'Data Integration', 'OLAP', 'ETL Processes', 'Data Mining'],
#     'Financial Technology (FinTech)': ['Payments', 'Blockchain', 'Robo-advisors', 'InsurTech', 'RegTech'],
#     'Legal Tech': ['Legal Research', 'Contract Management', 'eDiscovery', 'Document Automation', 'Case Management'],
#     'EdTech': ['Learning Management Systems', 'Online Education', 'eLearning Platforms', 'Educational Technology Tools', 'MOOCs'],
#     'GovTech': ['Digital Transformation', 'Smart Cities', 'Civic Tech', 'Open Data', 'Government Analytics'],
#     'AgTech': ['Precision Agriculture', 'Farm Management Software', 'AgriTech Startups', 'Vertical Farming', 'AgriTech Solutions'],
#     'Space Tech': ['Satellite Technology', 'Space Exploration', 'Rocket Science', 'Spacecraft Design', 'Astronautics'],
#     'Green Tech': ['Renewable Energy', 'Sustainability', 'Environmental Monitoring', 'CleanTech', 'Green Building'],
#     'MarTech': ['Marketing Automation', 'Digital Marketing Analytics', 'Customer Relationship Management (CRM)', 'Content Management Systems', 'SEO'],
#     'InsurTech': ['Digital Insurance Platforms', 'Claims Processing Automation', 'Underwriting Automation', 'Insurance Analytics', 'InsurTech Startups'],
#     'Retail Tech': ['E-commerce Platforms', 'Point-of-Sale Systems', 'Supply Chain Management', 'Retail Analytics', 'Omni-channel Retailing'],
#     'HR Tech': ['Human Resource Information Systems (HRIS)', 'Talent Acquisition Software', 'Employee Engagement Platforms', 'Performance Management Systems', 'HR Analytics'],
#     'Legal Tech': ['Legal Research', 'Contract Management', 'eDiscovery', 'Document Automation', 'Case Management'],
#     'AdTech': ['Programmatic Advertising', 'Ad Exchange Platforms', 'Ad Fraud Detection', 'Ad Targeting', 'Data Management Platforms'],
#     'Logistics Tech': ['Supply Chain Optimization', 'Fleet Management Systems', 'Last-Mile Delivery Solutions', 'Warehouse Management Systems', 'Freight Forwarding Software'],
#     'Food Tech': ['Food Delivery Platforms', 'Restaurant Management Software', 'Food Safety Technology', 'AgriTech Solutions', 'Food Analytics'],
#     'Sports Tech': ['Sports Analytics', 'Fitness Tracking Apps', 'Athlete Performance Management', 'Sports Biomechanics', 'eSports Platforms'],
#     'Travel Tech': ['Online Travel Agencies', 'Travel Management Software', 'Booking Platforms', 'Travel Expense Management', 'Hospitality Tech'],
#     'Telecom Tech': ['5G Networks', 'Fiber Optics', 'Mobile Infrastructure', 'Telecom Software', 'Network Security'],
#     'Entertainment Tech': ['Streaming Platforms', 'Gaming Consoles', 'Music Streaming Services', 'Video Production Tools', 'Virtual Reality (VR)'],
#     'Media Tech': ['Digital Publishing Platforms', 'Content Management Systems', 'Media Analytics', 'Broadcasting Technology', 'Media Distribution'],
#     'Fashion Tech': ['Virtual Try-On', 'Fashion E-commerce Platforms', 'AR Fashion Apps', 'Fashion Data Analytics', 'Sustainable Fashion Tech'],
#     'Real Estate Tech': ['Property Management Software', 'Real Estate Marketplaces', 'VR Property Tours', 'Real Estate CRM', 'Automated Valuation Models'],
#     'Health Tech': ['Telemedicine Platforms', 'Health Data Analytics', 'Medical Imaging AI', 'Patient Management Software', 'Health Wearables'],
#     'Construction Tech': ['Building Information Modeling (BIM)', 'Construction Management Software', 'Drones in Construction', 'IoT in Construction', 'Prefabrication Technology'],
#     'Energy Tech': ['Smart Grids', 'Renewable Energy Technology', 'Energy Storage Solutions', 'Energy Management Systems', 'Microgrid Technology']}




# def extract_text_from_pdf(file_binary):
#     text = ''
#     pdf_reader = PyPDF2.PdfFileReader(BytesIO(file_binary))
#     for page_num in range(pdf_reader.numPages):
#         text += pdf_reader.getPage(page_num).extract_text()
#     return text

# def extract_text_from_docx(file_binary):
#     doc = docx.Document(BytesIO(file_binary))
#     text = ''
#     for para in doc.paragraphs:
#         text += para.text
#     return text


# def extract_phone_number(text):
#     phone_regex = r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]'
#     phone_matches = re.findall(phone_regex, text)
#     return phone_matches[-1].strip() if phone_matches else "No phone number found"

# def extract_name(text):
#     lines = text.split('\n')
#     name_words = []

#     for line in lines[:5]:
#         if re.search(r'\b(phone|email)\b', line, re.IGNORECASE):
#             continue
        
#         # Extract potential name words
#         words = re.findall(r'\b[A-Za-z]+', line)
#         name_words.extend(words)

#         if len(name_words) >= 2:
#             return ' '.join(name_words[:3]).rstrip('.,')

#     return "No name found"

# def extract_email(text):
#     email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
#     email_matches = re.findall(email_regex, text)
#     return email_matches[-1].rstrip('.,') if email_matches else "No email found"


# # New functions for extracting candidate details
# # def extract_phone_number(text):
# #     phone_regex = r'\b\d{10}\b'
# #     phone_matches = re.findall(phone_regex, text)
# #     return phone_matches[-1] if phone_matches else "No phone number found"

# # def extract_name(text):
# #     lines = text.split('\n')
# #     name_words = []  # List to store the words of the name
    
# #     # Regular expressions to identify lines that are likely contact details
# #     phone_pattern = re.compile(r'\b(\+?\d[\d\-\.\s]+)?\d{10}\b')
# #     email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
    
# #     for line in lines[:5]:  # Look at the first five lines where the name is likely to appear
# #         # Skip lines that are likely to be contact details
# #         if phone_pattern.search(line) or email_pattern.search(line):
# #             continue
        
# #         # Remove common salutations and titles
# #         cleaned_line = re.sub(r'\b(Mr\.|Mrs\.|Ms\.|Miss|Dr\.|Sir|Madam)\b', '', line, flags=re.IGNORECASE).strip()
        
# #         # Extract names with up to three words
# #         words = cleaned_line.split()
# #         name_words.extend(words)  # Add words from the current line to the list
        
# #         if len(name_words) <= 2:
# #             continue  # Continue accumulating words if we have less than or equal to three words
# #         else:
# #             # Stop accumulating if we exceed three words and return the concatenated name
# #             return ' '.join(word.capitalize() for word in name_words[:3]).rstrip('.,')
    
# #     # Return the concatenated name if found within the first five lines
# #     if name_words:
# #         return ' '.join(word.capitalize() for word in name_words[:3]).rstrip('.,')
    
# #     return "No name found"

# # def extract_email(text):
# #     email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
# #     email_matches = re.findall(email_regex, text)
# #     return email_matches[-1].rstrip('.,') if email_matches else "No email found"


# def extract_text_from_docx1(file_binary):
#     """
#     Extract text from a DOCX file.
    
#     Parameters:
#         file (BytesIO): DOCX file-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         doc = Document(file)
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + '\n'
#     except Exception as e:
#         print(f"Error extracting text from DOCX: {e}")
#     return text

# def extract_text_from_resume(file_binary):
#     try:
#         return extract_text_from_pdf(file_binary)
#     except PyPDF2.utils.PdfReadError:
#         pass
    
#     try:
#         return extract_text_from_docx(file_binary)
#     except Exception:
#         raise ValueError("Unsupported file format")

# def preprocess_text(text):
#     text = re.sub(r'\W+', ' ', text).lower().strip()
#     return text

# def generate_sub_skills(skill):
#     return sub_categories.get(skill.strip(), [])

# def combine_skills_with_sub_categories(skills):
#     combined_skills = set()
#     for skill in skills:
#         combined_skills.add(skill.strip())
#         combined_skills.update(generate_sub_skills(skill.strip()))  # Include sub-skills
#     return list(combined_skills)

# def extract_skills_from_resume(resume_text, skills_list):
#     resume_text = preprocess_text(resume_text)
#     skills_found = set()
    
#     for skill in skills_list:
#         skill_lower = skill.lower()
        
#         # Check direct match
#         if re.search(r'\b' + re.escape(skill_lower) + r'\b', resume_text, re.IGNORECASE):
#             skills_found.add(skill)
#             continue
        
#         # Check sub-categories
#         sub_skills = generate_sub_skills(skill)
#         for sub_skill in sub_skills:
#             sub_skill_lower = sub_skill.lower()
#             if re.search(r'\b' + re.escape(sub_skill_lower) + r'\b', resume_text, re.IGNORECASE):
#                 skills_found.add(sub_skill)
#                 break
        
#         # Check key match and add sub-category skills
#         if skill_lower in sub_categories:
#             skills_found.update(sub_categories[skill_lower])
    
#     return list(skills_found)



# def calculate_skill_match_percentage(matched_skills, skills):
#     if not skills:
#         return 0
    
#     if all(skill in matched_skills for skill in skills):
#         return 100.0
#     else:
#         return (len(matched_skills) / len(skills)) * 100

# # def extract_skills_from_resume(resume_text, skills_list):
# #     resume_text = preprocess_text(resume_text)
# #     skills_found = set()

# #     for skill in skills_list:
# #         skill_lower = skill.lower()
        
# #         # Check direct match for complete words with word boundaries and case insensitivity
# #         if re.search(r'\b' + re.escape(skill_lower) + r'\b', resume_text, re.IGNORECASE):
# #             skills_found.add(skill)
# #             continue
        
# #         # Check sub-categories
# #         sub_skills = generate_sub_skills(skill)
# #         for sub_skill in sub_skills:
# #             sub_skill_lower = sub_skill.lower()
# #             if re.search(r'\b' + re.escape(sub_skill_lower) + r'\b', resume_text, re.IGNORECASE):
# #                 skills_found.add(sub_skill)  # Add sub-skill instead of the main skill
# #                 break
        
# #         # Optionally, you can also check for exact matches without word boundaries as a fallback
# #         # if skill_lower in resume_text.lower():
# #         #     skills_found.add(skill)
    
# #     return list(skills_found)



# # def calculate_skill_match_percentage(matched_skills, skills):
# #     if not skills:
# #         return 0
    
# #     # Normalize skills to lowercase for comparison
# #     normalized_matched_skills = {skill.lower() for skill in matched_skills}
# #     normalized_skills = {skill.lower() for skill in skills}
    
# #     if normalized_skills.intersection(normalized_matched_skills) == normalized_skills:
# #         return 100.0
# #     else:
# #         return (len(normalized_matched_skills) / len(normalized_skills)) * 100



# def get_job_details(job_id):
#     # Assuming JobPost model and retrieval function as defined earlier
#     job_post = JobPost.query.filter_by(id=job_id).first()
#     if not job_post:
#         return None
#     job_details = {
#         'client': job_post.client,
#         'detailed_jd': job_post.detailed_jd if job_post.detailed_jd else "",
#         'skills': job_post.skills.split(', ') if job_post.skills else [],  # Split skills into a list
#         'experience_min': job_post.experience_min,
#         'experience_max': job_post.experience_max
#     }
#     return job_details


  

# def extract_experience_from_resume(resume_text):
#     # Define patterns to match experience in years and months
#     experience_pattern_1 = r'(\d+(\.\d+)?)\s*(?:year|yr|years|yrs)?\s*(\d+)?\s*(?:month|months|mo|mos)?'
#     experience_pattern_2 = r'(\d+)\s*(?:year|yr|years|yrs)?\s*(\d+)?\s*months?'
#     experience_pattern_3 = r'(\d+(\.\d+)?)\s*(?:year|years|yr|yrs)?'
#     experience_pattern_4 = r'(\d+(\.\d+)?)\s*(?:-year)'
#     experience_pattern_5 = r'Demonstrated\s+(\d+(\.\d+)?)\s*(?:year|years|yr|yrs)'
#     experience_pattern_6 = r'(\d+(\.\d+)?)\s*(?:years|yr|yrs)?\s*of\s*experience'  # New pattern for "7.5 years of experience"
#     experience_pattern_7 = r'(\d+(\.\d+)?)\s*years?\s*of\s*(?:IT|technical)?\s*experience'  # New pattern for "6.3 years of IT experience"

#     # Define a pattern to identify internship-related terms
#     internship_pattern = r'\binternship\b|\bintern\b'

#     # Define pattern to match phone numbers
#     phone_pattern = re.compile(r'\b(\+?\d[\d\-\.\s]+)?\d{10}\b')

#     # Search for the internship pattern in the resume text
#     internship_match = re.search(internship_pattern, resume_text, re.IGNORECASE)

#     # If internship pattern is found, return 0 indicating no experience
#     if internship_match:
#         return 0

#     # Exclude phone numbers from the experience search
#     resume_text_no_phone = phone_pattern.sub('', resume_text)

#     # Search for experience patterns
#     patterns = [
#         experience_pattern_1,
#         experience_pattern_2,
#         experience_pattern_3,
#         experience_pattern_4,
#         experience_pattern_5,
#         experience_pattern_6,  # Include the new pattern in the search
#         experience_pattern_7   # Include the new pattern for "6.3 years of IT experience"
#     ]
    
#     for pattern in patterns:
#         match = re.search(pattern, resume_text_no_phone, re.IGNORECASE)
#         if match:
#             years = float(match.group(1))
#             months = int(match.group(3) or 0) if len(match.groups()) > 2 else 0
#             total_months = int(years * 12) + months
#             if total_months < 600:  # Assuming 50 years of experience as a reasonable upper limit
#                 return total_months

#     return 0  # Return 0 if no valid experience pattern is found or exceeds reasonable limit

# def extract_experience_from_dates(resume_text):
#     current_year = datetime.now().year
    
#     date_pattern = r'(\b\d{4}\b)[\s-]*(Present|\b\d{4}\b)'
#     date_matches = re.findall(date_pattern, resume_text, re.IGNORECASE)
    
#     total_experience_months = 0
#     experience_periods = []

#     for start_year, end_year in date_matches:
#         try:
#             start_year = int(start_year)
#             end_year = current_year if end_year.lower() == 'present' else int(end_year)
            
#             if start_year <= end_year:
#                 experience_periods.append((start_year, end_year))
#         except ValueError:
#             continue  # Skip invalid date ranges

#     # Merge overlapping periods and sum the total experience in months
#     merged_periods = merge_periods(experience_periods)
#     total_experience_months = sum((end - start + 1) * 12 for start, end in merged_periods)

#     return total_experience_months

# def merge_periods(periods):
#     # Sort periods by start year
#     periods.sort(key=lambda x: x[0])
    
#     merged = []
#     for current in periods:
#         if not merged:
#             merged.append(current)
#         else:
#             last = merged[-1]
#             if current[0] <= last[1]:  # Overlapping or contiguous
#                 merged[-1] = (last[0], max(last[1], current[1]))
#             else:
#                 merged.append(current)
#     return merged

# @app.route('/check_resume_match', methods=['POST'])
# def check_resume_match():
#     data = request.json
#     job_id = data.get('job_id')
#     user_id = data.get('user_id')
#     candidate_experience_str = data.get('candidate_experence')  # Fetch candidate experience as a string
#     print("candidate_experience_str: ",candidate_experience_str)
#     if 'resume' not in data:
#         return jsonify({'error': 'No resume provided in the request'}), 400

#     try:
#         resume_binary = base64.b64decode(data['resume'])
#     except Exception as e:
#         return jsonify({'error': f'Error decoding base64 resume: {str(e)}'}), 400

#     try:
#         resume_text = extract_text_from_resume(resume_binary)
#         resume_text = extract_text_from_docx1(resume_binary)
#     except ValueError as e:
#         return jsonify({'error': str(e)}), 400

#     # Extract candidate details
#     candidate_name = extract_name(resume_text)
#     candidate_phone = extract_phone_number(resume_text)
#     candidate_email = extract_email(resume_text)

#     job_details = get_job_details(job_id)
#     if not job_details:
#         return jsonify({'error': 'Job details not found'}), 404
    
#     combined_skills = job_details['skills']
#     matched_skills = extract_skills_from_resume(resume_text, combined_skills)
#     skill_match_percentage = calculate_skill_match_percentage(matched_skills, job_details['skills'])

#     # Convert candidate_experience_str to float or int if possible, or fallback to None
#     if candidate_experience_str is not None and candidate_experience_str.strip():  # Check if not None and not empty
#         try:
#             candidate_experience = float(candidate_experience_str)
#         except ValueError:
#             candidate_experience = None
#     else:
#         candidate_experience = None
    
#     # Calculate candidate experience only if candidate_experience is None or not provided
#     if candidate_experience is None or candidate_experience == 0.0:
#         print("resume_text :",resume_text)
#         candidate_experience_months = extract_experience_from_resume(resume_text)
#         candidate_experience_years = candidate_experience_months / 12
#         candidate_experience_formatted = f"{math.floor(candidate_experience_years)}.{candidate_experience_months % 12}"
#     else:
#         candidate_experience_months = int(float(candidate_experience) * 12)  # Convert years to months
#         candidate_experience_years = candidate_experience_months / 12
#         candidate_experience_formatted = f"{math.floor(candidate_experience_years)}.{candidate_experience_months % 12}"

#     experience_min_months = int(float(job_details['experience_min']) * 12)
#     experience_max_months = int(float(job_details['experience_max']) * 12)

#     # if candidate_experience_months >= experience_min_months and candidate_experience_months <= experience_max_months:
#     #     experience_match_percentage = 100
#     #     experience_unmatch_percentage = 0
#     # else:
#     #     if candidate_experience_months < experience_min_months:
#     #         experience_match_percentage = (candidate_experience_months / experience_min_months) * 100
#     #         experience_unmatch_percentage = 100 - experience_match_percentage
#     #     elif candidate_experience_months > experience_max_months:
#     #         experience_match_percentage = (experience_max_months / candidate_experience_months) * 100
#     #         experience_unmatch_percentage = 100 - experience_match_percentage


#     # if candidate_experience_months >= experience_min_months and candidate_experience_months <= experience_max_months:
#     #     experience_match_percentage = 100
#     #     experience_unmatch_percentage = 0
#     # elif candidate_experience_months < experience_min_months:
#     #     experience_match_percentage = (candidate_experience_months / experience_min_months) * 100
#     #     experience_unmatch_percentage = 100 - experience_match_percentage
#     # else:  # candidate_experience_months > experience_max_months
#     #     # Adjust the calculation to properly reflect the mismatch percentage
#     #     experience_match_percentage = (experience_max_months / candidate_experience_months) * 100
#     #     experience_unmatch_percentage = 100 - experience_match_percentage
#     if candidate_experience_months >= experience_min_months and candidate_experience_months <= experience_max_months:
#         experience_match_percentage = 100
#         experience_unmatch_percentage = 0
#     elif candidate_experience_months > experience_max_months:
#         experience_match_percentage = 100
#         experience_unmatch_percentage = 0
#     else:
#         if candidate_experience_months < experience_min_months:
#             experience_match_percentage = (candidate_experience_months / experience_min_months) * 100
#             experience_unmatch_percentage = 100 - experience_match_percentage
#         elif candidate_experience_months > experience_max_months:
#             experience_match_percentage = (experience_max_months / candidate_experience_months) * 100
#             experience_unmatch_percentage = 100 - experience_match_percentage
#     overall_match_percentage = (skill_match_percentage + experience_match_percentage) / 2

#     response_data = {
#         'client': job_details['client'],
#         'detailed_jd': job_details['detailed_jd'],
#         'experience_min': job_details['experience_min'],
#         'experience_max': job_details['experience_max'],
#         'skills': job_details['skills'],
#         'skill_match_percentage': skill_match_percentage,
#         'experience_match_percentage': experience_match_percentage,
#         'experience_unmatch_percentage': experience_unmatch_percentage,
#         'overall_match_percentage': overall_match_percentage,
#         'matched_skills': matched_skills,
#         'resume_skills': matched_skills,
#         'user_id': user_id,
#         'job_id': job_id,
#         'candidate_experience': candidate_experience_formatted,
#         'candidate_name' : candidate_name,
#         'candidate_phone' : candidate_phone,
#         'candidate_email' : candidate_email
#     }

#     return jsonify(response_data), 200

####################################################################################


def extract_text_from_docx2(file):
    """
    Extract text from a DOCX file.
    
    Parameters:
        file (BytesIO): DOCX file-like object.
    
    Returns:
        str: Extracted text.
    """
    text = ""
    try:
        doc = Document(file)
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
    return text


def get_job_details_candidate(job_id):
    job_post = JobPost.query.filter_by(id=job_id).first()
    if not job_post:
        return None

    job_details = {
        'client': job_post.client,
        'detailed_jd': job_post.detailed_jd if job_post.detailed_jd else "",
        'skills': job_post.skills.split(', ') if job_post.skills else [],
        'experience_min': job_post.experience_min,
        'experience_max': job_post.experience_max,
        'budget_min': job_post.budget_min,
        'budget_max': job_post.budget_max,
        'role': job_post.role,
        'location': job_post.location
    }

    return job_details


def clean_response(text):
    text = re.sub(r'[*"#]', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    text = text.replace('}, {', '},\n{').replace('},\n{', '},\n{')
    text = text.replace('```python', '').replace('```', '')
    text = text.replace('json ', '')  # Remove 'json ' prefix if present
    return text

def format_analyze_candidate_profile(text):
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        response_dict_str = match.group(0).replace("'", "\"")
        try:
            response_dict = json.loads(response_dict_str)
        except json.JSONDecodeError as e:
            print(f"Error parsing response dictionary: {e}")
            return []

        # Extract lists from the dictionary
        skills_domains = response_dict.get("Skills/Domain", [])
        experiences = response_dict.get("Candidate Experience", [])
        relevance_scores = response_dict.get("Relevance Score", [])

        if len(skills_domains) != len(experiences) or len(skills_domains) != len(relevance_scores):
            print("Length mismatch between skills/domains, experiences, and relevance scores.")
            return []

        # Create a formatted result as an array of dictionaries
        result = []
        for skill, experience, score in zip(skills_domains, experiences, relevance_scores):
            result.append({
                "Skill/Domain": skill,
                "Experience": experience,
                "Relevance Score": score
            })

        print("Formatted result : ", result)
        return result
    else:
        print("No matching pattern found in the response text.")
        return []


def format_job_info_text(text):
    # Find the dictionary-like substring within curly braces
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        # Get the matched string and replace single quotes with double quotes for JSON parsing
        response_dict_str = match.group(0).replace("'", "\"")
        try:
            # Parse the JSON string into a Python dictionary
            response_dict = json.loads(response_dict_str)
        except json.JSONDecodeError as e:
            print(f"Error parsing response dictionary: {e}")
            return []

        # Extract data from the dictionary, providing default empty string values
        candidate = response_dict.get('Candidate', [''])[0]
        experience = response_dict.get('Candidate Experience', [''])[0]
        experience_percentage = response_dict.get('Candidate Experience Percentage', [''])[0]
        min_budget = response_dict.get('Candidate Minimum Budget', [''])[0]
        max_budget = response_dict.get('Candidate Maximum Budget', [''])[0]
        min_job_desc_exp = response_dict.get('Job Description Min Experience', [''])[0]
        max_job_desc_exp = response_dict.get('Job Description Max Experience', [''])[0]
        min_job_desc_package = response_dict.get('Job Description Min Package (LPA)', [''])[0]
        max_job_desc_package = response_dict.get('Job Description Max Package (LPA)', [''])[0]
        job_desc_skills = response_dict.get('Job Description Skills', [])
        job_desc_skills_count = response_dict.get('Job Description Skills Count', [''])[0]
        matching_skills = response_dict.get('Matching Skills', [])
        resume_skills = response_dict.get('Resume Skills', [])
        resume_skills_count = response_dict.get('Resume Skills Count', [''])[0]
        skills_percentage = response_dict.get('Skills Matching Percentage', [''])[0]

        # Create a formatted result as a list of dictionaries
        result = [{
            "Candidate": candidate,
            "Candidate Experience": experience,
            "Candidate Experience Percentage": experience_percentage,
            "Candidate Minimum Budget": min_budget,
            "Candidate Maximum Budget": max_budget,
            "Job Description Min Experience": min_job_desc_exp,
            "Job Description Max Experience": max_job_desc_exp,
            "Job Description Min Package (LPA)": min_job_desc_package,
            "Job Description Max Package (LPA)": max_job_desc_package,
            "Job Description Skills": job_desc_skills,
            "Job Description Skills Count": job_desc_skills_count,
            "Matching Skills": matching_skills,
            "Resume Skills": resume_skills,
            "Resume Skills Count": resume_skills_count,
            "Skills Matching Percentage": skills_percentage
        }]

        print("Formatted result:", result)
        return result
    else:
        print("No matching pattern found in the response text.")
        return []

def parse_career_progress(text):
    # Define patterns for matching work experience details
    experience_patterns = [
        r'(\d+(\.\d+)?)\s*(?:year|yr|years|yrs)?\s*(\d+)?\s*(?:month|months|mo|mos)?',
        r'(\d+)\s*(?:year|yr|years|yrs)?\s*(\d+)?\s*months?',
        r'(\d+(\.\d+)?)\s*(?:year|years|yr|yrs)?',
        r'(\d+(\.\d+)?)\s*(?:-year)',
        r'Demonstrated\s+(\d+(\.\d+)?)\s*(?:year|years|yr|yrs)',
        r'(\d+(\.\d+)?)\s*(?:years|yr|yrs)?\s*of\s*experience',
        r'(\d+(\.\d+)?)\s*years?\s*of\s*(?:IT|technical)?\s*experience',
        r'(\d+)\s*-\s*(\d+)\s*(?:years|yrs|yr|months|mos)?\s*experience',
        r'experience\s*of\s*(\d+(\.\d+)?)\s*(?:years|yrs|yr|months|mos)',
        r'over\s*(\d+(\.\d+)?)\s*(?:years|yrs|yr|months|mos)?\s*experience',
        r'(?:years|yrs|yr|months|mos)\s*of\s*experience\s*(\d+(\.\d+)?)',
        r'(\d+)\s*(?:years|yrs|yr|months|mos)\s*\(to\s*date\)',
        r'total\s*experience\s*of\s*(\d+(\.\d+)?)\s*(?:years|yrs|yr|months|mos)',
        r'experience\s*(?:from|since)\s*\d{4}\s*to\s*\d{4}',
        r'around\s*(\d+(\.\d+)?)\s*(?:years|yrs|yr|months|mos)?\s*of\s*experience'
    ]

    # Define a pattern to match the work experience details in the given format
    pattern = r'\[\s*(.*?)\s*\]'
    match = re.search(pattern, text, re.DOTALL)
    
    if match:
        work_experience_str = match.group(1)
        
        # Define a pattern to match individual work experiences, including the 'Project' field
        experience_pattern = r"\{\s*'Company':\s*'(.*?)',\s*'Title':\s*'(.*?)',\s*'From Date':\s*'(.*?)',\s*'To Date':\s*'(.*?)',\s*'Total Duration of Work':\s*'(.*?)',\s*'Location':\s*'(.*?)',\s*'Project':\s*'(.*?)'\s*\}"
        experiences = re.findall(experience_pattern, work_experience_str)
        
        # Convert the extracted details into a list of dictionaries
        result = []
        for experience in experiences:
            company = experience[0].strip()
            title = experience[1].strip()
            from_date = experience[2].strip()
            to_date = experience[3].strip()
            total_duration = experience[4].strip()
            location = experience[5].strip()
            project = experience[6].strip()

            # Match and parse the total duration using the experience patterns
            duration_match = None
            for pattern in experience_patterns:
                duration_match = re.search(pattern, total_duration, re.IGNORECASE)
                if duration_match:
                    break

            # Format the duration to be consistent
            if duration_match:
                total_duration = duration_match.group(0)
            else:
                total_duration = " "

            result.append({
                "Company": company,
                "Title": title,
                "From Date": from_date,
                "To Date": to_date,
                "Total Duration of Work": total_duration,
                "Location": location,
                "Project": project
            })

        return result
    else:
        print("No matching pattern found in the career progress response text.")
        return []



def parse_expertise_text(expertise_text):
    # Define regex patterns for categories and domains
    category_pattern = re.compile(r'(\w[\w\s]*?):\s*\[([^\]]*?)\s*,\s*count:\s*(\d+)\s*\]')
    domain_pattern = re.compile(r'\{\s*([^\{\}:]+):\s*([^\}]+)\}', re.DOTALL)
    
    # Parse categories
    categories = []
    for match in category_pattern.finditer(expertise_text):
        category_name = match.group(1).strip()
        items_text = match.group(2).strip()
        # count = match.group(3).strip()

        items = [item.strip() for item in items_text.split(',') if item.strip()]
        categories.append({
            "Category": category_name,
            "Items": items
        })

    # Parse domains
    domains = []
    domain_section_pattern = re.compile(r'Domains:\s*\[(.*?)\]\s*}', re.DOTALL)
    domain_section_match = domain_section_pattern.search(expertise_text)
    
    if domain_section_match:
        domain_section = domain_section_match.group(1)
        domain_matches = domain_pattern.findall(domain_section)
        for match in domain_matches:
            domain_name = match[0].strip()
            description = match[1].strip()
            domains.append({
                "Domain": domain_name,
                "Description": description
            })

    return {
        "categories": categories,
        "domains": domains
    }



def convert_to_array(candidate_learning_text):
    # Initialize dictionaries and lists to store parsed data
    company_names = []
    technologies_used = {}
    certifications = []
    skills_domain = []
    working_periods = {}

    # Extract Company Names
    company_names_match = re.search(r"Company Names:\s*\[(.*?)\]", candidate_learning_text)
    if company_names_match:
        company_names = [name.strip() for name in company_names_match.group(1).split(",")]

    # Extract Technologies Used
    techs_match = re.search(r"Technologies Used:\s*\{(.*?)\}", candidate_learning_text, re.DOTALL)
    if techs_match:
        techs_text = techs_match.group(1).strip()
        techs_entries = re.findall(r'([\w\s]+): \[([^\]]+)\]', techs_text)
        for company, techs in techs_entries:
            technologies_used[company.strip()] = [tech.strip() for tech in techs.split(",")]

    # Extract Certifications
    certs_match = re.search(r"Certifications:\s*\[(.*?)\]", candidate_learning_text)
    if certs_match:
        certifications = [cert.strip() for cert in certs_match.group(1).split(",") if cert.strip()]

    # Extract Skills/Domain
    skills_match = re.search(r"Skills/Domain:\s*\[(.*?)\]", candidate_learning_text)
    if skills_match:
        skills_domain = [skill.strip() for skill in skills_match.group(1).split(",") if skill.strip()]

    # Extract Working Periods and Calculate Duration
    periods_match = re.search(r"Working Periods:\s*\{(.*?)\}", candidate_learning_text, re.DOTALL)
    if periods_match:
        periods_text = periods_match.group(1).strip()
        periods_entries = re.findall(r'([\w\s]+): \{from_date: ([^,]+), to_date: ([^,]+), duration: ([^}]+)\}', periods_text)
        for company, from_date, to_date, duration in periods_entries:
            working_periods[company.strip()] = {"from": from_date.strip(), "to": to_date.strip(), "duration": duration.strip()}

    return {
        "Company Names": company_names,
        "Technologies Used": technologies_used,
        "Certifications": certifications,
        "Skills/Domain": skills_domain,
        "Working Periods": working_periods
    }


def convert_to_array_textual_representation(candidate_learning_text):
    # Initialize dictionaries and lists to store parsed data
    bullet_points = {}
    summary_paragraph = ""
    tags = []

    # Extract Bullet Points
    bullet_points_match = re.search(r"BulletPoints:\s*\{(.*?)\}", candidate_learning_text, re.DOTALL)
    if bullet_points_match:
        bullet_points_text = bullet_points_match.group(1).strip()
        bullet_points_entries = re.findall(r'BulletPoint\d+:\s*(.*?)(?:,|$)', bullet_points_text, re.DOTALL)
        for idx, point in enumerate(bullet_points_entries):
            bullet_points[f"BulletPoint{idx + 1}"] = point.strip()

    # Extract Summary Paragraph
    summary_paragraph_match = re.search(r"SummaryParagraph:\s*(.*?)(?:,|$)", candidate_learning_text, re.DOTALL)
    if summary_paragraph_match:
        summary_paragraph = summary_paragraph_match.group(1).strip()

    # Extract Tags
    tags_match = re.search(r"Tags:\s*\[(.*?)\]", candidate_learning_text)
    if tags_match:
        tags = [tag.strip() for tag in tags_match.group(1).split(",") if tag.strip()]

    return {
        "BulletPoints": bullet_points,
        "SummaryParagraph": summary_paragraph,
        "Tags": tags
    }


@app.route('/candidate_over_view', methods=['POST'])
def candidate_over_view():
    data = request.json
    user_id = data.get('user_id')
    job_id = data.get('job_id')
    recruiter_prompt = data.get('recruiter_prompt')
    pdf_base64 = data.get('resume')

    if not pdf_base64:
        return jsonify({"error": "Resume not provided or invalid"}), 400

    try:
        # Decode the base64 PDF file
        pdf_bytes = base64.b64decode(pdf_base64)
        
    except Exception as e:
        return jsonify({"error": "Failed to decode base64 PDF"}), 400

    try:
        # Extract text from the PDF
        pdf_text = extract_text_from_pdf(pdf_bytes)
    except Exception as e:
        return jsonify({"error": "Failed to extract text from PDF"}), 500

    # Fetch job details based on job_id
    job_details = get_job_details_candidate(job_id)
    if not job_details:
        return jsonify({"error": "Job details not found"}), 404

    

    expertise_prompt = f"""
Analyze this {pdf_text} and provide an accurate and detailed list of the candidate's areas of expertise. Include specific skills, technologies, programming languages, frameworks, tools, and domains where the candidate has demonstrated proficiency. Ensure the list is precise and comprehensive.

Format the response as follows:
{{
  "categories": {{
      "sub_category_1": [
          "Topic1", "Topic2", "Topic3", "count: N"
       ],
      "sub_category_2": [
          "Topic1", "Topic2", "count: N"
       ],
      ...
   
  "Domains": [
    {{
      "Domain 1": "Briefly describe the candidate's expertise in this domain, highlighting key skills and experience."
    }},
    {{
      "Domain 2": "Briefly describe the candidate's expertise in this domain, highlighting key skills and experience."
    }},
    {{
      "Domain 3": "Briefly describe the candidate's expertise in this domain, highlighting key skills and experience."
    }},
    {{
      "Domain 4": "Briefly describe the candidate's expertise in this domain, highlighting key skills and experience."
    }},
    {{
      "Domain 5": "Briefly describe the candidate's expertise in this domain, highlighting key skills and experience."
    }}
    // Add more domains as needed
  ]
  }},
}}
Each item in the Domains array should be an object with a single key-value pair where the key is the domain name and the value is a string description of the candidate's expertise in that domain. The descriptions should be clear and concise, summarizing the candidate's proficiency and experience in each domain.
"""


    
    job_info_prompt = f"""
Analyze the following {pdf_text} and {job_details}. Provide the details in the format below with no **Explanation** or theoretical content.

Output format:

categories = {{
    'Candidate': ['Candidate Name'],
    'Candidate Experience': ['Total experience in years'],
    'Candidate Experience Percentage': ['Calculated as described below'],
    'Candidate Minimum Budget': ['Calculated based on overall matching percentage'],
    'Candidate Maximum Budget': ['Calculated based on overall matching percentage'],
    'Job Description Min Experience': ['Minimum experience required'],
    'Job Description Max Experience': ['Maximum experience required'],
    # 'Job Description Package (LPA)': ['Package range mentioned'],
    'Job Description Min Package (LPA)': ['Minimum Package mentioned'],
    'Job Description Max Package (LPA)': ['Maximum Package mentioned'],
    'Job Description Skills': ['Skills required'],
    'Job Description Skills Count': ['Number of required skills'],
    'Matching Skills': ['Skills that match between job description skills and resume skills, considering sub-skills and main skills as explained below'],
    'Resume Skills': ['Skills listed in the resume'],
    'Resume Skills Count': ['Total number of skills listed'],
    'Skills Matching Percentage': ['(Number of matching skills / Total number of job description skills) * 100']
}}

Candidate Experience Percentage Calculation:
- Extract the minimum and maximum values from the job description experience range (e.g., "5-10 years").
- Convert these values to numbers (e.g., "5" and "10").
- If the candidate's experience falls within this range:
  - Set Candidate Experience Percentage to 100%.
  - Proceed with budget calculations.
- If the candidate's experience is outside this range:
  - Set Candidate Experience Percentage to 0%.
  - Set both Candidate Minimum Budget and Candidate Maximum Budget to 0.

Skills Matching Percentage Calculation:
- Compare each skill from the job description with the skills listed in the resume.
- Count the number of matching skills.
- Consider sub-skills and main skills for matching:
  - If a job description skill matches a sub-skill in the resume, count it as a match.
  - If a resume skill matches a main skill in the job description, count it as a match.
- Calculate Skills Matching Percentage as (Number of matching skills / Total number of job description skills) * 100.
- Ensure this percentage does not exceed 100%.

Overall Matching Percentage Calculation:
- Calculate Overall Matching Percentage as the average of Candidate Experience Percentage and Skills Matching Percentage.

Budget Calculations (if Candidate Experience Percentage is 100%):
- **Candidate Minimum Budget**: min_package + ((overall_matching_percentage / 100) * (max_package - min_package) * 0.5)
- **Candidate Maximum Budget**: min_package + ((overall_matching_percentage / 100) * (max_package - min_package) * 1.5)

where:
- **min_package** is the minimum package from the job description.
- **max_package** is the maximum package from the job description.
- **overall_matching_percentage** is the average of Candidate Experience Percentage and Skills Matching Percentage.

Example Calculations:
Case 1: Experience within Range and Overall Matching Percentage
Candidate Experience: 5.4 years
Job Description Experience: 5-10 years
Candidate Experience Percentage: 100%
Skills Matching Percentage: 30%
Overall Matching Percentage: 65%
Candidate Minimum Budget: min_package + ((65 / 100) * (max_package - min_package) * 0.5)
Candidate Maximum Budget: min_package + ((65 / 100) * (max_package - min_package) * 1.5)

Case 2: Experience Outside Range
Candidate Experience: 9 years
Job Description Experience: 5-10 years
Candidate Experience Percentage: 0%
Skills Matching Percentage: 30%
Candidate Minimum Budget: 0
Candidate Maximum Budget: 0

"""


    carrer_progress = f"""
Analyze the following {pdf_text} and provide a detailed overview of the candidate's career progress. Focus on the progression of job titles, the duration of each position, and notable achievements in each role. Highlight the candidate's growth and development over time, including roles, promotions, location of company, project title  and achievements. Discuss any patterns or significant milestones, such as shifts in industry, increasing levels of responsibility, or specialized expertise development.

For each role, provide the following details in the array format (no responsibilities or detailed explanations):
[
    {{
        'Company': '',
        'Title': '',
        'From Date': '',
        'To Date': '',
        'Total Duration of Work': '',
        'Location': '',
        'Project': ''
    }}
]
    """



    candidate_learning = f"""
       Analyze {pdf_text}. Please provide the technologies used, any certifications mentioned, and the working periods organized by each company he has worked for, including the total experience duration at each company. Please ensure that the response remains consistent every time this prompt is used and do not modify the values each time:
    here technologies take from  (Environment or Tools and Technologies) 
*Example Arrays*:
- *Company Names*: ["World Pay India", "EPAM Systems India", "Global Logic", "Collabrera Technologies"]
- *Technologies Used*: {{
    "World Pay India": ["Java", "Spring Framework", "Spring Boot", ...],
    "EPAM Systems India": ["Java", "Spring Boot", "Hibernate", ...],
    "Global Logic": ["Java", "Spring Boot", "Hibernate", ...],
    "Collabrera Technologies": ["Java", "Spring Boot", "Hibernate", ...]
  }}
- *Certifications*: [" Java Full Stack Development Certification From Jspiders ", " ", ...]
- *Skills/Domain*: ["Core Java", "Spring Framework", "Spring Boot", ...]
- *Working Periods*: {{
    "World Pay India": {{"from_date": "", "to_date": "" , "duration": "X years Y months"}},
    "EPAM Systems India": {{"from_date": "", "to_date": "" , "duration": "X years Y months"}},
    "Global Logic": {{"from_date": "", "to_date": ""  "duration": "X years Y months"}},
    "Collabrera Technologies": {{"from_date": "", "to_date": "" , "duration": "X years Y months"}}
  }}

Please ensure each section is detailed and well-organized for clarity, with a specific focus on differentiating the technologies used by each company. Only provide the array responses without any theoretical explanations.
"""

    candidate_learning_textual_representation = f"""
Analyze {pdf_text} and provide a high-level summary of the candidate's learning attitude. Focus on indications of the candidate's openness to learning new things, such as certifications, courses, new technologies adopted, or any other learning initiatives mentioned. Ensure the response is concise and adaptable to various resume formats.

Format the response as follows:
{{
    "BulletPoints": {{
        "BulletPoint1": "Description of learning activity 1.",
        "BulletPoint2": "Description of learning activity 2.",
        "BulletPoint3": "Description of learning activity 3.",
        "BulletPoint4": "Description of learning activity 4.",
        "BulletPoint5": "Description of learning activity 5."
    }},
    "SummaryParagraph": "Brief paragraph summarizing the candidate's learning attitude.",
    "Tags": [
        "Tag1",
        "Tag2",
        "Tag3",
        "Tag4",
        "Tag5"
    ]
}}
"""

    Analyze_candidate_profile = f"""
Analyze the candidate profile using this {job_details} and {pdf_text}.
This is the flow:

1. Extract Key Skills and Domains: Identify the primary skills and domains from the {job_details}.

2. Analyze Candidate's Profile: Compare the extracted skills and domains from the {job_details} with the candidate's skills and experience.

3. Market Trends Analysis: Evaluate the relevance of the candidate’s skills and domains against current market trends and demands.

4. Scoring/Ranking: Assign a relevance score to each candidate based on how well their skills and experience match the {job_details} and market trends.

*Example Arrays*:
- *Skills/Domain*: ["Core Java", "Spring Boot", "Database", ...]
- *Candidate Experience*: ["High", "Low", "Moderate", ...]
- *Relevance Score*: [5, 4, 2, 3, ...]

*Output Format*: Provide the data as separate arrays in JSON format, without theoretical explanations, focusing on the scoring/ranking.
    """

    # Configure and use Generative AI

    api_key = config.api_key
    if api_key is None:
        raise ValueError("API_KEY environment variable not set")

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')

    try:
        # Generate content for both prompts
        expertise_response = model.generate_content(expertise_prompt)
        job_info_response = model.generate_content(job_info_prompt)
        carrer_progress_response = model.generate_content(carrer_progress)
        candidate_learning_response = model.generate_content(candidate_learning)
        candidate_learning_textual_representation_response = model.generate_content(candidate_learning_textual_representation)
        Analyze_candidate_profile_response = model.generate_content(Analyze_candidate_profile)
    except Exception as e:
        return jsonify({"error": "Failed to generate content using Generative AI"}), 500

    # Ensure response content is available
    expertise_text = getattr(expertise_response, 'text', '')
    job_info_text = getattr(job_info_response, 'text', '')
    carrer_progress_text = getattr(carrer_progress_response, 'text', '')
    candidate_learning_text = getattr(candidate_learning_response, 'text', '')
    candidate_learning_textual_representation_text = getattr(candidate_learning_textual_representation_response, 'text', '')
    
    # Check if response text is valid and not blocked
    try:
        Analyze_candidate_profile_text = getattr(Analyze_candidate_profile_response, 'text', '')
    except ValueError as e:
        # print(f"Error accessing response text: {e}")
        return jsonify({"error": "Failed to access Analyze Candidate Profile response text"}), 500

    formatted_expertise_text = clean_response(expertise_text)
    formatted_job_info_text = clean_response(job_info_text)
    formatted_career_progress_text = clean_response(carrer_progress_text)
    formatted_candidate_learning_text = clean_response(candidate_learning_text)
    formatted_candidate_learning_textual_representation_text_clean = clean_response(candidate_learning_textual_representation_text)


    formatted_Analyze_candidate_profile_text = format_analyze_candidate_profile(Analyze_candidate_profile_text)
    formatted_job_info_text = format_job_info_text(formatted_job_info_text)
    formatted_career_progress_text = parse_career_progress(formatted_career_progress_text)
    formatted_expertise_text = parse_expertise_text(formatted_expertise_text)
    formatted_candidate_learning_text = convert_to_array(formatted_candidate_learning_text)
    formatted_candidate_learning_textual_representation_text = convert_to_array_textual_representation(formatted_candidate_learning_textual_representation_text_clean)

    response_data = {
        'user_id': user_id,
        'expertise_response': formatted_expertise_text,
        'job_info_response': formatted_job_info_text,
        'career_progress_response': formatted_career_progress_text,
        'candidate_learning_response': formatted_candidate_learning_text,
        'candidate_learning_textual_representation':formatted_candidate_learning_textual_representation_text,
        'analyze_candidate_profile_response': formatted_Analyze_candidate_profile_text
    }

    return jsonify(response_data)



########################################################################################

@app.route('/check_candidate', methods=['POST'])
def check_candidate():
    data = request.json
    clients = []
    profiles = []
    dates = []
    job_ids = []
    status = []

    email = data.get('email')
    mobile = data.get('mobile')

    # Query the database to check for an existing candidate with the provided mobile or email
    existing_candidates = Candidate.query.filter(or_(Candidate.mobile == mobile, Candidate.email == email)).all()
    
    for candidate in existing_candidates:
        clients.append(candidate.client)
        profiles.append(candidate.profile)
        dates.append(candidate.date_created.strftime('%Y-%m-%d'))
        job_ids.append(candidate.job_id)
        status.append(candidate.status)

    if existing_candidates:
        response = {
            'message': "Candidate with this mobile or email already exists.",
            'clients': clients,
            'profiles': profiles,
            'dates': dates,
            'jobIds': job_ids,
            'status': status
        }
    else:
        response = {
            'message': "Mobile and email not available.",
            'clients': None,
            'profiles': None,
            'dates': None,
            'jobIds': None,
            'status': None
        }

    return jsonify(response)


# @app.route('/check_candidate', methods=['POST'])
# def check_candidate():
#     clients = []
#     profiles = []
#     dates=[]
#     job_ids=[]
#     status=[]
#     field = request.json['field']
#     value = request.json['value']

#     # Query the database to check for an existing candidate with the provided mobile or email
#     existing_candidate = Candidate.query.filter(or_(Candidate.mobile == value, Candidate.email == value)).all()
#     for i in existing_candidate:
#         clients.append(" " + i.client + " ")
#         profiles.append(" " + i.profile + " " )
#         dates.append(i.date_created.strftime('%Y-%m-%d'))
#         job_ids.append(i.job_id)
#         status.append(i.status)

#     # candidate = Candidate.query.filter_by(mobile=existing_candidate.mobile).first()
#     if existing_candidate:
#         response = {
#             'message' : f"Candidate with this {field} already exists.",
#             'client' : clients,
#             'profile' : profiles,
#             'dates':dates,
#             'jobId':job_ids,
#             'status':status
#         }

#     else:
#         response = {
#             'message': f"{field.capitalize()} is available.",
#             'client': None,
#             'profile': None,
#             'dates':None,
#             'jobId':None,
#             'status':None
#         }
#     return json.dumps(response)



@app.route('/forgot_password')
def forgot_password():
    return render_template('forgot_password.html')

def generate_6otp():
    digits = "0123456789"
    otp = "".join(random.choice(digits) for _ in range(6))
    return otp


@app.route('/generate_otp', methods=['POST'])
def generate_otp():
    if request.method == 'POST':
        data = request.json
        username = data.get('username')
        email = data.get('email')
        user = User.query.filter_by(username=username, email=email).first()
        if user:
            otp = generate_6otp()
            user.otp = otp
            db.session.commit()
            msg = Message('New OTP Generated', sender=config.sender_email, recipients=[email])
            msg.html = f'''
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        background-color: #f4f4f4;
                        color: #333;
                        margin: 0;
                        padding: 20px;
                    }}
                    .container {{
                        background-color: #ffffff;
                        max-width: 600px;
                        margin: 0 auto;
                        padding: 20px;
                        border: 1px solid #dddddd;
                        border-radius: 8px;
                        box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
                    }}
                    .header {{
                        font-size: 24px;
                        font-weight: bold;
                        margin-bottom: 20px;
                        color: #4CAF50;
                    }}
                    .content p {{
                        font-size: 16px;
                        line-height: 1.6;
                        margin: 10px 0;
                    }}
                    .otp-label {{
                        display: block;
                        margin-bottom: 10px;
                        font-weight: bold;
                    }}
                    .otp-input {{
                        width: 50%;
                        padding: 10px;
                       font-size: 20px;
                        font-weight: bold;
                        border: 1px solid #eeeeee;
                        border-radius: 5px;
                        background-color: #f9f9f9;
                        box-sizing: border-box;
                        outline: none; /* Remove the default outline */
                    }}
                    .footer {{
                        font-size: 12px;
                        color: #999;
                        margin-top: 20px;
                        text-align: center;
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="header">New OTP Generated</div>
                    <div class="content">
                        <p>Hi {user.name},</p>
                        <p>OTP for resetting your password:</p>
                        <p><input type="text" class="otp-input" id="otp" value="{otp}" readonly></p>
                    </div>
                    <div class="footer">
                        <p>If you did not request this change, please contact our support team immediately.</p>
                        <p><b>Makonis Talent Track Pro Team</b></p>
                    </div>
                </div>
            </body>
            </html>
            '''
            mail.send(msg)
            return jsonify({'status': 'success', 'message': 'OTP has been sent to your email.'})
        else:
            return jsonify({'status': 'error', 'message': 'User does not exist.'})
    else:
        return jsonify({'status': 'error', 'message': 'Invalid request method.'})





# @app.route('/generate_otp', methods=['POST'])
# def generate_otp():
#     if request.method == 'POST':
#         username = request.json.get('username')
#         email = request.json.get('email')
#         user = User.query.filter_by(username=username, email=email).first()
#         if user:
#             otp = generate_6otp()
#             user.otp = otp
#             db.session.commit()
#             msg = Message('Account Verification', sender=config.sender_email, recipients=[email])
#             msg.body = f'Hi {user.name},\n\n OTP for resetting your password {otp}.'
#             mail.send(msg)
#             return jsonify({'status': 'success', 'message': 'OTP has been sent to your email.'})
#         else:
#             return jsonify({'status': 'error', 'message': 'User does not exist.'})
#     else:
#         return jsonify({'status': 'error', 'message': 'Invalid request method.'})
    



@app.route('/reset_password', methods=['POST'])
def reset_password():
    if request.method == 'POST':
        data = request.json
        otp = data['otp']
        new_password = data.get('new_password')
        confirm_password = data.get('confirm_password')
        new_password_hashed = hashlib.sha256(new_password.encode()).hexdigest()

        user = User.query.filter_by(otp=otp).first()

        if user and user.otp == otp and new_password == confirm_password:
            # Check if the new password is different from the old password
            if new_password_hashed != user.password:  # comparing hashes
                user.password = new_password_hashed
                db.session.commit()
                # Send the updated password to the user's email
                msg = Message('Password Changed', sender=config.sender_email, recipients=[user.email])
                msg.html = f'''
                <!DOCTYPE html>
                <html lang="en">
                <head>
                    <meta charset="UTF-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1.0">
                    <style>
                        body {{
                            font-family: Arial, sans-serif;
                            background-color: #f4f4f4;
                            color: #333;
                            margin: 0;
                            padding: 20px;
                        }}
                        .container {{
                            background-color: #ffffff;
                            max-width: 600px;
                            margin: 0 auto;
                            padding: 20px;
                            border: 1px solid #dddddd;
                            border-radius: 8px;
                            box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
                        }}
                        .header {{
                            font-size: 24px;
                            font-weight: bold;
                            margin-bottom: 20px;
                            color: #4CAF50;
                        }}
                        .content {{
                            font-size: 16px;
                            line-height: 1.6;
                        }}
                        .credentials {{
                            background-color: #f9f9f9;
                            padding: 10px;
                            border: 1px solid #eeeeee;
                            border-radius: 5px;
                            margin-top: 10px;
                        }}
                        .footer {{
                            font-size: 12px;
                            color: #999;
                            margin-top: 20px;
                            text-align: center;
                        }}
                    </style>
                </head>
                <body>
                    <div class="container">
                        <div class="header">Password Changed</div>
                        <div class="content">
                            <p>Hello {user.name},</p>
                            <p>Your password has been successfully changed. </p>
                            <p>Here are your updated credentials:,</p>
                            <div class="credentials">
                                <p><strong>Username:</strong> {user.username}</p>
                                <p><strong>Password:</strong> {new_password}</p>
                            </div>
                        </div>
                        <div class="footer">
                            <p>If you did not request this change, please contact our support team immediately.</p>
                            <p><b>Makonis Talent Track Pro Team</b></p>
                        </div>
                    </div>
                </body>
                </html>
                '''
                mail.send(msg)

                return jsonify({'status': 'success', 'message': 'Password changed successfully.'})
            else:
                return jsonify({'status': 'error', 'message': 'New password is the same as the old password'})
        else:
            return jsonify({'status': 'error', 'message': 'Invalid OTP or password confirmation. Please try again.'})

    return jsonify({'status': 'error', 'message': 'Invalid request method.'})


# @app.route('/reset_password', methods=['POST'])
# def reset_password():
#     if request.method == 'POST':
#         otp = request.json['otp']
#         new_password = request.json.get('new_password')
#         confirm_password = request.json.get('confirm_password')
#         new_password_hashed = hashlib.sha256(new_password.encode()).hexdigest()

#         user = User.query.filter_by(otp=otp).first()

#         if user and user.otp == otp and new_password == confirm_password:
#             # Check if the new password is different from the old password
#             if new_password_hashed != user.password:  # comparing hashes
#                 user.password = new_password_hashed
#                 db.session.commit()
#                 # Send the updated password to the user's email
#                 msg = Message('Password Changed', sender=config.sender_email, recipients=[user.email])
#                 msg.body = f'Hello {user.name},\n\nYour password has been successfully changed. Here are your updated credentials:\n\nUsername: {user.username}\nPassword: {new_password}'
#                 mail.send(msg)

#                 return jsonify({'status': 'success', 'message': 'Password changed successfully.'})
#             else:
#                 return jsonify({'status': 'error', 'message': 'New password is the same as the old password'})
#         else:
#             return jsonify({'status': 'error', 'message': 'Invalid OTP or password confirmation. Please try again.'})

#     return jsonify({'status': 'error', 'message': 'Invalid request method.'})



def generate_html_message(message, redirect_url=None):
    html_message = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Message</title>
            <!-- Add CSS styles for your message -->
            <style>
                /* Example CSS styles */
                body {{
                    font-family: Arial, sans-serif;
                    background-color: #f4f4f4;
                    text-align: center;
                }}
                .message {{
                    margin-top: 50px;
                    background-color: #fff;
                    border-radius: 10px;
                    padding: 20px;
                    box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
                    display: inline-block;
                }}
                .message p {{
                    font-size: 20px;
                }}
                .button {{
                    display: inline-block;
                    padding: 10px 20px;
                    background-color: #007bff;
                    color: #fff;
                    text-decoration: none;
                    border-radius: 5px;
                    transition: background-color 0.3s;
                    margin-top: 20px;
                }}
                .button:hover {{
                    background-color: #0056b3;
                }}
            </style>
        </head>
        <body>
            <div class="message">
                <p>{message}</p>
                {'' if redirect_url is None else f'<a href="{redirect_url}" class="button">Login</a>'}
            </div>
        </body>
        </html>
    """
    return html_message

@app.route('/verify/<token>')
def verify(token):
    user_id = verify_token(token)
    if user_id:
        user = User.query.get(user_id)
        if user.is_verified:
            if user.user_type == 'management':
                message = 'Your Management Account is already Verified. Please Login!'
                return generate_html_message(message, redirect_url='https://ats-makonis.netlify.app/ManagementLogin')
            elif user.user_type == 'recruiter':
                message = 'Your Recruiter Account is already Verified. Please Login!'
                return generate_html_message(message, redirect_url='https://ats-makonis.netlify.app/RecruitmentLogin')
        else:
            user.is_verified = True
            db.session.commit()
            if user.user_type == 'management':
                message = 'Your Management Account has been Successfully Verified. Please Login!'
                return generate_html_message(message, redirect_url='https://ats-makonis.netlify.app/ManagementLogin')
            elif user.user_type == 'recruiter':
                message = 'Your Recruiter Account has been Successfully Verified. Please Login!'
                return generate_html_message(message, redirect_url='https://ats-makonis.netlify.app/RecruitmentLogin')
    else:
        message = 'Your verification link has expired. Please contact management to activate your account.'
        return generate_html_message(message)


# @app.route('/verify/<token>')
# def verify(token):
#     user_id = verify_token(token)
#     if user_id:
#         user = User.query.get(user_id)
#         user.is_verified = True
#         db.session.commit()
#         if user.user_type == 'management':
#             message = 'Your Management Account has been Successfully Verified. Please Login!'
#             return generate_html_message(message, redirect_url='https://ats-makonis.netlify.app/ManagementLogin')
#         elif user.user_type == 'recruiter':
#             message = 'Your Recruiter Account has been Successfully Verified!'
#             return generate_html_message(message, redirect_url='https://ats-makonis.netlify.app/RecruitmentLogin')
#     else:
#         message = 'Your verification link has expired. Please contact management to activate your account.'
#         return generate_html_message(message)


# @app.route('/verify/<token>')
# def verify(token):
#     user_id = verify_token(token)
#     if user_id:
#         user = User.query.get(user_id)
#         user.is_verified = True
#         db.session.commit()
#         if user.user_type == 'management':
#             return jsonify({'status': 'success', 'message': 'Account verified successfully!', 'redirect': url_for('management_login', verification_msg_manager='Your Account has been Successfully Verified. Please Login.')})
#         elif user.user_type == 'recruiter':
#             return jsonify({'status': 'success', 'message': 'Account verified successfully!', 'redirect': url_for('recruiter_index')})
#     else:
#         return jsonify({'status': 'error', 'message': 'Your verification link has expired. Please contact management to activate your account.'})
#     return jsonify({'status': 'error', 'message': 'An error occurred while verifying your account.'})



# Function to generate a random password
def generate_random_password(length=8):
    digits = string.digits
    password = ''.join(random.choice(digits) for _ in range(length - 3))
    return "Mak" + password


@app.route('/signup', methods=['POST'])
def signup():
    data = request.json
    user_id = data.get('user_id')  # Using get method to avoid KeyError
    user = User.query.filter_by(id=user_id).first()

    if not user:
        return jsonify({'status': 'error', 'message': 'Invalid user ID or user does not exist.'})

    user_type = user.user_type
    user_name = user.username

    if user_type == 'management':
        username = data.get('username')
        name = data.get('name')
        email = data.get('email')
        user_type = data.get('user_type')

        # Check if required fields are provided
        if not all([username, name, email, user_type]):
            return jsonify({'status': 'error', 'message': 'All fields are required'})

        # Generate a random password
        password = generate_random_password()
        hashed_password = hashlib.sha256(password.encode()).hexdigest()

        created_by = user_name

        existing_user = User.query.filter(or_(User.username == username, User.email == email, User.name == name)).first()

        if existing_user:
            return jsonify({'status': 'error', 'message': 'Account with the same Username, Email, or Name already exists.'})

        new_user = User(username=username, password=hashed_password, name=name, email=email, user_type=user_type, created_by=created_by)
        
        db.session.add(new_user)
        db.session.commit()

        # Generate a verification token
        verification_token = generate_verification_token(new_user.id)

        # Create the verification link
        verification_link = url_for('verify', token=verification_token, _external=True)

        # Send the verification email
        msg = Message('Account Verification', sender=config.sender_email, recipients=[new_user.email])
        
        msg.html = f'''
        <html>
        <head>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                }}
                .container {{
                    max-width: 600px;
                    margin: auto;
                    padding: 20px;
                    border: 1px solid #ddd;
                    border-radius: 5px;
                    box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
                }}
                h2 {{
                    color: #333;
                }}
                p {{
                    color: #555;
                }}
                ul {{
                    list-style-type: none;
                    padding: 0;
                }}
                ul li {{
                    background: #f9f9f9;
                    margin: 5px 0;
                    padding: 10px;
                    border: 1px solid #ddd;
                    border-radius: 3px;
                }}
                a {{
                    color: #1a73e8;
                    text-decoration: none;
                }}
                a:hover {{
                    text-decoration: underline;
                }}
                .footer {{
                    margin-top: 20px;
                    font-size: 0.9em;
                    color: #888;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <h2>Hello {new_user.name},</h2>
                <p>We are pleased to inform you that your account has been successfully created for the <strong>ATS Makonis Talent Track Pro</strong>. Here are your login credentials:</p>
                <ul>
                    <li><strong>Username:</strong> {new_user.username}</li>
                    <li><strong>Password:</strong> {password}</li>
                </ul>
                <p>Please note that the verification link will expire after 24 hours.</p>
                <p>To verify your account, please click on the following link:</p>
                <p><a href="{verification_link}">Verify Your Account</a></p>
                <p>After successfully verifying your account, you can access the application using the following link:</p>
                <p><a href="https://ats-makonis.netlify.app/">Application Link (Post Verification)</a></p>
                <p>If you have any questions or need assistance, please feel free to reach out.</p>
                <p>Best regards,</p>
                <p><strong>ATS Makonis Talent Track Pro Team</strong></p>
                <div class="footer">
                    <p>This is an automated message, please do not reply.</p>
                </div>
            </div>
        </body>
        </html>
        '''
        mail.send(msg)

        return jsonify({'status': 'success',
            'message': 'A verification email has been sent to your email address. Please check your inbox.',
            'success_message': 'Account created successfully'
        })
    else:
        return jsonify({'message': 'You do not have permission to create recruiter accounts.'})



# @app.route('/signup', methods=['POST'])
# def signup():
#     data = request.json
#     user_id = data.get('user_id')  # Using get method to avoid KeyError
#     user = User.query.filter_by(id=user_id).first()

#     if not user:
#         return jsonify({'status': 'error', 'message': 'Invalid user ID or user does not exist.'})

#     user_type = user.user_type
#     user_name = user.username

#     if user_type == 'management':
#         username = data.get('username')
#         name = data.get('name')
#         email = data.get('email')
#         user_type = data.get('user_type')

#         # Check if required fields are provided
#         if not all([username, name, email, user_type]):
#             return jsonify({'status': 'error', 'message': 'All fields are required'})

#         # Generate a random password
#         password = generate_random_password()
#         hashed_password = hashlib.sha256(password.encode()).hexdigest()

#         created_by = user_name

#         existing_user = User.query.filter(or_(User.username == username, User.email == email, User.name == name)).first()

#         if existing_user:
#             return jsonify({'status': 'error', 'message': 'Account with the same Username, Email, or Name already exists.'})

#         new_user = User(username=username, password=hashed_password, name=name, email=email, user_type=user_type, created_by=created_by)
        
#         db.session.add(new_user)
#         db.session.commit()

#         # Generate a verification token
#         verification_token = generate_verification_token(new_user.id)

#         # Create the verification link
#         verification_link = url_for('verify', token=verification_token, _external=True)

#         # Send the verification email
#         msg = Message('Account Verification', sender=config.sender_email, recipients=[new_user.email])
#         msg.body = f'Hello {new_user.name},\n\n We are pleased to inform you that your account has been successfully created for the ATS Makonis Talent Track Pro. Here are your login credentials:\n\nUsername: {new_user.username}\nPassword: {password}\n\n Please note that the verification link will expire after 24 hours. \n\n After successfully verifying your account, you can access the application using the following link : \n\n Application Link (Post Verification): https://ats-makonis.netlify.app/ \n\n To verify your account, please click on the following link: {verification_link} \n\n If you have any questions or need assistance, please feel free to reach out. \n\n Best regards, '
#         mail.send(msg)

#         # return jsonify({'message': 'A verification email has been sent to your email address. Please check your inbox.'})
#         return jsonify({'status': 'success',
#             'message': 'A verification email has been sent to your email address. Please check your inbox.',
#             'success_message': 'Account created successfully'
#             })
#     else:
#         return jsonify({'message': 'You do not have permission to create recruiter accounts.'})


import hashlib

@app.route('/signup-onetime', methods=['POST'])
def signup_onetime():
    data = request.json
    username = data.get('username')
    password = data.get('password')
    name = data.get('name')
    email = data.get('email')
    user_type = 'management'
    registration_completed = 'one_time'

    # Hash the password using SHA-256
    hashed_password = hashlib.sha256(password.encode()).hexdigest()

    user_onetime = User.query.filter_by(registration_completed='one_time').first()
    if user_onetime:
        return jsonify({'status': 'error','message': 'The one-time registration for this application has already been completed.'})

    new_user = User(username=username, password=hashed_password, name=name,
                    email=email, user_type=user_type, registration_completed=registration_completed)

    db.session.add(new_user)
    db.session.commit()

    # Generate a verification token
    verification_token = generate_verification_token(new_user.id)

    # Create the verification link
    verification_link = url_for('verify', token=verification_token, _external=True)

    # Send the verification email
    msg = Message('Account Verification', sender=config.sender_email, recipients=[new_user.email])
    
    msg.html = f'''
    <html>
    <head>
        <style>
            body {{
                font-family: 'Arial', sans-serif;
                background-color: #f7f7f7;
                margin: 0;
                padding: 0;
            }}
            .container {{
                max-width: 600px;
                margin: 20px auto;
                background-color: #ffffff;
                padding: 20px;
                border: 1px solid #e0e0e0;
                border-radius: 10px;
            }}
            h2 {{
                color: #333333;
                font-size: 24px;
                margin-bottom: 20px;
            }}
            p {{
                color: #666666;
                line-height: 1.6;
                margin: 10px 0;
            }}
            .credentials {{
                background-color: #f9f9f9;
                padding: 10px;
                border-radius: 5px;
                border: 1px solid #dddddd;
            }}
            .credentials li {{
                margin: 5px 0;
                padding: 5px;
            }}
            a {{
                color: #1a73e8;
                text-decoration: none;
            }}
            a:hover {{
                text-decoration: underline;
            }}
            .button {{
                display: inline-block;
                padding: 10px 20px;
                margin: 20px 0;
                font-size: 16px;
                color: #ffffff;
                background-color: #333333;
                border-radius: 5px;
                text-align: center;
                text-decoration: none;
            }}
            .footer {{
                margin-top: 20px;
                font-size: 0.9em;
                color: #999999;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <h2>Account Verification</h2>
            <p>Hello {new_user.name},</p>
            <p>We are pleased to inform you that your account has been successfully created for the <strong>ATS Makonis Talent Track Pro</strong>. Here are your login credentials:</p>
            <div class="credentials">
                <ul>
                    <li><strong>Username:</strong> {new_user.username}</li>
                    <li><strong>Password:</strong> {password}</li>
                </ul>
            </div>
            <p>Please note that the verification link will expire in 24 hours.</p>
            <p>To verify your account, please click the button below:</p>
            <p><a href="{verification_link}" class="button">Verify Your Account</a></p>
            <p>After verifying your account, you can access the application using the following link:</p>
            <p><a href="https://ats-makonis.netlify.app/" class="button">Go to ATS Makonis Talent Track Pro</a></p> 
            <p>If you have any questions or need assistance, please feel free to reach out to our support team.</p>
            <p>Best regards,</p>
            <p><strong>ATS Makonis Talent Track Pro Team</strong></p>
            <div class="footer">
                <p>This is an automated message, please do not reply.</p>
            </div>
        </div>
    </body>
    </html>
    '''
    mail.send(msg)

    return jsonify({'status': 'success',
        'message': 'A verification email has been sent to your email address. Please check your inbox.',
        'success_message': 'Account created successfully'
    }), 200


# @app.route('/signup-onetime', methods=['POST'])
# def signup_onetime():
#     if request.method == 'POST':
#         username = request.json.get('username')
#         password = request.json.get('password')
#         name = request.json.get('name')
#         email = request.json.get('email')
#         user_type = 'management'
#         registration_completed = 'one_time'

#         # Hash the password using SHA-256
#         hashed_password = hashlib.sha256(password.encode()).hexdigest()

#         user_onetime = User.query.filter_by(registration_completed='one_time').first()
#         if user_onetime:
#             return jsonify({'message': 'The one-time registration for this application has already been completed.'}),400

#         new_user = User(username=username, password=hashed_password, name=name,
#                         email=email, user_type=user_type, registration_completed=registration_completed)

#         db.session.add(new_user)
#         db.session.commit()

#         # Generate a verification token
#         verification_token = generate_verification_token(new_user.id)

#         # Create the verification link
#         verification_link = url_for('verify', token=verification_token, _external=True)

#         # Construct the message body with username and plaintext password
#         message_body = f'Hello {new_user.name},\n\nWe are pleased to inform you that your account has been successfully created for the ATS Makonis Talent Track Pro.\n\nYour login credentials:\n\nUsername: {new_user.username}\nPassword: {password}\n\nTo complete the account setup, kindly click on the verification link below:\n{verification_link}\n\nPlease note that the verification link will expire after 24 hours.\n\nAfter successfully verifying your account, you can access the application using the following link:\n\nApplication Link (Post Verification): https://ats-makonis.netlify.app/\n\nIf you have any questions or need assistance, please feel free to reach out.\n\nBest regards,'

#         # Send the verification email
#         msg = Message('Account Verification', sender=config.sender_email, recipients=[new_user.email])
#         msg.body = message_body
#         mail.send(msg)

#         # return jsonify({'message': 'A verification email has been sent to your email address. Please check your inbox.'})
#         return jsonify({
#             'message': 'A verification email has been sent to your email address. Please check your inbox.',
#             'success_message': 'Account created successfully'
#             }),200

#     return jsonify({'message': 'Invalid request method.'}),400

# @app.route('/login/recruiter', methods=['POST'])
# def recruiter_login():
#     verification_msg = request.args.get('verification_msg')
#     reset_message = request.args.get('reset_message')
#     session_timeout_msg = request.args.get("session_timeout_msg")
#     password_message = request.args.get('password_message')

#     if request.method == 'POST':
#         username = request.json.get('username')
#         password = request.json.get('password')

#         # Hash the entered password
#         hashed_password = hashlib.sha256(password.encode()).hexdigest()

#         # Check if the user exists and the password is correct
#         user = User.query.filter_by(username=username, password=hashed_password, user_type='recruiter').first()

#         if user:
#             if user.is_active:  # Check if the user is active
#                 if user.is_verified:
#                     # Set the user session variables
#                     session['user_id'] = user.id
#                     session['user_type'] = user.user_type
#                     session['username'] = user.username
#                     session['user_name'] = user.name
#                     session['JWT Token'] = secrets.token_hex(16)
#                     return jsonify({'status': 'success', 'redirect': url_for('dashboard'),'user_id': user.id})
#                 else:
#                     error = 'Your account is not verified yet. Please check your email for the verification link.'
#             else:
#                 error = 'Your account is not active. Please contact the administrator.'
#         else:
#             error = 'Invalid username or password'

#         return jsonify({'status': 'error', 'error': error})

#     # For GET requests, return necessary data
#     return jsonify({
#         'status': 'success',
#         'verification_msg': verification_msg,
#         'reset_message': reset_message,
#         'session_timeout_msg': session_timeout_msg,
#         'password_message': password_message
#     })


@app.route('/login/recruiter', methods=['POST'])
def recruiter_login():
    verification_msg = request.args.get('verification_msg')
    reset_message = request.args.get('reset_message')
    session_timeout_msg = request.args.get("session_timeout_msg")
    password_message = request.args.get('password_message')

    if request.method == 'POST':
        username = request.json.get('username')
        password = request.json.get('password')

        # Hash the entered password
        hashed_password = hashlib.sha256(password.encode()).hexdigest()

        # Check if the user exists and the password is correct
        user = User.query.filter_by(username=username, password=hashed_password, user_type='recruiter').first()

        if user:
            if user.is_active:  # Check if the user is active
                if user.is_verified:
                    # Set the user session variables
                    session['user_id'] = user.id
                    session['user_type'] = user.user_type
                    session['username'] = user.username
                    session['user_name'] = user.name
                    session['JWT Token'] = secrets.token_hex(16)
                    return jsonify({
                        'status': 'success',
                        'redirect': url_for('dashboard'),
                        'user_id': user.id,
                        'username': user.username,
                        'name': user.name,
                        'email':user.email
                    })
                else:
                    message = 'Your account is not verified yet. Please check your email for the verification link.'
            else:
                message = 'Your account is not active. Please contact the administrator.'
        else:
            message= 'Invalid username or password'

        return jsonify({'status': 'error', 'message': message})

    # For GET requests, return necessary data
    return jsonify({
        'status': 'success',
        'verification_msg': verification_msg,
        'reset_message': reset_message,
        'session_timeout_msg': session_timeout_msg,
        'password_message': password_message
    })


import hashlib

import hashlib

@app.route('/login/management', methods=['POST'])
def management_login():
    username = request.json.get('username')
    password = request.json.get('password')
    verification_msg_manager = request.args.get('verification_msg_manager')
    
    # Check if the user exists
    user = User.query.filter_by(username=username, user_type='management').first()
    
    if user:
        # Hash the provided password using the same hash function and parameters used to hash the passwords in the database
        hashed_password = hashlib.sha256(password.encode()).hexdigest()
        
        # Compare the hashed password with the hashed password stored in the database
        if hashed_password == user.password:
            if user.is_active:  # Check if the user is active
                if user.is_verified:
                    # Set the user session variables
                    session['user_id'] = user.id
                    session['user_type'] = user.user_type
                    session['username'] = user.username
                    session['user_name'] = user.name
                    session['JWT Token'] = secrets.token_hex(16)
                    return jsonify({'status': 'success', 'redirect': url_for('dashboard'),'user_id':user.id,'email':user.email,'name':user.name})
                else:
                    message = 'Your account is not verified yet. Please check your email for the verification link.'
            else:
                message = 'Your account is not active. Please contact the administrator.'
        else:
            message = 'Invalid username or password'
    else:
        message = 'Invalid username or password'

    return jsonify({'status': 'error', 'message': message, 'verification_msg_manager': verification_msg_manager})

# @app.route('/get_recruiters', methods=['GET'])   
# def get_recruiters_list():
#     recruiters = User.query.filter_by(user_type='recruiter').all()
    
#     # Assuming you want to return a list of dictionaries containing user details
#     recruiters_list = []
#     for recruiter in recruiters:
#         recruiter_dict = {
#             'id': recruiter.id,
#             'username': recruiter.username,
#             'user_type': recruiter.user_type
#             # Add more fields if needed
#         }
#         recruiters_list.append(recruiter_dict)
    
#     return jsonify(recruiters_list)


@app.route('/get_recruiters', methods=['GET'])
def get_recruiters_list():
    recruiters = User.query.filter_by(user_type='recruiter').all()
    management = User.query.filter_by(user_type='management').all()

    # Extracting only usernames
    recruiter_usernames = [recruiter.username for recruiter in recruiters]
    management_usernames = [manager.username for manager in management]

    return jsonify({
        'recruiters': recruiter_usernames,
        'management': management_usernames
    })


@app.route('/get_recruiters_candidate', methods=['POST'])
def get_recruiters_candidate():
    data = request.json
    
    if not data or 'user_name' not in data:
        return jsonify({'status': 'error', 'message': 'Invalid input'}), 400
    
    username = data['user_name']
    print("username :",username)
    
    
    # Find the user with the given username who is either a recruiter or in management
    user = User.query.filter((User.name == username) & (User.user_type.in_(['recruiter', 'management']))).first()
    print("user data :" ,user)
    
    if user:
        # If the user is a recruiter, find all candidates linked with the user's username
        if user.user_type == 'recruiter':
            candidates = Candidate.query.filter_by(recruiter=username).all()
        # If the user is in management, find all candidates where recruiter matches the username
        elif user.user_type == 'management':
            candidates = Candidate.query.filter(
                (Candidate.recruiter == username) | (Candidate.management == username)
            ).all()
        else:
            return jsonify({'status': 'error', 'message': 'User type not authorized'}), 403
        
        # Prepare response data
        candidates_list = [
            {
                'id': candidate.id,
                'username': candidate.name,
                'status': candidate.status,
                'client':candidate.client,
                'profile': candidate.profile,
                'recruiter': candidate.recruiter,
                'management': candidate.management
            } 
            for candidate in candidates
        ]
        return jsonify(candidates_list)
    else:
        return jsonify({'status': 'error', 'message': 'User not found or not authorized'}), 404


############ Here i am converting the recruiters usernamr to name   ###################

# @app.route('/get_recruiters_candidate', methods=['POST'])
# def get_recruiters_candidate():
#     data = request.json
    
#     if not data or 'user_name' not in data:
#         return jsonify({'status': 'error', 'message': 'Invalid input'}), 400
    
#     username = data['user_name']
    
#     # Find the user with the given username who is either a recruiter or in management
#     user = User.query.filter((User.username == username) & (User.user_type.in_(['recruiter', 'management']))).first()
#     user_name=user.name
#     print("user_name :",user_name)
#     if user:
#         # If the user is a recruiter, find all candidates linked with the user's username
#         if user.user_type == 'recruiter':
#             candidates = Candidate.query.filter_by(recruiter=user_name).all()
#         # If the user is in management, find all candidates where recruiter matches the username
#         elif user.user_type == 'management':
#             candidates = Candidate.query.filter(
#                 (Candidate.recruiter == user_name) | (Candidate.management == user_name)
#             ).all()
#         else:
#             return jsonify({'status': 'error', 'message': 'User type not authorized'}), 403
        
#         # Prepare response data
#         candidates_list = [
#             {
#                 'id': candidate.id,
#                 'username': candidate.name,
#                 'status': candidate.status,
#                 'client':candidate.client,
#                 'profile': candidate.profile,
#                 'recruiter': candidate.recruiter,
#                 'management': candidate.management
#             } 
#             for candidate in candidates
#         ]
#         return jsonify(candidates_list)
#     else:
#         return jsonify({'status': 'error', 'message': 'User not found or not authorized'}), 404



from flask_mail import Message


def assign_candidates_notification(recruiter_email, new_recruiter_name, candidates_data):
    html_body = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                color: #333;
                line-height: 1.6;
                background-color: #f4f4f4;
                margin: 0;
                padding: 0;
            }}
            .container {{
                padding: 20px;
                margin: 20px auto;
                max-width: 600px;
                background-color: #ffffff;
                border: 1px solid #ddd;
                border-radius: 8px;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                text-align: center;
                font-size: 20px;
                border-radius: 8px 8px 0 0;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin-top: 10px;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #4CAF50;
                color: white;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            p {{
                margin: 10px 0;
            }}
            .footer {{
                margin-top: 20px;
                font-size: 12px;
                color: #777;
                text-align: center;
                border-top: 1px solid #ddd;
                padding-top: 10px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                Candidate Assignment Notification
            </div>
            <p>Hi {new_recruiter_name},</p>
            <p>Candidate data has been transferred to your ATS account.</p>
            <p>Please find the details below:</p>
            <table>
                <tr>
                    <th>Job ID</th>
                    <th>Client Name</th>
                    <th>Role/Profile</th>
                    <th>Candidate Name</th>
                    <th>Previous Recruiter</th>
                </tr>
                {candidates_data}
            </table>
            <p>Please check <b>ATS Dashboard</b> page for more details.</p>
            <p>Best Regards,</p>
            <p><b>Makonis Talent Track Pro Team</b></p>
        </div>
    </body>
    </html>
    """

    msg = Message(
        'Candidate Assignment Notification',
        sender=config.sender_email,
        recipients=[recruiter_email]
    )
    msg.html = html_body
    mail.send(msg)


def job_transfered_to_new_recruiter_notification(recruiter_email, new_recruiter_name, job_data):
    html_body = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                color: #333;
                line-height: 1.6;
                background-color: #f4f4f4;
                margin: 0;
                padding: 0;
            }}
            .container {{
                padding: 20px;
                margin: 20px auto;
                max-width: 600px;
                background-color: #ffffff;
                border: 1px solid #ddd;
                border-radius: 8px;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                text-align: center;
                font-size: 20px;
                border-radius: 8px 8px 0 0;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin-top: 10px;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #4CAF50;
                color: white;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            p {{
                margin: 10px 0;
            }}
            .footer {{
                margin-top: 20px;
                font-size: 12px;
                color: #777;
                text-align: center;
                border-top: 1px solid #ddd;
                padding-top: 10px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                New Job Requirement Assigned 
            </div>
            <p>Dear {new_recruiter_name},</p>
            <p>A new requirement has been assigned while transfering candidates to you.</p>
            <p> Please find the details below:</p>
            <table>
                <tr>
                    <th style="width: 20%;">Job ID</th>
                    <th style="width: 30%;">Client</th>
                    <th style="width: 30%;">Role/Profile</th>
                    <th style="width: 30%;">Location</th>
                </tr>
                {job_data}
            </table>
            <p>Please check in Job Listing page for more details.</p>
            <p>Regards,</p>
            <p><b>Makonis Talent Track Pro Team</b></p>
        </div>
    </body>
    </html>
    """

    msg = Message(
        f'New Job Requirement Assigned',
        sender=config.sender_email,
        recipients=[recruiter_email]
    )
    msg.html = html_body
    mail.send(msg)


@app.route('/assign_candidate_new_recuriter', methods=['POST'])
def assign_candidate_to_a_new_recruiter():
    data = request.json
    print("data ",data)

    try:
        candidates_data = ""
        current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
        new_recruiter_email = None
        new_recruiter_name = ""

        for candidate_data in data['candidates']:
            candidate_id = candidate_data.get('candidate_id')
            new_recruiter_name = candidate_data.get('new_recruiter')
            current_recruiter_username = candidate_data.get('current_recruiter')

            if not candidate_id or not new_recruiter_name or not current_recruiter_username:
                return jsonify({"error": "Candidate ID, new recruiter username, or current recruiter username not provided"}), 400

            # Fetch candidate details from the database
            candidate = Candidate.query.filter(
                Candidate.id == candidate_id,
                or_(
                    Candidate.recruiter == current_recruiter_username,
                    Candidate.management == current_recruiter_username
                )
            ).first()

            print("Candidate data:", candidate)  # Debugging: Check candidate details

            if candidate is None:
                return jsonify({"error": f"Candidate with ID {candidate_id} not found or not assigned to current recruiter/management {current_recruiter_username}"}), 404

            # Fetch job_id associated with the candidate
            job_id = candidate.job_id

            # Update the recruiter for the candidate
            candidate.recruiter = new_recruiter_name
            candidate.data_updated_date = current_datetime.date()
            candidate.data_updated_time = current_datetime.time()

            # Remove the current recruiter from the management field if it matches
            if candidate.management == current_recruiter_username:
                candidate.management = None

            # Append candidate details to the candidates_data string
            candidates_data += f"<tr><td>{candidate.job_id}</td><td>{candidate.client}</td><td>{candidate.profile}</td><td>{candidate.name}</td><td>{escape(current_recruiter_username)}</td></tr>"

            # Update recruiter column in job_post table
            job_post = JobPost.query.filter_by(id=job_id).first()
            if job_post:
                recruiters_list = job_post.recruiter.split(", ") if job_post.recruiter else []
                print("recruiters_list :", recruiters_list)

                if current_recruiter_username in recruiters_list:
                    # Check if there are any candidates still linked with this job post and the current recruiter
                    linked_candidates = Candidate.query.filter(
                        Candidate.job_id == job_id,
                        Candidate.recruiter == current_recruiter_username
                    ).count()
                    if linked_candidates == 0:
                        recruiters_list.remove(current_recruiter_username)

                if new_recruiter_name not in recruiters_list:
                    recruiters_list.append(new_recruiter_name)
                    job_post.recruiter = ", ".join(recruiters_list)

                    # Add new notification for the new recruiter
                    new_notification = Notification(job_post_id=job_id, recruiter_name=new_recruiter_name)
                    db.session.add(new_notification)
                    new_notification.num_notification = 1

                    # Fetch new recruiter details
                    new_recruiter = User.query.filter_by(name=new_recruiter_name).first()  # Using `username` to fetch recruiter
                    print("New Recruiter Name: :", new_recruiter)
                    if new_recruiter:
                        new_recruiter_name = new_recruiter.name
                        new_recruiter_email = new_recruiter.email
                        print("New Recruiter Email: ", new_recruiter_email)  # Debugging: Print email
                        print(f"New Recruiter Name: {new_recruiter_name}")  # Debugging: Print name
                        job_data = f"""
                        <tr>
                            <td>{job_post.id}</td>
                            <td>{job_post.client}</td>
                            <td>{job_post.role}</td>
                            <td>{job_post.location}</td>
                        </tr>
                        """
                        job_transfered_to_new_recruiter_notification(new_recruiter_email, new_recruiter_name, job_data)
                    else:
                        print("New Recruiter not found in the database.")  # Debugging: Recruiter not found

        # Commit changes to the database
        db.session.commit()

        # Send notification email to the new recruiter for candidate assignment
        if candidates_data and new_recruiter:
            print(f"Sending Candidate Assignment Notification to {new_recruiter_email} ({new_recruiter_name})")  # Debugging: Confirm email sending
            assign_candidates_notification(new_recruiter_email, new_recruiter_name, candidates_data)
        else:
            print("Candidates data or new recruiter email is missing; email notification not sent.")  # Debugging: Missing data

        return jsonify({'status': 'success', "message": "Candidates assigned successfully."})
    except Exception as e:
        db.session.rollback()
        print(f"Error assigning candidates: {str(e)}")  # Debugging: Print error
        return jsonify({'status': 'error', "error": f"Error assigning candidates: {str(e)}"}), 500


    

# @app.route('/assign_candidate_new_recuriter', methods=['POST'])
# def assign_candidate_to_a_new_recruiter():
#     data = request.json

#     try:
#         candidates_data = ""
#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#         new_recruiter = None
#         new_recruiter_name = ""

#         for candidate_data in data['candidates']:
#             candidate_id = candidate_data.get('candidate_id')
#             new_recruiter_username = candidate_data.get('new_recruiter')
#             current_recruiter_username = candidate_data.get('current_recruiter')

#             if not candidate_id or not new_recruiter_username or not current_recruiter_username:
#                 return jsonify({"error": "Candidate ID, new recruiter username, or current recruiter username not provided"}), 400

#             # Fetch candidate details from the database
#             candidate = Candidate.query.filter(
#                 Candidate.id == candidate_id,
#                 or_(
#                     Candidate.recruiter == current_recruiter_username,
#                     Candidate.management == current_recruiter_username
#                 )
#             ).first()
#             print("Candidate data :",candidate)
#             if candidate is None:
#                 return jsonify({"error": f"Candidate with ID {candidate_id} not found or not assigned to current recruiter/management {current_recruiter_username}"}), 404

#             # Fetch job_id associated with the candidate
#             job_id = candidate.job_id

#             # Update the recruiter for the candidate
#             candidate.recruiter = new_recruiter_username
#             candidate.data_updated_date = current_datetime.date()
#             candidate.data_updated_time = current_datetime.time()

#             # Remove the current recruiter from the management field if it matches
#             if candidate.management == current_recruiter_username:
#                 candidate.management = None

#             # Append candidate details to the candidates_data string
#             candidates_data += f"<tr><td>{candidate.job_id}</td><td>{candidate.client}</td><td>{candidate.profile}</td><td>{candidate.name}</td><td>{escape(current_recruiter_username)}</td></tr>"

#             # Update recruiter column in job_post table
#             job_post = JobPost.query.filter_by(id=job_id).first()
#             if job_post:
#                 recruiters_list = job_post.recruiter.split(", ") if job_post.recruiter else []

#                 if current_recruiter_username in recruiters_list:
#                     # Check if there are any candidates still linked with this job post and the current recruiter
#                     linked_candidates = Candidate.query.filter(
#                         Candidate.job_id == job_id,
#                         Candidate.recruiter == current_recruiter_username
#                     ).count()
#                     if linked_candidates == 0:
#                         recruiters_list.remove(current_recruiter_username)

#                 if new_recruiter_username not in recruiters_list:
#                     recruiters_list.append(new_recruiter_username)
#                     job_post.recruiter = ", ".join(recruiters_list)

#                     # Add new notification for the new recruiter
#                     new_notification = Notification(job_post_id=job_id, recruiter_name=new_recruiter_username)
#                     db.session.add(new_notification)
#                     new_notification.num_notification = 1

#                     # Send new job post notification to the new recruiter
#                     new_recruiter = User.query.filter_by(name=new_recruiter_username).first()
#                     new_recruiter_name = new_recruiter.name if new_recruiter else "New Recruiter"
#                     job_data = f"""
#                     <tr>
#                         <td>{job_post.id}</td>
#                         <td>{job_post.client}</td>
#                         <td>{job_post.role}</td>
#                         <td>{job_post.location}</td>
#                     </tr>
#                     """
#                     job_transfered_to_new_recruiter_notification(new_recruiter.email, new_recruiter_name, job_data)
                    

#         # Commit changes to the database
#         db.session.commit()

#         # Send notification email to the new recruiter for candidate assignment
#         if candidates_data and new_recruiter:
#             assign_candidates_notification(new_recruiter.email, new_recruiter_name, candidates_data)

#         return jsonify({'status': 'success', "message": "Candidates assigned successfully."})
#     except Exception as e:
#         db.session.rollback()
#         return jsonify({'status': 'error', "error": f"Error assigning candidates: {str(e)}"}), 500


############# Here i am converting the recruiters username to name ###################

# @app.route('/assign_candidate_new_recuriter', methods=['POST'])
# def assign_candidate_to_a_new_recruiter():
#     data = request.json

#     if not data or 'candidates' not in data:
#         return jsonify({"error": "Invalid input"}), 400

#     candidates_data = ""
#     current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#     new_recruiter_name = ""

#     try:
#         for candidate_data in data['candidates']:
#             candidate_id = candidate_data.get('candidate_id')
#             new_recruiter_username = candidate_data.get('new_recruiter')
#             current_recruiter_username = candidate_data.get('current_recruiter')

#             if not candidate_id or not new_recruiter_username or not current_recruiter_username:
#                 return jsonify({"error": "Candidate ID, new recruiter username, or current recruiter username not provided"}), 400

#             # Fetch the current recruiter's details from the database
#             current_recruiter = User.query.filter_by(username=current_recruiter_username).first()
#             if not current_recruiter:
#                 return jsonify({"error": f"Current recruiter with username {current_recruiter_username} not found"}), 404
#             current_recruiter_name = current_recruiter.name

#             # Fetch the new recruiter's details from the database
#             new_recruiter = User.query.filter_by(username=new_recruiter_username).first()
#             if not new_recruiter:
#                 return jsonify({"error": f"New recruiter with username {new_recruiter_username} not found"}), 404
#             new_recruiter_name = new_recruiter.name

#             # Fetch candidate details from the database
#             candidate = Candidate.query.filter(
#                 Candidate.id == candidate_id,
#                 or_(
#                     Candidate.recruiter == current_recruiter_name,
#                     Candidate.management == current_recruiter_name
#                 )
#             ).first()

#             if candidate is None:
#                 return jsonify({"error": f"Candidate with ID {candidate_id} not found or not assigned to current recruiter/management {current_recruiter_name}"}), 404

#             # Fetch job_id associated with the candidate
#             job_id = candidate.job_id

#             # Update the recruiter for the candidate
#             candidate.recruiter = new_recruiter_name
#             candidate.data_updated_date = current_datetime.date()
#             candidate.data_updated_time = current_datetime.time()

#             # Remove the current recruiter from the management field if it matches
#             if candidate.management == current_recruiter_name:
#                 candidate.management = None

#             # Append candidate details to the candidates_data string
#             candidates_data += f"<tr><td>{candidate.job_id}</td><td>{candidate.client}</td><td>{candidate.profile}</td><td>{candidate.name}</td><td>{escape(current_recruiter_name)}</td></tr>"

#             # Update recruiter column in job_post table
#             job_post = JobPost.query.filter_by(id=job_id).first()
#             if job_post:
#                 recruiters_list = job_post.recruiter.split(", ") if job_post.recruiter else []
#                 print("recruiters_list  :",recruiters_list)

#                 if current_recruiter_name in recruiters_list:
#                     # Check if there are any candidates still linked with this job post and the current recruiter
#                     linked_candidates = Candidate.query.filter(
#                         Candidate.job_id == job_id,
#                         Candidate.recruiter == current_recruiter_name
#                     ).count()
#                     if linked_candidates == 0:
#                         recruiters_list.remove(current_recruiter_name)

#                 if new_recruiter_name not in recruiters_list:
#                     recruiters_list.append(new_recruiter_name)
#                     job_post.recruiter = ", ".join(recruiters_list)

#                 # Commit changes to the job post recruiter list
#                 db.session.add(job_post)

#                 # Add new notification for the new recruiter
#                 new_notification = Notification(job_post_id=job_id, recruiter_name=new_recruiter_name)
#                 db.session.add(new_notification)
#                 new_notification.num_notification = 1

#                 # Send new job post notification to the new recruiter
#                 job_data = f"""
#                 <tr>
#                     <td>{job_post.id}</td>
#                     <td>{job_post.client}</td>
#                     <td>{job_post.role}</td>
#                     <td>{job_post.location}</td>
#                 </tr>
#                 """
#                 job_transfered_to_new_recruiter_notification(new_recruiter.email, new_recruiter_name, job_data)

#         # Commit changes to the database
#         db.session.commit()

#         # Send notification email to the new recruiter for candidate assignment
#         if candidates_data and new_recruiter:
#             assign_candidates_notification(new_recruiter.email, new_recruiter_name, candidates_data)

#         return jsonify({'status': 'success', "message": "Candidates assigned successfully."})
#     except Exception as e:
#         db.session.rollback()
#         return jsonify({'status': 'error', "error": f"Error assigning candidates: {str(e)}"}), 500





from flask import Flask, jsonify, request, Response
from datetime import date
import json
from sqlalchemy import and_


# # Define your utility functions
# def get_common_candidate_data(candidate):
#     return {
#         'id': candidate.id,
#         'job_id': candidate.job_id,
#         'name': candidate.name,
#         'mobile': candidate.mobile,
#         'email': candidate.email,
#         'client': candidate.client,
#         'current_company': candidate.current_company,
#         'position': candidate.position,
#         'profile': candidate.profile,
#         'current_job_location': candidate.current_job_location,
#         'preferred_job_location': candidate.preferred_job_location,
#         'qualifications': candidate.qualifications,
#         'experience': candidate.experience,
#         'relevant_experience': candidate.relevant_experience,
#         'current_ctc': candidate.current_ctc,
#         'expected_ctc': candidate.expected_ctc,
#         'notice_period': candidate.notice_period,
#         'linkedin': candidate.linkedin_url,
#         'reason_for_job_change': candidate.reason_for_job_change,
#         'holding_offer': candidate.holding_offer,
#         'recruiter': candidate.recruiter,
#         'management': candidate.management,
#         'status': candidate.status,
#         'remarks': candidate.remarks,
#         'skills': candidate.skills,
#         'resume': candidate.resume if candidate.resume is not None else "",
#         'serving_notice_period': candidate.notice_period,
#         'period_of_notice': candidate.period_of_notice,
#         'last_working_date': candidate.last_working_date,
#         'total_offers': candidate.total,
#         'highest_package_in_lpa': candidate.package_in_lpa,
#         'buyout': candidate.buyout,
#         'date_created': candidate.date_created,
#         'time_created': candidate.time_created,
#         'data_updated_date': candidate.data_updated_date,
#         'data_updated_time': candidate.data_updated_time,
#         'resume_present': candidate.resume_present
#     }

# def get_common_job_data(job):
#     return {
#         'id': job.id,
#         'client': job.client,
#         'experience_min': job.experience_min,
#         'experience_max': job.experience_max,
#         'budget_min': job.budget_min,
#         'budget_max': job.budget_max,
#         'location': job.location,
#         'shift_timings': job.shift_timings,
#         'notice_period': job.notice_period,
#         'role': job.role,
#         'detailed_jd': job.detailed_jd,
#         'jd_pdf': job.jd_pdf,
#         'mode': job.mode,
#         'recruiter': job.recruiter,
#         'management': job.management,
#         'date_created': job.date_created,
#         'time_created': job.time_created,
#         'job_status': job.job_status,
#         'job_type': job.job_type,
#         'contract_in_months': job.contract_in_months,
#         'skills': job.skills,
#         'notification': job.notification,
#         'data_updated_date': job.data_updated_date,
#         'data_updated_time': job.data_updated_time,
#         'jd_pdf_present': job.jd_pdf_present,
#         'no_of_positions': job.no_of_positions
#     }

# def date_handler(obj):
#     if isinstance(obj, datetime):
#         return obj.strftime('%H:%M:%S')  # Format time without microseconds
#     elif isinstance(obj, date):
#         return obj.isoformat()
#     else:
#         return None

# def fetch_recruiter_data(user_id):
#     with app.app_context():
#         recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
#         if recruiter:
#             recruiters = recruiter.username.split(',')
#             candidates = Candidate.query.filter(
#                 and_(Candidate.recruiter == recruiter.name, Candidate.reference.is_(None))
#             ).order_by(
#                 desc(case((Candidate.data_updated_date != None, Candidate.data_updated_date), else_=Candidate.date_created)),
#                 desc(case((Candidate.data_updated_time != None, Candidate.data_updated_time), else_=Candidate.time_created)),
#                 desc(Candidate.id)
#             ).all()
#             jobs_query = JobPost.query.filter(
#                 or_(*[JobPost.recruiter.like(f"%{recruiter}%") for recruiter in recruiters])
#             )
#             jobs = jobs_query.all()
#             return recruiter, candidates, jobs
#         return None, [], []

# def fetch_management_data():
#     with app.app_context():
#         users = User.query.all()
#         candidates = Candidate.query.filter(Candidate.reference.is_(None)).order_by(
#             desc(case((Candidate.data_updated_date != None, Candidate.data_updated_date), else_=Candidate.date_created)),
#             desc(case((Candidate.data_updated_time != None, Candidate.data_updated_time), else_=Candidate.time_created)),
#             desc(Candidate.id)
#         ).all()
#         jobs = JobPost.query.all()
#         return users, candidates, jobs

# def fetch_other_user_data(user_name):
#     with app.app_context():
#         return Candidate.query.filter_by(recruiter=user_name).order_by(
#             desc(case((Candidate.data_updated_date != None, Candidate.data_updated_date), else_=Candidate.date_created)),
#             desc(case((Candidate.data_updated_time != None, Candidate.data_updated_time), else_=Candidate.time_created)),
#             desc(Candidate.id)
#         ).all()

# @app.route('/dashboard', methods=['POST'])
# def dashboard():
#     data = request.json
#     user_id = data.get('user_id')

#     if user_id is None:
#         return jsonify({"message": "User ID missing"}), 400

#     user = User.query.filter_by(id=user_id).first()
#     if user is None:
#         return jsonify({"message": "User not found"}), 404

#     user_type = user.user_type
#     user_name = user.username
#     name = user.name

#     response_data = {}

#     with ThreadPoolExecutor() as executor:
#         if user_type == 'recruiter':
#             future_recruiter_data = executor.submit(fetch_recruiter_data, user_id)
#             recruiter, candidates, jobs = future_recruiter_data.result()
#             if recruiter:
#                 response_data = {
#                     'user': {
#                         'id': recruiter.id,
#                         'name': recruiter.name,
#                         'user_type': recruiter.user_type,
#                         'email': recruiter.email
#                     },
#                     'user_type': user_type,
#                     'user_name': user_name,
#                     'name': name,
#                     'candidates': [get_common_candidate_data(candidate) for candidate in candidates],
#                     'jobs': [get_common_job_data(job) for job in jobs],
#                     'edit_candidate_message': data.get('edit_candidate_message'),
#                     'page_no': data.get('page_no'),
#                 }
#             else:
#                 return jsonify({"message": "Recruiter not found"}), 404

#         elif user_type == 'management':
#             future_management_data = executor.submit(fetch_management_data)
#             users, candidates, jobs = future_management_data.result()
#             response_data = {
#                 'users': [{'id': user.id, 'name': user.name, 'user_type': user.user_type, 'email': user.email} for user in users],
#                 'user_type': user_type,
#                 'user_name': user_name,
#                 'candidates': [get_common_candidate_data(candidate) for candidate in candidates],
#                 'jobs': [get_common_job_data(job) for job in jobs],
#                 'signup_message': data.get('signup_message'),
#                 'job_message': data.get('job_message'),
#                 'page_no': data.get('page_no'),
#                 'edit_candidate_message': data.get('edit_candidate_message'),
#             }

#         else:
#             future_other_user_data = executor.submit(fetch_other_user_data, user_name)
#             candidates = future_other_user_data.result()
#             response_data = {
#                 'user': {
#                     'id': user.id,
#                     'name': user.name,
#                     'user_type': user.user_type,
#                     'email': user.email
#                 },
#                 'user_type': user_type,
#                 'user_name': user_name,
#                 'candidates': [get_common_candidate_data(candidate) for candidate in candidates],
#             }

#     response_json = json.dumps(response_data, default=date_handler)  # Use date_handler to handle date serialization
#     return Response(response=response_json, status=200, mimetype='application/json')

########################################################################################################

def date_handler(obj):
    if isinstance(obj, datetime):
        return obj.strftime('%H:%M:%S')  # Format time without microseconds
    elif isinstance(obj, date):
        return obj.isoformat()
    else:
        return None

def get_common_candidate_data(candidate):
    return {
        'id': candidate.id,
        'job_id': candidate.job_id,
        'name': candidate.name,
        'mobile': candidate.mobile,
        'email': candidate.email,
        'client': candidate.client,
        'current_company': candidate.current_company,
        'position': candidate.position,
        'profile': candidate.profile,
        'current_job_location': candidate.current_job_location,
        'preferred_job_location': candidate.preferred_job_location,
        'qualifications': candidate.qualifications,
        'experience': candidate.experience,
        'relevant_experience': candidate.relevant_experience,
        'current_ctc': candidate.current_ctc,
        'expected_ctc': candidate.expected_ctc,
        'notice_period': candidate.notice_period,
        'linkedin': candidate.linkedin_url,
        'reason_for_job_change': candidate.reason_for_job_change,
        'holding_offer': candidate.holding_offer,
        'recruiter': candidate.recruiter,
        'management': candidate.management,
        'status': candidate.status,
        'remarks': candidate.remarks,
        'skills': candidate.skills,
        'resume': candidate.resume if candidate.resume is not None else "",
        'serving_notice_period': candidate.notice_period,
        'period_of_notice': candidate.period_of_notice,
        'last_working_date': candidate.last_working_date,
        'total_offers': candidate.total,
        'highest_package_in_lpa': candidate.package_in_lpa,
        'buyout': candidate.buyout,
        'date_created': candidate.date_created.isoformat() if candidate.date_created else None,
        'time_created': candidate.time_created.strftime('%H:%M:%S') if candidate.time_created else None,
        'data_updated_date': candidate.data_updated_date.isoformat() if candidate.data_updated_date else None,
        'data_updated_time': candidate.data_updated_time.strftime('%H:%M:%S') if candidate.data_updated_time else None,
        'resume_present': candidate.resume_present
    }

def get_common_job_data(job):
    return {
        'id': job.id,
        'client': job.client,
        'experience_min': job.experience_min,
        'experience_max': job.experience_max,
        'budget_min': job.budget_min,
        'budget_max': job.budget_max,
        'location': job.location,
        'shift_timings': job.shift_timings,
        'notice_period': job.notice_period,
        'role': job.role,
        'detailed_jd': job.detailed_jd,
        'jd_pdf': job.jd_pdf,
        'mode': job.mode,
        'recruiter': job.recruiter,
        'management': job.management,
        'date_created': job.date_created.isoformat() if job.date_created else None,
        'time_created': job.time_created.strftime('%H:%M:%S') if job.time_created else None,
        'job_status': job.job_status,
        'job_type': job.job_type,
        'contract_in_months': job.contract_in_months,
        'skills': job.skills,
        'notification': job.notification,
        'data_updated_date': job.data_updated_date.isoformat() if job.data_updated_date else None,
        'data_updated_time': job.data_updated_time.strftime('%H:%M:%S') if job.data_updated_time else None,
        'jd_pdf_present': job.jd_pdf_present,
        'no_of_positions': job.no_of_positions
    }

@app.route('/dashboard', methods=['POST'])
def dashboard():
    data = request.json
    user_id = data.get('user_id')

    if user_id is None:
        return jsonify({"message": "User ID missing"}), 400

    user = User.query.filter_by(id=user_id).first()
    if user is None:
        return jsonify({"message": "User not found"}), 404

    user_type = user.user_type
    user_name = user.username
    name = user.name

    # Conditional ordering
    conditional_order_date = case(
        (Candidate.data_updated_date != None, Candidate.data_updated_date),
        else_=Candidate.date_created
    )
    conditional_order_time = case(
        (Candidate.data_updated_time != None, Candidate.data_updated_time),
        else_=Candidate.time_created
    )

    response_data = {}
    
    if user_type == 'recruiter':
        recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
        if recruiter is None:
            return jsonify({"message": "Recruiter not found"}), 404

        recruiters = recruiter.username.split(',')

        candidates = Candidate.query.filter(and_(Candidate.recruiter == recruiter.name, Candidate.reference.is_(None)))\
            .order_by(
                desc(conditional_order_date),
                desc(conditional_order_time),
                desc(Candidate.id)
            ).all()

        jobs_query = JobPost.query.filter(
            or_(*[JobPost.recruiter.like(f"%{recruiter}%") for recruiter in recruiters])
        )
        jobs = jobs_query.all()

        response_data = {
            'user': {
                'id': recruiter.id,
                'name': recruiter.name,
                'user_type': recruiter.user_type,
                'email': recruiter.email
            },
            'user_type': user_type,
            'user_name': user_name,
            'name': name,
            'candidates': [get_common_candidate_data(candidate) for candidate in candidates],
            'jobs': [get_common_job_data(job) for job in jobs],
            'edit_candidate_message': data.get('edit_candidate_message'),
            'page_no': data.get('page_no'),
        }

    elif user_type == 'management':
        users = User.query.all()

        candidates = Candidate.query.filter(Candidate.reference.is_(None))\
            .order_by(
                desc(conditional_order_date),
                desc(conditional_order_time),
                desc(Candidate.id)
            ).all()

        jobs = JobPost.query.all()
        
        response_data = {
            'users': [{'id': user.id, 'name': user.name, 'user_type': user.user_type, 'email': user.email} for user in users],
            'user_type': user_type,
            'user_name': user_name,
            'candidates': [get_common_candidate_data(candidate) for candidate in candidates],
            'jobs': [get_common_job_data(job) for job in jobs],
            'signup_message': data.get('signup_message'),
            'job_message': data.get('job_message'),
            'page_no': data.get('page_no'),
            'edit_candidate_message': data.get('edit_candidate_message'),
        }

    else:
        candidates = Candidate.query.filter_by(recruiter=user.name)\
            .order_by(
                desc(conditional_order_date),
                desc(conditional_order_time),
                desc(Candidate.id)
            ).all()

        response_data = {
            'user': {
                'id': user.id,
                'name': user.name,
                'user_type': user.user_type,
                'email': user.email
            },
            'user_type': user_type,
            'user_name': user_name,
            'candidates': [get_common_candidate_data(candidate) for candidate in candidates],
        }

    response_json = json.dumps(response_data, default=date_handler)
    return Response(response=response_json, status=200, mimetype='application/json')

###########################################################################################################


# def date_handler(obj):
#     if isinstance(obj, datetime):
#         return obj.strftime('%H:%M:%S')  # Format time without microseconds
#     elif isinstance(obj, date):
#         return obj.isoformat()
#     else:
#         return None
        
# @app.route('/dashboard', methods=['POST'])
# def dashboard():
#     print("Dashboard  :")
#     data = request.json
#     print(data)  # Just to verify if data is received properly
#     edit_candidate_message = data.get('edit_candidate_message')
#     page_no = data.get('page_no')
#     candidate_message = data.get('candidate_message')
#     signup_message = data.get('signup_message')
#     job_message = data.get('job_message')
#     update_candidate_message = data.get('update_candidate_message')
#     delete_message = data.get("delete_message")

#     user_id = data.get('user_id')
#     if user_id is None:
#         return jsonify({"message": "User ID missing"}), 400

#     user = User.query.filter_by(id=user_id).first()
#     if user is None:
#         return jsonify({"message": "User not found"}), 404

#     user_type = user.user_type
#     user_name = user.username
#     name = user.name
#     print("main name :",name)

#     response_data = {}

#     # Define case statements for conditional ordering
#     conditional_order_date = case(
#         (Candidate.data_updated_date != None, Candidate.data_updated_date),
#         (Candidate.date_created != None, Candidate.date_created),
#         else_=Candidate.date_created
#     )

#     conditional_order_time = case(
#         (Candidate.data_updated_time != None, Candidate.data_updated_time),
#         (Candidate.time_created != None, Candidate.time_created),
#         else_=Candidate.time_created
#     )

#     if user_type == 'recruiter':
#         recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
#         if recruiter is None:
#             return jsonify({"message": "Recruiter not found"}), 404

#         user_name = recruiter.username
#         recruiters = user_name.split(',')  # Splitting the recruiter usernames separated by commas

#         print("Recruiter usernames:", recruiters)  # Debugging statement to check the recruiter usernames

#         candidates = Candidate.query.filter(and_(Candidate.recruiter == recruiter.name, Candidate.reference.is_(None)))\
#             .order_by(
#                 desc(conditional_order_date),
#                 desc(conditional_order_time),
#                 desc(Candidate.id)  # Ensure newer candidates appear first if dates are equal
#             )\
#             .all()

#         # for candidate in candidates:
#         #     print(f"Candidate ID: {candidate.id}, Time Created: {candidate.time_created}")

#         jobs_query = JobPost.query.filter(
#             or_(*[JobPost.recruiter.like(f"%{recruiter}%") for recruiter in recruiters])
#          )
#         jobs = jobs_query.all()

#         print("Jobs retrieved:", jobs)  # Debugging statement to check the jobs retrieved

#         response_data = {
#             'user': {
#                 'id': recruiter.id,
#                 'name': recruiter.name,
#                 'user_type': recruiter.user_type,
#                 'email': recruiter.email
#             },
#             'user_type': user_type,
#             'user_name': user_name,
#             'name':name,
#             'candidates': [{
#                 'id': candidate.id,
#                 'job_id': candidate.job_id,
#                 'name': candidate.name,
#                 'mobile': candidate.mobile,
#                 'email': candidate.email,
#                 'client': candidate.client,
#                 'current_company': candidate.current_company,
#                 'position': candidate.position,
#                 'profile': candidate.profile,
#                 'current_job_location': candidate.current_job_location,
#                 'preferred_job_location': candidate.preferred_job_location,
#                 'qualifications': candidate.qualifications,
#                 'experience': candidate.experience,
#                 'relevant_experience': candidate.relevant_experience,
#                 'current_ctc': candidate.current_ctc,
#                 'expected_ctc': candidate.expected_ctc,
#                 'notice_period': candidate.notice_period,
#                 'linkedin': candidate.linkedin_url,
#                 'reason_for_job_change':candidate.reason_for_job_change,
#                 'holding_offer': candidate.holding_offer,
#                 'recruiter': candidate.recruiter,
#                 'management': candidate.management,
#                 'status': candidate.status,
#                 'remarks': candidate.remarks,
#                 'skills': candidate.skills,
#                 'resume': candidate.resume if candidate.resume is not None else "",
#                 'serving_notice_period': candidate.notice_period,
#                 'period_of_notice': candidate.period_of_notice,
#                 'last_working_date': candidate.last_working_date,
#                 'total_offers': candidate.total,
#                 'highest_package_in_lpa': candidate.package_in_lpa,
#                 'buyout': candidate.buyout,
#                 'date_created': candidate.date_created.isoformat() if candidate.date_created else None,
#                 'time_created': candidate.time_created.strftime('%H:%M:%S') if candidate.time_created else None,
#                 'data_updated_date': candidate.data_updated_date.isoformat() if candidate.data_updated_date else None,
#                 'data_updated_time': candidate.data_updated_time.strftime('%H:%M:%S') if candidate.data_updated_time else None,
#                 'resume_present': candidate.resume_present
#             } for candidate in candidates],
#             'jobs': [{
#                 'id': job.id,
#                 'client': job.client,
#                 'experience_min': job.experience_min,
#                 'experience_max': job.experience_max,
#                 'budget_min': job.budget_min,
#                 'budget_max': job.budget_max,
#                 'location': job.location,
#                 'shift_timings': job.shift_timings,
#                 'notice_period': job.notice_period,
#                 'role': job.role,
#                 'detailed_jd': job.detailed_jd,
#                 'jd_pdf': job.jd_pdf,
#                 'mode': job.mode,
#                 'recruiter': job.recruiter,
#                 'management': job.management,
#                 'date_created': job.date_created.isoformat() if job.date_created else None,
#                 'time_created': job.time_created.strftime('%H:%M:%S') if job.time_created else None,
#                 'job_status': job.job_status,
#                 'job_type': job.job_type,
#                 'contract_in_months': job.contract_in_months,
#                 'skills': job.skills,
#                 'notification': job.notification,
#                 'data_updated_date': job.data_updated_date.isoformat() if job.data_updated_date else None,
#                 'data_updated_time': job.data_updated_time.strftime('%H:%M:%S') if job.data_updated_time else None,
#                 'jd_pdf_present': job.jd_pdf_present,
#                 'no_of_positions' :job.no_of_positions
#             } for job in jobs],
#             'edit_candidate_message': edit_candidate_message,
#             'page_no': page_no,
#         }
        
#     elif user_type == 'management':
#         # Define case statements for conditional ordering
#         conditional_order_date = case(
#         (Candidate.date_created != None, Candidate.date_created),
#         else_=Candidate.date_created
#         )

#         conditional_order_time = case(
#         (Candidate.time_created != None, Candidate.time_created),
#         else_=Candidate.time_created
#         )
#         users = User.query.all()
        
#         candidates = Candidate.query.filter(Candidate.reference.is_(None))\
#             .order_by(
#                 desc(conditional_order_date),
#                 desc(conditional_order_time),
#                 desc(Candidate.id)
#             )\
#             .all()

#         # for candidate in candidates:
#         #     print(f"Candidate ID: {candidate.id}, Time Created: {candidate.time_created}")

#         jobs = JobPost.query.all()
        
#         response_data = {
#             'users': [{
#                 'id': user.id,
#                 'name': user.name,
#                 'user_type': user.user_type,
#                 'email': user.email
#             } for user in users],
#             'user_type': user_type,
#             'user_name': user_name,
#             'candidates': [{
#                 'id': candidate.id,
#                 'job_id': candidate.job_id,
#                 'name': candidate.name,
#                 'mobile': candidate.mobile,
#                 'email': candidate.email,
#                 'client': candidate.client,
#                 'current_company': candidate.current_company,
#                 'position': candidate.position,
#                 'profile': candidate.profile,
#                 'current_job_location': candidate.current_job_location,
#                 'preferred_job_location': candidate.preferred_job_location,
#                 'qualifications': candidate.qualifications,
#                 'experience': candidate.experience,
#                 'relevant_experience': candidate.relevant_experience,
#                 'current_ctc': candidate.current_ctc,
#                 'expected_ctc': candidate.expected_ctc,
#                 'notice_period': candidate.notice_period,
#                 'reason_for_job_change':candidate.reason_for_job_change,
#                 'linkedin': candidate.linkedin_url,
#                 'holding_offer': candidate.holding_offer,
#                 'recruiter': candidate.recruiter,
#                 'management': candidate.management,
#                 'status': candidate.status,
#                 'remarks': candidate.remarks,
#                 'skills': candidate.skills,
#                 'resume': candidate.resume if candidate.resume is not None else "",
#                 'serving_notice_period': candidate.notice_period,
#                 'period_of_notice': candidate.period_of_notice,
#                 'last_working_date': candidate.last_working_date,
#                 'total_offers': candidate.total,
#                 'highest_package_in_lpa': candidate.package_in_lpa,
#                 'buyout': candidate.buyout,
#                 'date_created': candidate.date_created.isoformat() if candidate.date_created else None,
#                 'time_created': candidate.time_created.strftime('%H:%M:%S') if candidate.time_created else None,
#                 'data_updated_date': candidate.data_updated_date.isoformat() if candidate.data_updated_date else None,
#                 'data_updated_time': candidate.data_updated_time.strftime('%H:%M:%S') if candidate.data_updated_time else None,
#                 'resume_present': candidate.resume_present
#             } for candidate in candidates],
#             'jobs': [{
#                 'id': job.id,
#                 'client': job.client,
#                 'experience_min': job.experience_min,
#                 'experience_max': job.experience_max,
#                 'budget_min': job.budget_min,
#                 'budget_max': job.budget_max,
#                 'location': job.location,
#                 'shift_timings': job.shift_timings,
#                 'notice_period': job.notice_period,
#                 'role': job.role,
#                 'detailed_jd': job.detailed_jd,
#                 'jd_pdf': job.jd_pdf,
#                 'mode': job.mode,
#                 'recruiter': job.recruiter,
#                 'management': job.management,
#                 'date_created': job.date_created.isoformat() if job.date_created else None,
#                 'time_created': job.time_created.strftime('%H:%M:%S') if job.time_created else None,
#                 'job_status': job.job_status,
#                 'job_type': job.job_type,
#                 'contract_in_months': job.contract_in_months,
#                 'skills': job.skills,
#                 'notification': job.notification,
#                 'data_updated_date': job.data_updated_date.isoformat() if job.data_updated_date else None,
#                 'data_updated_time': job.data_updated_time.strftime('%H:%M:%S') if job.data_updated_time else None,
#                 'jd_pdf_present': job.jd_pdf_present,
#                  'no_of_positions' :job.no_of_positions
#             } for job in jobs],
#             'signup_message': signup_message,
#             'job_message': job_message,
#             'page_no': page_no,
#             'edit_candidate_message': edit_candidate_message
#         }
        
#     else:
#         candidates = Candidate.query.filter_by(recruiter=user.name)\
#             .order_by(
#                 desc(conditional_order_date),
#                 desc(conditional_order_time),
#                 desc(Candidate.id)
#             )\
#             .all()

#         # for candidate in candidates:
#         #     print(f"Candidate ID: {candidate.id}, Time Created: {candidate.time_created}")
        
#         response_data = {
#             'user': {
#                 'id': user.id,
#                 'name': user.name,
#                 'user_type': user.user_type,
#                 'email': user.email
#             },
#             'user_type': user_type,
#             'user_name': user_name,
#             'candidates': [{
#                 'id': candidate.id,
#                 'job_id': candidate.job_id,
#                 'name': candidate.name,
#                 'mobile': candidate.mobile,
#                 'email': candidate.email,
#                 'client': candidate.client,
#                 'current_company': candidate.current_company,
#                 'position': candidate.position,
#                 'profile': candidate.profile,
#                 'current_job_location': candidate.current_job_location,
#                 'preferred_job_location': candidate.preferred_job_location,
#                 'qualifications': candidate.qualifications,
#                 'experience': candidate.experience,
#                 'relevant_experience': candidate.relevant_experience,
#                 'current_ctc': candidate.current_ctc,
#                 'expected_ctc': candidate.expected_ctc,
#                 'notice_period': candidate.notice_period,
#                 'linkedin': candidate.linkedin_url,
#                 'reason_for_job_change':candidate.reason_for_job_change,
#                 'holding_offer': candidate.holding_offer,
#                 'recruiter': candidate.recruiter,
#                 'management': candidate.management,
#                 'status': candidate.status,
#                 'remarks': candidate.remarks,
#                 'skills': candidate.skills,
#                 'resume': candidate.resume if candidate.resume is not None else "",
#                 'serving_notice_period': candidate.notice_period,
#                 'period_of_notice': candidate.period_of_notice,
#                 'last_working_date': candidate.last_working_date,
#                 'buyout': candidate.buyout,
#                 'total_offers': candidate.total,
#                 'highest_package_in_lpa': candidate.package_in_lpa,
#                 'date_created': candidate.date_created.isoformat() if candidate.date_created else None,
#                 'time_created': candidate.time_created.strftime('%H:%M:%S') if candidate.time_created else None,
#                 'data_updated_date': candidate.data_updated_date.isoformat() if candidate.data_updated_date else None,
#                 'data_updated_time': candidate.data_updated_time.strftime('%H:%M:%S') if candidate.data_updated_time else None,
#                 'resume_present': candidate.resume_present
#             } for candidate in candidates],
#         }
    
#     # Convert response_data to JSON string
#     response_json = json.dumps(response_data, default=date_handler)

#     # Create the response
#     return response_json







# Mocked function for demonstration
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'pdf', 'doc', 'docx'}


import binascii   

@app.route('/add_candidate', methods=['POST'])
def add_candidate():
    try:
        
        # Retrieve request data from JSON
        data = request.json
        user_id = data['user_id']
        user = User.query.filter_by(id=user_id).first()
        user_type = user.user_type
        user_name = user.username
        

        job_id = data.get('job_id')
        client = data.get('client')
        name = data.get('name')
        mobile = data.get('mobile')
        email = data.get('email')
        profile = data.get('profile')
        skills = data.get('skills')
        current_company = data.get('current_company')
        position = data.get('position')
        current_job_location = data.get('current_job_location')
        preferred_job_location = data.get('preferred_job_location')
        
        # notice_period = data.get('notice_period')  # yes no completed
        # period_of_notice = data.get('period_of_notice')
        
        qualifications = data.get('qualifications')
        experience = data.get('experience')
        experience_months=data.get('experience')
        relevant_experience = data.get('relevant_experience')
        relevant_experience_months=data.get('relevant_experience_months')
        reason_for_job_change=data.get('reason_for_job_change')
        current_ctc = data.get('current_ctc')
        expected_ctc = data.get('expected_ctc')
        linkedin = data.get('linkedin')
        
        resume = data.get('resume')
        resume_binary = base64.b64decode(resume)
        print("Resume : ",type(resume_binary))

        # Set jd_pdf_present based on the presence of jd_pdf
        if resume_binary is not None:
            resume_present = True
        else:
            resume_present = False

        notice_period = data.get('serving_notice_period')
        last_working_date = None
        buyout = False
        period_of_notice = None
        if notice_period == 'yes':
            last_working_date=data.get('last_working_date')
            buyout=data.get('buyout')
        elif notice_period == 'no':
            period_of_notice = data.get('period_of_notice')
            buyout=data.get('buyout')
        # elif notice_period == 'completed':
        #     last_working_date=data.get('last_working_date')

        holding_offer = data.get('holding_offer')
        if holding_offer == 'yes':
            total_offers=data.get('total_offers')
            if total_offers == '':
                total_offers = 0
            else:
                total_offers=data.get('total_offers')
            highest_package_lpa=data.get('highest_package')
            if highest_package_lpa == '':
                highest_package_lpa = 0
            else:
                highest_package_lpa=data.get('highest_package')
        else:
            total_offers = None
            highest_package_lpa = None

        # # Check if the user is logged in
        if request.method == 'POST':
            # Retrieve the recruiter and management names based on user type
            if user_type == 'recruiter':
                # recruiter = User.query.get(user_id).name
                recruiter=user.name
                management = None
            elif user_type == 'management':
                recruiter = None
                # management = User.query.get(user_id).name
                management=user.name
            else:
                recruiter = None
                management = None

            # Check if the job_id is provided and job is active
            matching_job_post = JobPost.query.filter(and_(JobPost.id == job_id, JobPost.job_status == 'Active')).first()
            if not matching_job_post:
                return jsonify({'status': 'error',"message": "Job on hold"})

            # Create new candidate object
            new_candidate = Candidate(
                user_id=user_id,
                job_id=job_id,
                name=name,
                mobile=mobile,
                email=email,
                client=client,
                current_company=current_company,
                position=position,
                profile=profile,
                current_job_location=current_job_location,
                preferred_job_location=preferred_job_location,
                qualifications=qualifications,
                experience=experience,
                relevant_experience=relevant_experience,
                current_ctc=current_ctc,
                expected_ctc=expected_ctc,
                linkedin_url=linkedin,
                holding_offer=holding_offer,
                recruiter=recruiter,
                management=management,
                status='SCREENING',
                remarks=data.get('remarks'),
                skills=skills,
                resume=resume_binary,
                # serving_notice_period=serving_notice_period,
                notice_period=notice_period,
                period_of_notice=period_of_notice,
                # last_working_date=data.get('last_working_date') if notice_period in {'yes', 'completed'} else None,
                last_working_date=last_working_date,
                buyout=buyout,
                package_in_lpa=highest_package_lpa,
                total=total_offers,
                reason_for_job_change=reason_for_job_change,
                resume_present=resume_present
                # buyout='buyout' in data
            )

            print("Hello !!")
            
            # new_candidate.date_created = date.today()
            # new_candidate.time_created = datetime.now().time()
    
            # Created data and time
            current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
            new_candidate.date_created = current_datetime.date()
            new_candidate.time_created = current_datetime.time()


            db.session.add(new_candidate)
            db.session.commit()

            return jsonify({'status': 'success',"message": "Candidate Added Successfully", "candidate_id": new_candidate.id})

        return jsonify({"error_message": "Method not found"})

    except Exception as e:
        print(e)
        return jsonify({'status': 'error',"message": "Candidate unable to add"})
        



@app.route('/get_job_role', methods=['GET'])
def get_job_role():
    job_id = request.args.get('job_id')

    job_post = JobPost.query.filter_by(id=job_id).first()
    if job_post:
        return jsonify({"role": job_post.role})
    else:
        return jsonify({"role": ""})

@app.route('/delete_candidate/<int:candidate_id>', methods=["POST"])
def delete_candidate(candidate_id):
    data = request.json
    user_id = data['user_id']
    user = User.query.filter_by(id=user_id).first()
    user_type = user.user_type
    username = user.username

    if user_type == 'management':
        candidate = Candidate.query.filter_by(id=candidate_id).first()

        if candidate:
            if request.method == "POST":
                # Check for null values
                if None in [candidate.name, candidate.email, candidate.client, candidate.profile, candidate.status]:
                    # If any value is null, delete the candidate
                    Candidate.query.filter_by(id=candidate_id).delete()
                    db.session.commit()
                    return jsonify({'status': 'success', "message": "Candidate details deleted successfully"})

                # Save deletion details before deleting the candidate
                deleted_candidate = Deletedcandidate(
                    username=username,
                    candidate_name=candidate.name,
                    candidate_email=candidate.email,
                    client=candidate.client,
                    profile=candidate.profile,
                    status=candidate.status
                )
                db.session.add(deleted_candidate)
                db.session.commit()

                # Delete the candidate
                Candidate.query.filter_by(id=candidate_id).delete()
                db.session.commit()

                return jsonify({'status': 'success', "message": "Candidate details deleted successfully"})

            return jsonify({
                "candidate": {
                    "id": candidate.id,
                    "name": candidate.name,
                    "email": candidate.email,
                    "client": candidate.client,
                    "profile": candidate.profile,
                    "status": candidate.status
                },
                "user_name": username
            })

        else:
            return jsonify({'status': 'error', "message": "Candidate not found"})

    return jsonify({'status': 'error', "message": "Unauthorized: Only management can delete candidates"})


# @app.route('/delete_candidate/<int:candidate_id>', methods=["POST"])
# def delete_candidate(candidate_id):
#     data = request.json
#     user_id = data['user_id']
#     user = User.query.filter_by(id=user_id).first()
#     user_type = user.user_type
#     username=user.username
    
#     if user_type == 'management':
#         candidate = Candidate.query.filter_by(id=candidate_id).first()

#         if candidate:
#             if request.method == "POST":
#                 # Save deletion details before deleting the candidate
#                 deleted_candidate = Deletedcandidate(
#                     username=username,
#                     candidate_name=candidate.name,
#                     candidate_email=candidate.email,
#                     client=candidate.client,
#                     profile=candidate.profile,
#                     status=candidate.status
#                 )
#                 db.session.add(deleted_candidate)
#                 db.session.commit()

#                 # Delete the candidate
#                 Candidate.query.filter_by(id=candidate_id).delete()
#                 db.session.commit()

#                 return jsonify({'status': 'success',"message": "Candidate details deleted successfully"})

#             return jsonify({
#                 "candidate": {
#                     "id": candidate.id,
#                     "name": candidate.name,
#                     "email": candidate.email,
#                     "client": candidate.client,
#                     "profile": candidate.profile,
#                     "status": candidate.status
#                 },
#                 "user_name": username
#             })

#         else:
#             return jsonify({'status': 'error',"message": "Candidate not found"}), 404

#     return jsonify({'status': 'error',"message": "Unauthorized: Only management can delete candidates"}), 401





def verify_token(token):
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    try:
        user_id = serializer.loads(token, max_age=86400)  
        return user_id
    except BadSignature:
        return None  
    except Exception as e:
        return None

# Search String Changed
# @app.route('/update_candidate/<int:candidate_id>/<page_no>/<search_string>', methods=['GET', 'POST'])
from flask import session, jsonify

@app.route('/update_candidate/<int:candidate_id>', methods=['POST'])
def update_candidate(candidate_id):
    data = request.json

    user_id = data['user_id']
    user = User.query.filter_by(id=user_id).first()
    user_type = user.user_type
    user_name = user.username
    count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                      Notification.recruiter_name == user_name).count()
    career_count_notification_no = Career_notification.query.filter(
        Career_notification.notification_status == 'false',
        Career_notification.recruiter_name == user_name).count()
    
    if user_type == 'recruiter':
        recruiter = User.query.get(user_id).name
        management = None
    elif user_type == 'management':
        recruiter = None
        management = User.query.get(user_id).name
    else:
        recruiter = None
        management = None

    if user_type == 'recruiter':
        user_email = User.query.get(user_id).email
        management_email = None
    elif user_type == 'management':
        user_email = None
        management_email = User.query.get(user_id).email
    else:
        user_email = None
        management_email = None

    candidate = Candidate.query.filter_by(id=candidate_id).first()
    if not candidate:
        return jsonify({'status': 'error',"message": "Candidate not found"})
    
    previous_status = candidate.status

    candidate_status = request.json.get('candidate_status')
    candidate_comment = request.json.get('comments')

    candidate.status = candidate_status
    candidate.comments = candidate_comment

    # Update data_updated_date and data_updated_time
    current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
    candidate.data_updated_date = current_datetime.date()
    candidate.data_updated_time = current_datetime.time()

   # If candidate is on-boarded, decrement the job post's no_of_positions count
    if candidate_status == "ON-BOARDED":
        job_post = JobPost.query.filter_by(id=candidate.job_id).first()
        if job_post and job_post.no_of_positions is not None and int(job_post.no_of_positions) > 0:
            job_post.no_of_positions = str(int(job_post.no_of_positions) - 1)
            db.session.commit()
    db.session.commit()

    # db.session.commit

    if candidate_status in [
            "SCREENING", "SCREEN REJECTED", "NO SHOW", "DROP", "CANDIDATE HOLD", "OFFERED - DECLINED", "DUPLICATE", "SCREENING SELECTED",
            "L1-SCHEDULE", "L1-FEEDBACK", "L1-SELECTED", "L1-REJECTED", "CANDIDATE RESCHEDULE", "PANEL RESCHEDULE", "L2-SCHEDULE", 
            "L2-FEEDBACK", "L2-SELECTED", "L2-REJECTED", "HR-ROUND", "MANAGERIAL ROUND", "NEGOTIATION", "SELECTED", "OFFER-REJECTED",
            "OFFER-DECLINED", "ON-BOARDED", "HOLD", "CANDIDATE NO-SHOW"
            ]:
        candidate_name = candidate.name
        candidate_position = candidate.position
        candidate_email = candidate.email

        if candidate_position:
            candidate_position = candidate_position.upper()
        else:
            candidate_position = ""

        if candidate.client:
            client = candidate.client.upper()
        else:
            client = ""

        if candidate_status in ["SCREENING", "SCREEN REJECTED"]:
            message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
        else:
            message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Previous Status : "{previous_status}"\n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
    else:
        message = ""
        candidate_name = ""
        candidate_position = ""
        candidate_email = ""

    return jsonify({
        'status': 'success',
        "message": "Candidate Status Updated Successfully",
        "user_id": user_id,
        "user_type": user_type,
        "user_name": user_name,
        "count_notification_no": count_notification_no,
        "career_count_notification_no": career_count_notification_no,
        "recruiter": recruiter,
        "management": management,
        "recruiter_email": user_email,
        "management_email": management_email,
        "candidate_name": candidate_name,
        "candidate_position": candidate_position,
        "candidate_email": candidate_email,
        "message_body": message 
    })


# @app.route('/update_candidate/<int:candidate_id>', methods=['POST'])
# def update_candidate(candidate_id):
#     data = request.json

#     user_id = data['user_id']
#     user = User.query.filter_by(id=user_id).first()
#     user_type = user.user_type
#     user_name = user.username
#     count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
#                                                       Notification.recruiter_name == user_name).count()
#     career_count_notification_no = Career_notification.query.filter(
#         Career_notification.notification_status == 'false',
#         Career_notification.recruiter_name == user_name).count()
    
#     if user_type == 'recruiter':
#         recruiter = User.query.get(user_id).name
#         management = None
#     elif user_type == 'management':
#         recruiter = None
#         management = User.query.get(user_id).name
#     else:
#         recruiter = None
#         management = None

#     if user_type == 'recruiter':
#         user_email = User.query.get(user_id).email
#         management_email = None
#     elif user_type == 'management':
#         user_email = None
#         management_email = User.query.get(user_id).email
#     else:
#         user_email = None
#         management_email = None

#     candidate = Candidate.query.filter_by(id=candidate_id).first()
#     if not candidate:
#         return jsonify({"error_message": "Candidate not found"}), 500
    
#     previous_status = candidate.status

#     candidate_status = request.json.get('candidate_status')
#     candidate_comment = request.json.get('comments')

#     candidate.status = candidate_status
#     candidate.comments = candidate_comment

#     # Update data_updated_date and data_updated_time
#     current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#     candidate.data_updated_date = current_datetime.date()
#     candidate.data_updated_time = current_datetime.time()

#     db.session.commit()

#     if candidate_status in [
#             "SCREENING", "SCREEN REJECTED", "NO SHOW", "DROP", "CANDIDATE HOLD", "OFFERED - DECLINED", "DUPLICATE", "SCREENING SELECTED",
#             "L1-SCHEDULE", "L1-FEEDBACK", "L1-SELECTED", "L1-REJECTED", "CANDIDATE RESCHEDULE", "PANEL RESCHEDULE", "L2-SCHEDULE", 
#             "L2-FEEDBACK", "L2-SELECTED", "L2-REJECTED", "HR-ROUND", "MANAGERIAL ROUND", "NEGOTIATION", "SELECTED", "OFFER-REJECTED",
#             "OFFER-DECLINED", "ON-BOARDED", "HOLD", "CANDIDATE NO-SHOW"
#             ]:
#         candidate_name = candidate.name
#         candidate_position = candidate.position
#         candidate_email = candidate.email

#         if candidate_position:
#             candidate_position = candidate_position.upper()
#         else:
#             candidate_position = ""

#         if candidate.client:
#             client = candidate.client.upper()
#         else:
#             client = ""

#         if candidate_status in ["SCREENING", "SCREEN REJECTED"]:
#             message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
#         else:
#             message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Previous Status : "{previous_status}"\n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
#     else:
#         message = ""
#         candidate_name = ""
#         candidate_position = ""
#         candidate_email = ""

#     return jsonify({
#         "message": "Candidate Status Updated Successfully",
#         "user_id": user_id,
#         "user_type": user_type,
#         "user_name": user_name,
#         "count_notification_no": count_notification_no,
#         "career_count_notification_no": career_count_notification_no,
#         "recruiter": recruiter,
#         "management": management,
#         "recruiter_email": user_email,
#         "management_email": management_email,
#         "candidate_name": candidate_name,
#         "candidate_position": candidate_position,
#         "candidate_email": candidate_email,
#         "message_body": message 
#     })

# @app.route('/update_candidate/<int:candidate_id>', methods=['POST'])
# def update_candidate(candidate_id):
#     data = request.json

#     user_id = data['user_id']
#     user = User.query.filter_by(id=user_id).first()
#     user_type = user.user_type
#     user_name = user.username
#     count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
#                                                       Notification.recruiter_name == user_name).count()
#     career_count_notification_no = Career_notification.query.filter(
#         Career_notification.notification_status == 'false',
#         Career_notification.recruiter_name == user_name).count()
#     if request.method == 'POST':
#         if user_type == 'recruiter':
#             recruiter = User.query.get(user_id).name
#             management = None
#         elif user_type == 'management':
#             recruiter = None
#             management = User.query.get(user_id).name
#         else:
#             recruiter = None
#             management = None

#         if user_type == 'recruiter':
#             user_email = User.query.get(user_id).email
#             management_email = None
#         elif user_type == 'management':
#             user_email = None
#             management_email = User.query.get(user_id).email
#         else:
#             user_email = None
#             management_email = None

#         candidate = Candidate.query.filter_by(id=candidate_id).first()
#         print(candidate)
        
#         previous_status = candidate.status

#         candidate_status = request.json.get('candidate_status')
#         candidate_comment = request.json.get('comments')

#         candidate.status = candidate_status
#         candidate.comments = candidate_comment

#         db.session.commit()

#         if candidate_status in [
#                 "SCREENING", "SCREEN REJECTED", "NO SHOW", "DROP", "CANDIDATE HOLD", "OFFERED - DECLINED", "DUPLICATE", "SCREENING SELECTED",
#                 "L1-SCHEDULE", "L1-FEEDBACK", "L1-SELECTED", "L1-REJECTED", "CANDIDATE RESCHEDULE", "PANEL RESCHEDULE", "L2-SCHEDULE", 
#                 "L2-FEEDBACK", "L2-SELECTED", "L2-REJECTED", "HR-ROUND", "MANAGERIAL ROUND", "NEGOTIATION", "SELECTED", "OFFER-REJECTED",
#                 "OFFER-DECLINED", "ON-BOARDED", "HOLD", "CANDIDATE NO-SHOW"
#                 ]:
#             candidate_name = candidate.name
#             candidate_position = candidate.position
#             candidate_email = candidate.email

#             if candidate_position:
#                 candidate_position = candidate_position.upper()
#             else:
#                 candidate_position = ""

#             if candidate.client:
#                 client = candidate.client.upper()
#             else:
#                 client = ""

#             if candidate_status in ["SCREENING", "SCREEN REJECTED"]:
#                 message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
#             else:
#                 message = f'Dear {candidate_name}, \n\nGreetings! \n\nWe hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate_position} position and participating in the recruitment process. \n\nWe are writing to inform you about the latest update we received from our client {client} regarding your interview. \n\n        Previous Status : "{previous_status}"\n\n        Current Status :  "{candidate_status}"\n\nThank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. \n\nIf you have any questions or need further information, please feel free to reach out to us. \n\nThanks,\n'
#         else:
#             message = ""
#             candidate_name = ""
#             candidate_position = ""
#             candidate_email = ""

#         return jsonify({
#         "message": "Candidate Status Updated Successfully",
#         "user_id": user_id,
#         "user_type": user_type,
#         "user_name": user_name,
#         "count_notification_no": count_notification_no,
#         "career_count_notification_no": career_count_notification_no,
#         "recruiter": recruiter,
#         "management": management,
#         "recruiter_email": user_email,
#         "management_email": management_email,
#         "candidate_name": candidate_name,
#         "candidate_position": candidate_position,
#         "candidate_email": candidate_email,
#         # "message": message
#         "message_body": message 
#     })

@app.route('/update_candidate_careers/<int:candidate_id>/<page_no>/<search_string>', methods=['GET', 'POST'])
@app.route('/update_candidate_careers/<int:candidate_id>/<page_no>', methods=['GET', 'POST'])
def update_candidate_careers(candidate_id, page_no):
    if 'user_id' in session and 'user_type' in session:
        user_id = session['user_id']
        user_type = session['user_type']
        user_name = session['user_name']
        count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                          Notification.recruiter_name == user_name).count()
        career_count_notification_no = Career_notification.query.filter(
            Career_notification.notification_status == 'false',
            Career_notification.recruiter_name == user_name).count()
        if request.method == 'POST':
            # Retrieve the logged-in user's ID and user type from the session
            user_id = session['user_id']
            user_type = session['user_type']

            # Retrieve the recruiter and management names based on user type
            if user_type == 'recruiter':
                recruiter = User.query.get(user_id).name
                management = None
            elif user_type == 'management':
                recruiter = None
                management = User.query.get(user_id).name
            else:
                recruiter = None
                management = None

            if user_type == 'recruiter':
                recruiter_email = User.query.get(user_id).email
                management = None
            elif user_type == 'management':
                recruiter = None
                recruiter_email = User.query.get(user_id).email
            else:
                recruiter_email = None
                management_email = None

            # Retrieve the form data for the candidate
            candidate = Candidate.query.get(candidate_id)
            previous_status = candidate.status
            # candidate.recruiter = recruiter
            # candidate.management = management

            # Get the selected candidate status from the form
            candidate_status = request.form.get('candidate_status')
            candidate_comment = request.form.get('comments')

            # Update the candidate status field
            candidate.status = candidate_status

            candidate.comments = candidate_comment

            db.session.commit()

            if candidate_status == "SCREENING" or candidate_status == "SCREEN REJECTED":
                candidate_name = candidate.name
                candidate_position = candidate.position

                # Retrieve the candidate's email
                candidate_email = candidate.email

                # Determine if the logged-in user is a recruiter or management
                user_type = session.get('user_type')

                if user_type == 'recruiter' or user_type == 'management':
                    # Retrieve the corresponding user's email
                    user_email = User.query.get(session.get('user_id')).email

                    message = Message(f'Job Application Status - {candidate_position}',
                                      sender=config.sender_email, recipients=[candidate_email])

                    if user_type == 'management':
                        management_email = user_email
                        message.cc = [management_email]
                    elif user_type == 'recruiter':
                        recruiter_email = user_email
                        message.cc = [recruiter_email]
                    message.body = f'''Dear {candidate.name}, 

Greetings! 

We hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate.position.upper()} position and participating in the recruitment process. 

We are writing to inform you about the latest update we received from our client {candidate.client.upper()} regarding your interview. 

        Current Status :  "{candidate.status}"

Thank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. 

If you have any questions or need further information, please feel free to reach out to us. 

Thanks, 
                            '''
                #mail.send(message)
                pass
            elif candidate_status == "NO SHOW" or candidate_status == "DROP" or candidate_status == "CANDIDATE HOLD" or candidate_status == "OFFERED - DECLINED" or candidate_status == "DUPLICATE":
                pass
            else:
                candidate_name = candidate.name
                candidate_position = candidate.position
                candidate_email = candidate.email

                user_type = session.get('user_type')

                if user_type == 'recruiter' or user_type == 'management':
                    user_email = User.query.get(session.get('user_id')).email

                    message = Message(f'Job Application Status - {candidate_position}',
                                      sender=config.sender_email, recipients=[candidate_email])

                    if user_type == 'management':
                        management_email = user_email
                        message.cc = [management_email]
                    elif user_type == 'recruiter':
                        recruiter_email = user_email
                        message.cc = [recruiter_email]
                    message.body = f'''Dear {candidate.name}, 

Greetings! 

We hope this email finds you well. We wanted to extend our thanks for showing your interest in the {candidate.position.upper()} position and participating in the recruitment process. 

We are writing to inform you about the latest update we received from our client {candidate.client.upper()} regarding your interview. 

        Previous Status : "{previous_status}"

        Current Status :  "{candidate.status}"

Thank you once again for considering this opportunity with us. We wish you all the best in your future endeavors. 

If you have any questions or need further information, please feel free to reach out to us. 

Thanks, 
                            '''
                #mail.send(message)
                pass

            return redirect(
                url_for('career_dashboard', update_candidate_message='Candidate Status Updated Sucessfully', page_no=page_no))

        candidate = Candidate.query.get(candidate_id)
        candidate_data = {
            'id': candidate.id,
            'name': candidate.name,
            'mobile': candidate.mobile,
            'email': candidate.email,
            'client': candidate.client,
            'current_company': candidate.current_company,
            'position': candidate.position,
            'profile': candidate.profile,
            'current_job_location': candidate.current_job_location,
            'preferred_job_location': candidate.preferred_job_location,
            'resume': candidate.resume,
            'qualifications': candidate.qualifications,
            'experience': candidate.experience,
            'relevant_experience': candidate.relevant_experience,
            'current_ctc': candidate.current_ctc,
            'expected_ctc': candidate.expected_ctc,
            'notice_period': candidate.notice_period,
            'last_working_date': candidate.last_working_date,
            'buyout': candidate.buyout,
            'holding_offer': candidate.holding_offer,
            'total': candidate.total,
            'package_in_lpa': candidate.package_in_lpa,
            'reason_for_job_change': candidate.reason_for_job_change,
            'remarks': candidate.remarks,
            'candidate_status': candidate.status,
        }

        return render_template('update_candidate.html', candidate_data=candidate_data, user_id=user_id,
                               user_type=user_type, user_name=user_name, candidate=candidate,
                               count_notification_no=count_notification_no,
                               career_count_notification_no=career_count_notification_no)

    return redirect(url_for('career_dashboard'))


from flask import jsonify

@app.route('/logout', methods=['POST'])
def logout():
    data = request.json
    
    if data:
        user_id = data.get('user_id')
        
        if user_id:
            user = User.query.filter_by(id=user_id).first()
            
            if user:
                user_type = user.user_type
                user_name = user.username
                
                return jsonify({"message": "Logged out successfully"}), 200
            
            return jsonify({"message": "User not found"}), 404
        else:
            return jsonify({"message": "'user_id' not provided in JSON data"}), 400
    
    return jsonify({"message": "No JSON data provided"}), 400



from datetime import datetime

@app.route('/edit_candidate/<int:candidate_id>', methods=['POST'])
def edit_candidate(candidate_id):
    try:
        data = request.json

        user_id = data.get('user_id')
        if not user_id:
            return jsonify({"error_message": "User ID is required"}), 400

        user = User.query.filter_by(id=user_id).first()
        if not user:
            return jsonify({"error_message": "User not found"}), 404

        user_name = user.username
        count_notification_no = Notification.query.filter(
            Notification.notification_status == 'false',
            Notification.recruiter_name == user_name
        ).count()
        career_count_notification_no = Career_notification.query.filter(
            Career_notification.notification_status == 'false',
            Career_notification.recruiter_name == user_name
        ).count()

        # Retrieve the candidate object
        candidate = Candidate.query.get(candidate_id)
        if not candidate:
            return jsonify({"error_message": "Candidate not found"}), 404

        # Update the candidate fields with the new data
        candidate.name = data.get('name')
        candidate.mobile = data.get('mobile')
        candidate.email = data.get('email')
        candidate.client = data.get('client')
        candidate.current_company = data.get('current_company')
        candidate.position = data.get('position')
        candidate.profile = data.get('profile')
        candidate.current_job_location = data.get('current_job_location')
        candidate.preferred_job_location = data.get('preferred_job_location')
        candidate.qualifications = data.get('qualifications')
        candidate.experience = data.get('experience')
        candidate.relevant_experience = data.get('relevant_experience')
        candidate.current_ctc = data.get('current_ctc')
        candidate.expected_ctc = data.get('expected_ctc')    
        candidate.reason_for_job_change = data.get('reason_for_job_change')
        candidate.linkedin_url = data.get('linkedin')
        candidate.remarks = data.get('remarks')
        candidate.skills = data.get('skills')
        # candidate.holding_offer = data.get('holding_offer')
        candidate.total = data.get('total_offers')
        candidate.package_in_lpa = data.get('highest_package')
        
        
        # Handle resume decoding
        resume_data = data.get('resume')
        if resume_data is not None:
            try:
                resume_binary = base64.b64decode(resume_data)
                candidate.resume = resume_binary
                candidate.resume_present = True  # Update resume_present to True if resume is provided and valid
            except (binascii.Error, TypeError) as e:
                return jsonify({"error_message": "Invalid resume format"}), 500

        # Serving notice period logic
        notice_period = data.get('serving_notice_period')
        candidate.notice_period = notice_period
        if notice_period == 'yes':
            candidate.last_working_date = data.get('last_working_date')
            candidate.buyout = data.get('buyout')
        elif notice_period == 'no':
            candidate.period_of_notice = data.get('period_of_notice')
            candidate.buyout = data.get('buyout')
        # elif notice_period == 'completed':
        #     candidate.last_working_date = data.get('last_working_date')

        # Holding offer logic
        holding_offer = data.get('holding_offer')
        candidate.holding_offer = holding_offer
        if holding_offer == 'yes':
            total_offers = data.get('total_offers')
            candidate.total = 0 if total_offers == '' else total_offers
            highest_package = data.get('highest_package')
            candidate.package_in_lpa = 0 if highest_package == '' else highest_package
        else:
            candidate.total = None
            candidate.package_in_lpa = None

        # Update data_updated_date and data_updated_time
        # current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
        # candidate.data_updated_date = current_datetime.date()
        # candidate.data_updated_time = current_datetime.time()

        db.session.commit()
        return jsonify({'status': 'success', "message": "Candidate Details Edited Successfully"})

    except Exception as e:
        print(e)
        return jsonify({'status': 'error', "message": "Candidate Details not Edited "})





@app.route('/edit_candidate_careers/<int:candidate_id>/<int:page_no>/<search_string>', methods=['GET', 'POST'])
@app.route('/edit_candidate_careers/<int:candidate_id>/<int:page_no>', methods=['GET', 'POST'])
def edit_candidate_careers(candidate_id, page_no):
    if 'user_id' in session and 'user_type' in session:
        user_id = session['user_id']
        user_type = session['user_type']
        user_name = session['user_name']
        count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                          Notification.recruiter_name == user_name).count()
        career_count_notification_no = Career_notification.query.filter(
            Career_notification.notification_status == 'false',
            Career_notification.recruiter_name == user_name).count()

        if request.method == 'POST':
            # Retrieve the logged-in user's ID and user type from the session
            user_id = session['user_id']
            user_type = session['user_type']

            # Retrieve the form data for the candidate
            candidate = Candidate.query.get(candidate_id)

            # Update the candidate information based on user type
            if user_type == 'recruiter':
                candidate.recruiter = User.query.get(user_id).name
            elif user_type == 'management':
                candidate.management = User.query.get(user_id).name

            # Update the candidate fields with the new form data
            candidate.name = request.form.get('name')
            candidate.mobile = request.form.get('mobile')
            candidate.email = request.form.get('email')
            candidate.client = request.form.get('client')
            candidate.current_company = request.form.get('current_company')
            candidate.position = request.form.get('position')
            candidate.profile = request.form.get('profile')
            candidate.current_job_location = request.form.get('current_job_location')
            candidate.preferred_job_location = request.form.get('preferred_job_location')
            candidate.qualifications = request.form.get('qualifications')
            experience = request.form.get('experience')
            exp_months = request.form.get('exp_months')
            candidate.experience = experience +'.'+exp_months
            relevant_experience = request.form.get('relevant_experience')
            relevant_exp_months = request.form.get('relevant_exp_months')
            candidate.relevant_experience = relevant_experience + '.' + relevant_exp_months
            candidate.current_ctc = request.form.get('current_ctc')
            candidate.expected_ctc = request.form.get('expected_ctc')
            currency_type_current = request.form['currency_type_current']
            currency_type_except = request.form['currency_type_except']
            candidate.current_ctc = currency_type_current + " " + request.form['current_ctc']
            candidate.expected_ctc = currency_type_except + " " + request.form['expected_ctc']
            candidate.notice_period = request.form.get('notice_period')
            candidate.reason_for_job_change = request.form.get('reason_for_job_change')
            candidate.linkedin_url = request.form.get('linkedin')
            candidate.remarks = request.form.get('remarks')
            candidate.skills = request.form.get('skills')
            candidate.holding_offer = request.form.get('holding_offer')
            candidate.total = request.form.get('total')
            candidate.package_in_lpa = request.form.get('package_in_lpa')
            candidate.period_of_notice = request.form.get('period_of_notice')

            # Handle the resume file upload
            resume_file = request.files['resume']
            if resume_file.filename != '':
                # Save the new resume to the candidate's resume field as bytes
                candidate.resume = resume_file.read()

            holding_offer = request.form.get('holding_offer')
            if holding_offer == 'yes':
                total = request.form.get('total')
                package_in_lpa = request.form.get('package_in_lpa')

                candidate.total = total
                candidate.package_in_lpa = package_in_lpa
            elif holding_offer in ['no', 'pipeline']:
                candidate.total = None
                candidate.package_in_lpa = None

            notice_period = request.form.get('notice_period')
            if notice_period == 'yes':
                last_working_date = request.form['last_working_date']
                buyout = 'buyout' in request.form
                candidate.last_working_date = last_working_date
                candidate.buyout = buyout
            elif notice_period == 'no':
                period_of_notice = request.form['months']
                buyout = 'buyout' in request.form
                candidate.period_of_notice = period_of_notice
                candidate.buyout = buyout
            elif notice_period == 'completed':
                last_working_date = request.form['last_working_date']
                candidate.last_working_date = last_working_date

            db.session.commit()

            return redirect(
                url_for('career_dashboard', page_no=page_no, edit_candidate_message='Candidate Details Edited Successfully'))

        candidate = Candidate.query.get(candidate_id)
        candidate_data = {
            'id': candidate.id,
            'name': candidate.name,
            'mobile': candidate.mobile,
            'email': candidate.email,
            'client': candidate.client,
            'current_company': candidate.current_company,
            'position': candidate.position,
            'profile': candidate.profile,
            'current_job_location': candidate.current_job_location,
            'preferred_job_location': candidate.preferred_job_location,
            'qualifications': candidate.qualifications,
            'experience': candidate.experience,
            'relevant_experience': candidate.relevant_experience,
            'current_ctc': candidate.current_ctc,
            'expected_ctc': candidate.expected_ctc,
            'notice_period': candidate.notice_period,
            'reason_for_job_change': candidate.reason_for_job_change,
            'remarks': candidate.remarks,
            'candidate_status': candidate.status,
            'linkedin_url': candidate.linkedin_url,
            'skills': candidate.skills,
            'resume': candidate.resume,
            'holding_offer': candidate.holding_offer,
            'total': candidate.total,
            'package_in_lpa': candidate.package_in_lpa,
            'last_working_date': candidate.last_working_date,
            'buyout': candidate.buyout,
            'period_of_notice': candidate.period_of_notice,
        }

        return render_template('edit_candidate_careers.html', candidate_data=candidate_data, user_id=user_id,
                               user_type=user_type, user_name=user_name, count_notification_no=count_notification_no,
                               page_no=page_no,career_count_notification_no=career_count_notification_no)

    return redirect(url_for('career_dashboard'))


# Function to check if a filename has an allowed extension
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


from flask import send_file

import io
import os
import base64
from flask import redirect, url_for, send_file

@app.route('/download_resume/<int:candidate_id>', methods=['GET','POST'])
def download_resume(candidate_id):
    candidate = Candidate.query.get(candidate_id)

    if candidate is None or candidate.resume is None:
        return redirect(url_for('dashboard'))

    # Decode the base64 encoded resume
    resume_data = candidate.resume.split(',')[1]  # Get the data part after the comma
    resume_bytes = base64.b64decode(resume_data)

    # Determine the file extension
    if candidate.resume.startswith("data:application/pdf"):
        resume_filename = f"{candidate.name}_resume.pdf"
    else:
        resume_filename = f"{candidate.name}_resume.docx"

    # Send the resume data for download
    return send_file(io.BytesIO(resume_bytes),
                     attachment_filename=resume_filename,
                     as_attachment=True)



# def send_notification(recruiter_email):
#     msg = Message('New Job Posted', sender=config.sender_email, recipients=[recruiter_email])
#     msg.body = 'A new job has been posted. Check your dashboard for more details.'
#     mail.send(msg)

def post_job_send_notification(recruiter_email, new_recruiter_name, job_data):
    html_body = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                color: #333;
                line-height: 1.6;
                background-color: #f4f4f4;
                margin: 0;
                padding: 0;
            }}
            .container {{
                padding: 20px;
                margin: 20px auto;
                max-width: 600px;
                background-color: #ffffff;
                border: 1px solid #ddd;
                border-radius: 8px;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                text-align: center;
                font-size: 20px;
                border-radius: 8px 8px 0 0;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin-top: 10px;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #4CAF50;
                color: white;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            p {{
                margin: 10px 0;
            }}
            .footer {{
                margin-top: 20px;
                font-size: 12px;
                color: #777;
                text-align: center;
                border-top: 1px solid #ddd;
                padding-top: 10px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                New Job Posted
            </div>
            <p>Dear {new_recruiter_name},</p>
            <p>A new requirement has been assigned to you.</p>
            <p> Please find the details below:</p>
            <table>
                <tr>
                    <th style="width: 20%;">Job ID</th>
                    <th style="width: 30%;">Client</th>
                    <th style="width: 30%;">Role/Profile</th>
                    <th style="width: 30%;">Location</th>
                </tr>
                {job_data}
            </table>
            <p>Please check in Job Listing page for more details.</p>
            <p>Regards,</p>
            <p><b>Makonis Talent Track Pro Team</b></p>
        </div>
    </body>
    </html>
    """

    msg = Message(
        'New Requirement Assigned',
        sender=config.sender_email,
        recipients=[recruiter_email]
    )
    msg.html = html_body
    try:
        mail.send(msg)
    except Exception as e:
        print("mail error", str(e))
        return f'Failed to send mail: {str(e)}'
    return None



import re

def is_valid_email(email):
    """ Validate email format using regex. """
    regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(regex, email) is not None

@app.route('/post_job', methods=['POST'])
def post_job():
    data = request.json

    user_id = data.get('user_id')
    if not user_id:
        return jsonify({'status': 'error', 'message': 'user_id is required'}), 400

    user = User.query.filter_by(id=user_id).first()
    if not user:
        return jsonify({'status': 'error', 'message': 'User not found'}), 404

    if user.user_type != 'management':
        return jsonify({'status': 'error', 'message': 'You do not have permission to post a job'}), 401

    try:
        job_details = {
            'client': data['client'],
            'experience_min': data['experience_min'],
            'experience_max': data['experience_max'],
            'budget_min': f"{data['currency_type_min']} {data['budget_min']}",
            'budget_max': f"{data['currency_type_max']} {data['budget_max']}",
            'location': data['location'],
            'shift_timings': data['shift_timings'],
            'notice_period': data['notice_period'],
            'role': data['role'],
            'detailed_jd': data['detailed_jd'],
            'mode': data['mode'],
            'job_status': data['job_status'],
            'skills': data['skills'],
            'job_type': data['Job_Type'],
            'no_of_positions': data['no_of_positions'],
            'contract_in_months': data['Job_Type_details'] if data['Job_Type'] == 'Contract' else None
        }
    except KeyError as e:
        return jsonify({'status': 'error', 'message': f'Missing required field: {e}'}), 400

    jd_pdf = data.get('jd_pdf')
    jd_binary = None
    jd_pdf_present = False

    if jd_pdf:
        try:
            jd_binary = base64.b64decode(jd_pdf)
            jd_pdf_present = bool(jd_binary)
        except Exception:
            return jsonify({'status': 'error', 'message': 'Error decoding base64 PDF file'}), 400

    new_job_post = JobPost(
        client=job_details['client'],
        experience_min=job_details['experience_min'],
        experience_max=job_details['experience_max'],
        budget_min=job_details['budget_min'],
        budget_max=job_details['budget_max'],
        location=job_details['location'],
        shift_timings=job_details['shift_timings'],
        notice_period=job_details['notice_period'],
        role=job_details['role'],
        detailed_jd=job_details['detailed_jd'],
        recruiter=', '.join(data.get('recruiter', [])),
        management=user.username,
        mode=job_details['mode'],
        job_status=job_details['job_status'],
        job_type=job_details['job_type'],
        skills=job_details['skills'],
        contract_in_months=job_details['contract_in_months'],
        jd_pdf=jd_binary,
        jd_pdf_present=jd_pdf_present,
        no_of_positions=job_details['no_of_positions']
    )

    try:
        current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
        new_job_post.date_created = current_datetime.date()
        new_job_post.time_created = current_datetime.time()

        db.session.add(new_job_post)
        db.session.commit()

        job_post_id = new_job_post.id
        for recruiter_name in data.get('recruiter', []):
            notification = Notification(
                job_post_id=job_post_id,
                recruiter_name=recruiter_name.strip(),
                notification_status=False,
                num_notification=1
            )
            db.session.add(notification)

        db.session.commit()

        return jsonify({'status': 'success', 'message': 'Job posted successfully', 'job_post_id': job_post_id}), 200
    except Exception as e:
        db.session.rollback()
        print(e)
        return jsonify({'status': 'error', 'message': f'Failed to post job: {str(e)}'}), 500

@app.route('/send_notifications', methods=['POST'])
def send_notifications():
    data = request.json
    job_post_id = data.get('job_post_id')

    if not job_post_id:
        return jsonify({'status': 'error', 'message': 'job_post_id is required'}), 400

    job_post = JobPost.query.filter_by(id=job_post_id).first()
    if not job_post:
        return jsonify({'status': 'error', 'message': 'Job post not found'}), 404

    job_data = f"<tr><td>{job_post.id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"

    invalid_emails = []
    for recruiter_name in job_post.recruiter.split(','):
        recruiter = User.query.filter_by(name=recruiter_name.strip()).first()
        if recruiter:
            if is_valid_email(recruiter.email):
                error_msg = post_job_send_notification(recruiter.email, recruiter.name, job_data)
                if error_msg:
                    return jsonify({'status': 'error', 'message': error_msg}), 500
            else:
                invalid_emails.append(recruiter.email)

    if invalid_emails:
        print("Invalid Emails:", invalid_emails)  # Print invalid_emails list
        return jsonify({'status': 'error', 'message': f'Invalid email format for: {", ".join(invalid_emails)}'}), 400

    return jsonify({'status': 'success', 'message': 'Notifications sent successfully'}), 200



# def is_valid_email(email):
#     """ Validate email format using regex. """
#     regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
#     return re.match(regex, email) is not None
    
# @app.route('/post_job', methods=['POST'])
# def post_job():
#     data = request.json

#     user_id = data.get('user_id')
#     if not user_id:
#         return jsonify({'status': 'error', 'message': 'user_id is required'}), 400

#     user = User.query.filter_by(id=user_id).first()
#     if not user:
#         return jsonify({'status': 'error', 'message': 'User not found'}), 404

#     if user.user_type != 'management':
#         return jsonify({'status': 'error', 'message': 'You do not have permission to post a job'}), 401

#     try:
#         job_details = {
#             'client': data['client'],
#             'experience_min': data['experience_min'],
#             'experience_max': data['experience_max'],
#             'budget_min': f"{data['currency_type_min']} {data['budget_min']}",
#             'budget_max': f"{data['currency_type_max']} {data['budget_max']}",
#             'location': data['location'],
#             'shift_timings': data['shift_timings'],
#             'notice_period': data['notice_period'],
#             'role': data['role'],
#             'detailed_jd': data['detailed_jd'],
#             'mode': data['mode'],
#             'job_status': data['job_status'],
#             'skills': data['skills'],
#             'job_type': data['Job_Type'],
#             'contract_in_months': data['Job_Type_details'] if data['Job_Type'] == 'Contract' else None
#         }
#     except KeyError as e:
#         return jsonify({'status': 'error', 'message': f'Missing required field: {e}'}), 400

#     jd_pdf = data.get('jd_pdf')
#     jd_binary = None
#     jd_pdf_present = False

#     if jd_pdf:
#         try:
#             jd_binary = base64.b64decode(jd_pdf)
#             jd_pdf_present = bool(jd_binary)
#         except Exception:
#             return jsonify({'status': 'error', 'message': 'Error decoding base64 PDF file'}), 400

#     new_job_post = JobPost(
#         client=job_details['client'],
#         experience_min=job_details['experience_min'],
#         experience_max=job_details['experience_max'],
#         budget_min=job_details['budget_min'],
#         budget_max=job_details['budget_max'],
#         location=job_details['location'],
#         shift_timings=job_details['shift_timings'],
#         notice_period=job_details['notice_period'],
#         role=job_details['role'],
#         detailed_jd=job_details['detailed_jd'],
#         recruiter=', '.join(data.get('recruiter', [])),
#         management=user.username,
#         mode=job_details['mode'],
#         job_status=job_details['job_status'],
#         job_type=job_details['job_type'],
#         skills=job_details['skills'],
#         contract_in_months=job_details['contract_in_months'],
#         jd_pdf=jd_binary,
#         jd_pdf_present=jd_pdf_present
#     )

#     try:
#         current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
#         new_job_post.date_created = current_datetime.date()
#         new_job_post.time_created = current_datetime.time()

#         db.session.add(new_job_post)
#         db.session.commit()

#         job_post_id = new_job_post.id
#         job_data = f"<tr><td>{job_post_id}</td><td>{new_job_post.client}</td><td>{new_job_post.role}</td><td>{new_job_post.location}</td></tr>"
        
#         invalid_emails = []
#         for recruiter_name in data.get('recruiter', []):
#             recruiter = User.query.filter_by(username=recruiter_name.strip()).first()
#             if recruiter:
#                 if is_valid_email(recruiter.email):
#                     error_msg = post_job_send_notification(recruiter.email, recruiter.username, job_data)
#                     # pass
#                     if error_msg:
#                         return jsonify({'status': 'error', 'message': error_msg}), 500
#                         # pass
#                 else:
#                     invalid_emails.append(recruiter.email)

#         if invalid_emails:
#             print("Invalid Emails:", invalid_emails)  # Print invalid_emails list
#             return jsonify({'status': 'error', 'message': f'Invalid email format for: {", ".join(invalid_emails)}'}), 400

#         return jsonify({'status': 'success', 'message': 'Job posted successfully', 'job_post_id': job_post_id}), 200
#     except Exception as e:
#         db.session.rollback()
#         print(e)
#         return jsonify({'status': 'error', 'message': f'Failed to post job: {str(e)}'}), 500




@app.route('/recruiter_job_posts', methods=['POST'])
def recruiter_job_posts():
    data = request.json
    user_id = data.get('user_id')  # Using get() to avoid KeyError if 'user_id' is missing
    if not user_id:
        return jsonify({"error": "User ID is missing"}), 400

    # Validate user existence
    recruiter = User.query.get(user_id)
    if not recruiter:
        return jsonify({"error": "Recruiter not found"}), 404

    recruiter_name = recruiter.name

    # Filter unread notifications based on recruiter name
    unread_notifications = Career_notification.query.filter(
        Career_notification.recruiter_name == recruiter_name,
        Career_notification.notification_status == False
    ).all()

    # Filter active and on-hold job posts
    active_job_posts = JobPost.query.filter(
        JobPost.recruiter == recruiter_name,  # Filtering based on the recruiter's name
        JobPost.job_status == 'Active'
    ).order_by(JobPost.id).all()

    on_hold_job_posts = JobPost.query.filter(
        JobPost.recruiter == recruiter_name,  # Filtering based on the recruiter's name
        JobPost.job_status == 'Hold'
    ).order_by(JobPost.id).all()

    # Update notification statuses after retrieving them
    for notification in unread_notifications:
        notification.notification_status = True
    db.session.commit()

    # Construct JSON response
    response_data = {
        "count_notification_no": len(unread_notifications),
        "job_posts": [job_post_to_dict(job_post) for job_post in active_job_posts],
        "user_name": recruiter_name,
        "job_posts_hold": [job_post_to_dict(job_post) for job_post in on_hold_job_posts],
        "redirect_url": url_for('add_candidate'),  # Optional, include if needed
        "no_doc_message": request.args.get('no_doc_message'),  # Optional, include if needed
        "career_count_notification_no": 0  # Placeholder, implement career notification logic
    }

    return jsonify(response_data)

# Helper function to convert JobPost object to dictionary
def job_post_to_dict(job_post):
    data_updated_date_str = job_post.data_updated_date.strftime('%Y-%m-%d') if job_post.data_updated_date else None
    data_updated_time_str = job_post.data_updated_time.strftime('%H:%M:%S') if job_post.data_updated_time else None

    return {
        "id": job_post.id,
        "client": job_post.client,
        "experience_min": job_post.experience_min,
        "experience_max": job_post.experience_max,
        "budget_min": job_post.budget_min,
        "budget_max": job_post.budget_max,
        "location": job_post.location,
        "shift_timings": job_post.shift_timings,
        "notice_period": job_post.notice_period,
        "role": job_post.role,
        "detailed_jd": job_post.detailed_jd,
        "mode": job_post.mode,
        "recruiter": job_post.recruiter,
        "management": job_post.management,
        "date_created": job_post.date_created.strftime('%Y-%m-%d'),
        "time_created": job_post.time_created.strftime('%H:%M:%S'),
        "job_status": job_post.job_status,
        "job_type": job_post.job_type,
        "skills": job_post.skills,
        "notification": job_post.notification,
        "data_updated_date": data_updated_date_str,
        "data_updated_time": data_updated_time_str
    }


from flask import jsonify

@app.route('/update_job_status/<int:job_id>', methods=['POST'])
def update_job_status(job_id):
    data = request.json
    user_id = data['user_id']
    user = User.query.filter_by(id=user_id).first()

    if not user:
        return jsonify({'status': 'error', "message": "User not found"}), 404

    user_type = user.user_type
    username = user.username

    # Retrieve the job post from the database based on the provided job_id
    job_post = JobPost.query.get(job_id)

    if job_post:
        try:
            # Extract the new job status from the form data
            new_job_status = data['new_job_status']

            # Update the job status
            job_post.job_status = new_job_status
            
            # Update data_updated_date and data_updated_time
            # current_datetime = datetime.now(pytz.timezone('Asia/Kolkata')) 
            # job_post.data_updated_date = current_datetime.date()
            # job_post.data_updated_time = current_datetime.time()
            
            # Commit the changes to the database
            db.session.commit()

            # Return a JSON response indicating success
            return jsonify({'status': 'success', "message": "Job status updated successfully"})
        
        except KeyError:
            # If 'new_job_status' key is missing in form data
            return jsonify({'status': 'error', "message": "Missing 'new_job_status' in form data"})
        
        except Exception as e:
            # Handle other exceptions
            db.session.rollback()  # Rollback any changes made to the session
            return jsonify({'status': 'error', "message": str(e)})

    # If job_post is None (job not found)
    return jsonify({'status': 'error', "message": "Job post not found"})




import base64

@app.route('/view_all_jobs', methods=['POST'])
def view_all_jobs():
    data = request.json
    user_name = data['username']

    # Define case statements for conditional ordering
    conditional_order_date = case(
        (JobPost.data_updated_date != None, JobPost.data_updated_date),
        (JobPost.date_created != None, JobPost.date_created),
        else_=JobPost.date_created
    )

    conditional_order_time = case(
        (JobPost.data_updated_time != None, JobPost.data_updated_time),
        (JobPost.time_created != None, JobPost.time_created),
        else_=JobPost.time_created
    )

    # Retrieve all job posts with conditional ordering
    job_posts_active = JobPost.query.filter_by(job_status='Active')\
        .order_by(
            desc(conditional_order_date),
            desc(conditional_order_time),
            desc(JobPost.id)  # Ensure newer posts appear first if dates are equal
        )\
        .all()

    job_posts_hold = JobPost.query.filter_by(job_status='Hold')\
        .order_by(
            desc(conditional_order_date),
            desc(conditional_order_time),
            desc(JobPost.id)  # Ensure newer posts appear first if dates are equal
        )\
        .all()

    # Construct JSON response
    response_data = {
        "user_name": user_name,
        "job_posts_active": [
            {
                "id": job_post.id,
                "client": job_post.client,
                "role": job_post.role,
                "experience_min": job_post.experience_min,
                "experience_max": job_post.experience_max,
                "budget_min": job_post.budget_min,
                "budget_max": job_post.budget_max,
                "location": job_post.location,
                "shift_timings": job_post.shift_timings,
                "notice_period": job_post.notice_period,
                "detailed_jd": job_post.detailed_jd,
                "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
                "mode": job_post.mode,
                "recruiter": job_post.recruiter,
                "management": job_post.management,
                "job_status": job_post.job_status,
                "job_type": job_post.job_type,
                "contract_in_months": job_post.contract_in_months,
                "skills": job_post.skills,
                # "date_created": str(job_post.date_created),
                # "time_created": str(job_post.time_created),
                # "data_updated_date": str(job_post.data_updated_date) if job_post.data_updated_date else None,
                # "data_updated_time": str(job_post.data_updated_time) if job_post.data_updated_time else None,
                "date_created": job_post.date_created.isoformat() if job_post.date_created else None,
                "time_created": job_post.time_created.strftime('%H:%M:%S') if job_post.time_created else None,
                "data_updated_date": job_post.data_updated_date.isoformat() if job_post.data_updated_date else None,
                "data_updated_time": job_post.data_updated_time.strftime('%H:%M:%S') if job_post.data_updated_time else None,
                "jd_pdf_present": job_post.jd_pdf_present,  # Added line
                "no_of_positions":job_post.no_of_positions
            }
            for job_post in job_posts_active
        ],
        "job_posts_hold": [
            {
                "id": job_post.id,
                "client": job_post.client,
                "role": job_post.role,
                "experience_min": job_post.experience_min,
                "experience_max": job_post.experience_max,
                "budget_min": job_post.budget_min,
                "budget_max": job_post.budget_max,
                "location": job_post.location,
                "shift_timings": job_post.shift_timings,
                "notice_period": job_post.notice_period,
                "detailed_jd": job_post.detailed_jd,
                "jd_pdf": base64.b64encode(job_post.jd_pdf).decode('utf-8') if job_post.jd_pdf else None,
                "mode": job_post.mode,
                "recruiter": job_post.recruiter,
                "management": job_post.management,
                "job_status": job_post.job_status,
                "job_type": job_post.job_type,
                "contract_in_months": job_post.contract_in_months,
                "skills": job_post.skills,
                # "date_created": str(job_post.date_created),
                # "time_created": str(job_post.time_created),
                # "data_updated_date": str(job_post.data_updated_date) if job_post.data_updated_date else None,
                # "data_updated_time": str(job_post.data_updated_time) if job_post.data_updated_time else None,
                "date_created": job_post.date_created.isoformat() if job_post.date_created else None,
                "time_created": job_post.time_created.strftime('%H:%M:%S') if job_post.time_created else None,
                "data_updated_date": job_post.data_updated_date.isoformat() if job_post.data_updated_date else None,
                "data_updated_time": job_post.data_updated_time.strftime('%H:%M:%S') if job_post.data_updated_time else None,
                "jd_pdf_present": job_post.jd_pdf_present,  # Added line
                 "no_of_positions":job_post.no_of_positions
            }
            for job_post in job_posts_hold
        ]
    }

    # Return JSON response
    return jsonify(response_data)
    



# def send_notification(recruiter_email):
#     msg = Message('New Job Posted', sender=config.sender_email, recipients=[recruiter_email])
#     msg.body = 'A new job has been posted. Check your dashboard for more details.'
#     mail.send(msg)

@app.route('/other_job_posts', methods=['GET'])
def other_job_posts():
    if 'user_id' in session and 'user_type' in session:
        if session['user_type'] == 'recruiter':
            # Retrieve the logged-in user's ID from the session
            user_id = session['user_id']

            # Retrieve the recruiter's name based on user ID
            recruiter_name = User.query.get(user_id).name

            job_posts = JobPost.query.filter(JobPost.recruiter != recruiter_name).distinct(JobPost.client).all()

            return render_template('other_job_posts.html', job_posts=job_posts)

    # Redirect or render an appropriate page if the conditions are not met
    return redirect(url_for('login'))



import base64
import io

from flask import Flask, send_file, request

@app.route('/view_resume/<int:candidate_id>', methods=['GET'])
def view_resume(candidate_id):
    # Retrieve the resume data from the database using SQLAlchemy
    candidate = Candidate.query.filter_by(id=candidate_id).first()
    if not candidate:
        return 'Candidate not found', 404

    if candidate.resume is None:
        return 'Resume not found for this candidate', 404

    try:
        # If the resume data is a base64 encoded string, decode it
        if isinstance(candidate.resume, str):
            resume_binary = base64.b64decode(candidate.resume)
        elif isinstance(candidate.resume, bytes):
            resume_binary = candidate.resume
        else:
            return 'Invalid resume format', 400

        # Determine the mimetype based on the file content
        is_pdf = resume_binary.startswith(b"%PDF")
        mimetype = 'application/pdf' if is_pdf else 'application/msword'

        # Send the file as a response
        return send_file(
            io.BytesIO(resume_binary),
            mimetype=mimetype,
            as_attachment=False
        )
    except Exception as e:
        return f'Error processing resume: {str(e)}', 500



@app.route('/upload_user_image/<int:user_id>', methods=['POST'])
def upload_user_image(user_id):
    try:
        # Extract data from the request
        data = request.json
        image_base64 = data['image']
        filename = data['filename']
        image_delete_status = data['image_delete_status']

        # Decode the base64 image
        image_binary = base64.b64decode(image_base64)

        # Find the user by user_id
        user = User.query.get(user_id)
        if not user:
            return jsonify({'error': 'User not found'}), 404

        # Update user's filename and image content
        user.filename = filename
        user.image_file = image_binary  # Store image content as binary data
        user.image_deleted = image_delete_status

        # Commit changes to the database
        db.session.commit()

        return jsonify({'message': 'Image updated successfully'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

#################################################################################################

import base64
import io
# @app.route('/image_status/<int:user_id>', methods=['GET'])
# def image_status(user_id):
#     user = User.query.filter_by(id=user_id).first()
#     if not user or not user.image_file:
#         return jsonify({'error': 'Image not found'}), 404

#     return jsonify({'message': user.image_file}), 200


from PIL import Image
import mimetypes
import imghdr

@app.route('/user_image/<int:user_id>', methods=['GET'])
def user_image(user_id):
    # Retrieve the user data from the database
    user = User.query.filter_by(id=user_id).first()
    if not user or not user.image_file:
        return jsonify({'message': 'Image not found'}), 404
    
    # Get the image data from the user object
    image_data = user.image_file
    
    # Determine the image format dynamically
    image_format = imghdr.what(None, h=image_data)
    if not image_format:
        return jsonify({'error': 'Unknown image format'}), 500
    
    # Determine the MIME type based on the image format
    if image_format == 'jpeg':
        mimetype = 'image/jpeg'
    elif image_format == 'png':
        mimetype = 'image/png'
    else:
        # Handle other image formats as needed
        return jsonify({'error': 'Unsupported image format'}), 500
    
    # Send the file as a response
    return send_file(
        io.BytesIO(image_data),
        mimetype=mimetype,
        as_attachment=False
    )



@app.route('/delete_user_image/<int:user_id>', methods=['POST'])
def delete_user_image(user_id):
    data = request.json
    profile_image = data['profileImage']
    image_delete_status=data['image_delete_status']
    if not profile_image:
        return jsonify({"error": "Profile image must be specified"}), 400

    user = User.query.filter_by(id=user_id).first()
    
    if not user:
        return jsonify({"error": "User not found"}), 404

    user.image_file = None
    user.filename = None
    user.image_deleted = image_delete_status
    db.session.commit()

    return jsonify({"message": "Image file deleted successfully"}), 200



    

@app.route('/viewfull_jd/<int:id>')
def viewfull_jd(id):
    user_type = session['user_type']
    job_post = JobPost.query.get(id)
    return render_template('viewfull_jd.html', job_post=job_post,user_type=user_type)

@app.route('/add_candidate_view')
def add_candidate_view():
    user_id = session['user_id']
    user_type = session['user_type']
    user_name = session['user_name']

    if user_type == 'recruiter':
        recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
        if recruiter:
            candidates = Candidate.query.filter_by(
                recruiter=recruiter.name).all()  # Filter candidates by recruiter's name
            # data = json.dumps(candidates, sort_keys=False)
            results = db.session.query(JobPost.client, JobPost.recruiter).filter(
                JobPost.recruiter.contains(user_name)).all()
            client_names = sorted(list(set([result.client for result in results])))
            count_notification_no = Notification.query.filter(Notification.notification_status == 'false',
                                                              Notification.recruiter_name == user_name).count()
            return render_template('add_candidate_view.html', user=recruiter, user_type=user_type, user_name=user_name,
                                   candidates=candidates, count_notification_no=count_notification_no,
                                   client_names=client_names)
    elif user_type == 'management':
        users = User.query.all()
        candidates = Candidate.query.all()
        JobsPosted = JobPost.query.all()
        clients = db.session.query(JobPost.client).all()
        client_names = list(set([client[0] for client in clients]))

        return render_template('add_candidate_view.html', users=users, user_type=user_type, user_name=user_name,
                               JobsPosted=JobsPosted, client_names=client_names)

import os
import shutil
from flask import Flask, request, send_file, redirect, url_for
from zipfile import ZipFile

@app.route('/download_resumes')
def download_resumes():
    candidate_ids = request.args.getlist('candidate_ids')
    
    # Create a temporary directory to store resume files
    temp_dir = 'temp_resumes'
    os.makedirs(temp_dir, exist_ok=True)
    
    resume_paths = []

    for candidate_id in candidate_ids:
        candidate = Candidate.query.get(candidate_id)
        if candidate is None or candidate.resume is None:
            continue
        
        resume_file = io.BytesIO(candidate.resume)
        is_pdf = resume_file.getvalue().startswith(b"%PDF")
        if is_pdf : 
            resume_filename = f"{candidate.name}_resume.pdf" 
            resume_path = os.path.join(temp_dir, resume_filename)
            with open(resume_path, 'wb') as file:
                file.write(candidate.resume)
            
            resume_paths.append(resume_path)
        else:
            resume_filename = f"{candidate.name}_resume.docx" 
            resume_path = os.path.join(temp_dir, resume_filename)
            with open(resume_path, 'wb') as file:
                file.write(candidate.resume)
            
            resume_paths.append(resume_path)

    # Create a zip file containing all resume files
    zip_filename = 'resumes.zip'
    with ZipFile(zip_filename, 'w') as zipf:
        for resume_path in resume_paths:
            zipf.write(resume_path, os.path.basename(resume_path))
    
    # Clean up temporary directory
    shutil.rmtree(temp_dir)
    
    # Send the zip file for download
    return send_file(zip_filename, as_attachment=True)


@app.route('/assign_job/<int:job_id>', methods=['POST'])
def assign_job(job_id):
    data = request.json
    user_id = data['user_id']
    user = User.query.filter_by(id=user_id).first()

    if not user:
        return jsonify({'status': 'error',"message": "User not found"})

    user_type = user.user_type
    username = user.username
    job_post = JobPost.query.get(job_id)  # Retrieve the job post by its ID

    if not job_post:
        return jsonify({'status': 'success',"message": "Job not found"})

    current_recruiters = job_post.recruiter.split(', ') if job_post.recruiter else []

    if request.method == 'POST':
        new_recruiter_names = data.get('recruiters', [])
        
        # Modification: Remove duplicate recruiters by combining lists and converting to a set
        updated_recruiter_names = list(set(current_recruiters + new_recruiter_names))
        
        # Join the recruiter names into a single string
        joined_recruiters = ', '.join(updated_recruiter_names)
        job_post.recruiter = joined_recruiters
        
        # Update data_updated_date and data_updated_time
        # current_datetime = datetime.now(pytz.timezone('Asia/Kolkata'))
        # job_post.data_updated_date = current_datetime.date()
        # job_post.data_updated_time = current_datetime.time()

        db.session.commit()

        # Send notification emails to the newly assigned recruiters
        new_recruiter_emails = [recruiter.email for recruiter in
                                User.query.filter(User.name.in_(new_recruiter_names),
                                                  User.user_type == 'recruiter')]
        for email in new_recruiter_emails:
            send_notification(email)

        # Define an empty list to hold Notification instances
        notifications = []

        for recruiter_name in updated_recruiter_names:
            if recruiter_name.strip() in new_recruiter_names:
                notification_status = False  # Set the initial status
                notification = Notification(
                    job_post_id=job_post.id,
                    recruiter_name=recruiter_name.strip(),
                    notification_status=notification_status
                )
                # Append each Notification instance to the notifications list
                notifications.append(notification)

        # Commit the notifications to the database session
        db.session.add_all(notifications)
        db.session.commit()

        return jsonify({'status': 'success',"message": "Job re-assigned successfully"}), 200

    recruiter_names = [recruiter.name for recruiter in User.query.filter_by(user_type='recruiter')]
    return jsonify({
        "user_name": username,
        "job_post": job_post.serialize(),
        "current_recruiters": current_recruiters,
        "recruiters": recruiter_names
    })



from flask import jsonify

@app.route('/disable_user', methods=['POST'])
def disable_user():
    data = request.json
    user_id = data['user_id']
    user_status = data['user_status']
    user_name = data['user_name']

    if user_id is None or user_status is None or user_name is None:
        return jsonify({'message': 'User ID, user status, and user name are required'}), 400

    # Find the user making the request
    request_user = User.query.get(user_id)

    if request_user is None or request_user.user_type != 'management':
        return jsonify({'message': 'Unauthorized access'}), 403

    # Find the user to be updated
    user = User.query.filter_by(username=user_name).first()

    if user is None:
        return jsonify({'message': 'User not found'}), 404

    # Change verification status for the user
    user.is_verified = user_status

    # If the user is a recruiter, change verification status for management user with the same username
    if user.user_type == 'recruiter':
        management_user = User.query.filter_by(username=user_name, user_type='management').first()
        if management_user:
            management_user.is_verified = user_status

    try:
        db.session.commit()
        # Return different messages based on user_type
        if user.user_type == 'management':
            if user_status:
                return jsonify({'message': 'Verification status updated for management account'}), 200
            else:
                return jsonify({'message': 'Verification status updated to unverified for management account'}), 200
        elif user.user_type == 'recruiter':
            if user_status:
                return jsonify({'message': 'Verification status updated for recruiter account'}), 200
            else:
                return jsonify({'message': 'Verification status updated to unverified for recruiter account'}), 200
    except Exception as e:
        # Log the exception or return an error message
        db.session.rollback()
        return jsonify({'message': 'Failed to update verification status'}), 500

@app.route('/active_users', methods=['POST'])
def update_user_status():
    data = request.json
    username = data.get('user_name')
    new_status = data.get('new_status')

    try:
        user = User.query.filter_by(username=username).first()
        if user:
            # user.is_verified = new_status
            # db.session.commit()

            # Fetch updated active users list
            active_users_manager = User.query.filter_by(user_type='management').all()
            active_users_manager = sorted(active_users_manager, key=lambda user: user.id)
            active_users_recruiter = User.query.filter_by(user_type='recruiter').all()
            active_users_recruiter = sorted(active_users_recruiter, key=lambda user: user.id)

            return jsonify({
                "message": "User status updated successfully",
                "username": username,
                "active_users_manager": [user.serialize() for user in active_users_manager],
                "active_users_recruiter": [user.serialize() for user in active_users_recruiter]
            })
        else:
            return jsonify({"message": "User not found"}), 404
    except Exception as e:
        db.session.rollback()  # Rollback changes in case of error
        return jsonify({"message": "Error updating user status", "error": str(e)}), 500
        

from flask import jsonify

@app.route('/deactivate_user', methods=['POST'])
def deactivate_user():
    data = request.json
    management_user_id = data.get('user_id')
    username = data.get('user_name')
    user_status = data.get('user_status')

    if management_user_id:
        # Find the management user
        management_user = User.query.get(management_user_id)

        if management_user and management_user.user_type == 'management':
            messages = []

            if username:
                # Find the target user by username
                target_user = User.query.filter_by(username=username).first()

                if target_user:
                    if target_user.user_type == 'management' or target_user.user_type == 'recruiter':
                        # Update user account status
                        target_user.is_active = user_status
                        db.session.commit()

                        # Determine the message based on user_status
                        if user_status:
                            messages.append(f'{target_user.user_type.capitalize()} account {username} has been successfully activated.')
                        else:
                            messages.append(f'{target_user.user_type.capitalize()} account {username} has been successfully deactivated.')
                    else:
                        messages.append('User is neither a management nor a recruiter account.')
                else:
                    messages.append('User not found.')

            if messages:
                # Get all user records
                all_users = User.query.all()
                user_data = [{'id': user.id, 'username': user.username, 'is_active': user.is_active} for user in all_users]
                return jsonify({'messages': messages, 'users': user_data})
            else:
                return jsonify({'message': 'No valid username provided.'})
        else:
            return jsonify({'message': 'Management user not found or not a management user.'})
    else:
        return jsonify({'message': 'Management user_id is required.'})



        
# @app.route('/verify_checkbox', methods=['POST'])
# def verify_checkbox():
#     data = request.json
#     user_id = data.get('userId')
#     checked = data.get('checked')
#     user = User.query.get(user_id)
#     user.is_verified = checked
#     db.session.commit()
#     return redirect(url_for('active_users'))

import hashlib
from flask_mail import Message


@app.route('/change_password', methods=['POST'])
def change_password():
    data = request.json

    if not data:
        return jsonify({'status': 'error', 'message': 'No JSON data provided.'})

    user_id = data.get('user_id')
    username = data.get('username')
    old_password = data.get('old_password')
    new_password = data.get('new_password')
    confirm_password = data.get('confirm_password')

    user = User.query.filter_by(id=user_id).first()

    if not user:
        return jsonify({'status': 'error', 'message': 'User not found.'})

    if username != user.username:
        return jsonify({'status': 'error', 'message': 'Logged in user does not match the provided username.'})

    hashed_old_password = hashlib.sha256(old_password.encode()).hexdigest()

    if user.password != hashed_old_password:
        return jsonify({'status': 'error', 'message': 'Invalid old password.'})

    if old_password == new_password:
        return jsonify({'status': 'error', 'message': f'New password cannot be the same as the old password.'})

    if new_password != confirm_password:
        return jsonify({'status': 'error', 'message': f'New password and confirm password is not matching.'})

    hashed_new_password = hashlib.sha256(new_password.encode()).hexdigest()
    user.password = hashed_new_password
    db.session.commit()

    msg = Message('Password Changed', sender=config.sender_email, recipients=[user.email])
    msg.html = f'''
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <style>
            body {{
                font-family: Arial, sans-serif;
                background-color: #f4f4f4;
                color: #333;
                margin: 0;
                padding: 20px;
            }}
            .container {{
                background-color: #ffffff;
                max-width: 600px;
                margin: 0 auto;
                padding: 20px;
                border: 1px solid #dddddd;
                border-radius: 8px;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                font-size: 24px;
                font-weight: bold;
                margin-bottom: 20px;
                color: #4CAF50;
            }}
            .content {{
                font-size: 16px;
                line-height: 1.6;
            }}
            .credentials {{
                background-color: #f9f9f9;
                padding: 10px;
                border: 1px solid #eeeeee;
                border-radius: 5px;
                margin-top: 10px;
            }}
            .footer {{
                font-size: 12px;
                color: #999;
                margin-top: 20px;
                text-align: center;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">Password Changed</div>
            <div class="content">
                <p>Hello {user.username},</p>
                <p>Your password has been successfully changed.</p>
                <p>Here are your updated credentials:</p>
                <div class="credentials">
                    <p><strong>Username:</strong> {user.username}</p>
                    <p><strong>Password:</strong> {new_password}</p>
                </div>
            </div>
            <div class="footer">
                <p>If you did not request this change, please contact our support team immediately.</p>
                <p><b>Makonis Talent Track Pro Team</b></p>
            </div>
        </div>
    </body>
    </html>
    '''
    mail.send(msg)

    if user.user_type == 'management':
        return jsonify({'status': 'success', 'message': f'Password changed successfully for management {user.username}.'})
    elif user.user_type == 'recruiter':
        return jsonify({'status': 'success', 'message': f'Password changed successfully for recruiter {user.username}.'})
    else:
        return jsonify({'status': 'success', 'message': 'Password changed successfully.'})





@app.route('/delete_job_post_message/<int:job_id>')
def delete_job_post_message(job_id):
    job_post = JobPost.query.get(job_id)
    id = job_post.id
    client = job_post.client
    role = job_post.role
    return redirect(url_for('view_all_jobs',client=client,role=role,id=id))


@app.route('/delete_job_post/<int:job_id>', methods=['POST'])
def delete_job_post(job_id):
    # Fetch the job post
    job_post = JobPost.query.get(job_id)
    
    if not job_post:
        return jsonify({'status': 'error',"message": "Job Post not found"}), 404

    # Fetch all notifications related to the job post
    notifications = Notification.query.filter_by(job_post_id=job_id).all()

    if notifications:
        try:
            # Delete all associated notifications first
            for notification in notifications:
                db.session.delete(notification)
            db.session.commit()
            
            # Now delete the job post
            db.session.delete(job_post)
            db.session.commit()
            
            return jsonify({'status': 'success',"message": "Job Post and Notifications Deleted Successfully"})
        except Exception as e:
            # Handle any potential exceptions
            db.session.rollback()
            return jsonify({'status': 'error',"message": "An error occurred while deleting job post and notifications"})
    else:
        # No notifications found for the job post
        # Delete only the job post
        db.session.delete(job_post)
        db.session.commit()
        return jsonify({'status': 'success',"message": "Job Post and Notifications Deleted Successfully"})
        # return jsonify({'status': 'success',"message": "Job Post Deleted Successfully. No associated notifications found."}), 200


# @app.route('/delete_job_post/<int:job_id>', methods=['POST'])
# def delete_job_post(job_id):
#     # Fetch the job post
#     job_post = JobPost.query.get(job_id)
    
#     if not job_post:
#         return jsonify({"error": "Job Post not found"}), 404

#     # Fetch all notifications related to the job post
#     notifications = Notification.query.filter_by(job_post_id=job_id).all()

#     if notifications:
#         # Delete the job post and all associated notifications
#         db.session.delete(job_post)
#         for notification in notifications:
#             db.session.delete(notification)
#         db.session.commit()
#         return jsonify({"message": "Job Post and Notifications Deleted Successfully"}), 200
#     else:
#         # No notifications found for the job post
#         # Delete only the job post
#         db.session.delete(job_post)
#         db.session.commit()
#         return jsonify({"message": "Job Post Deleted Successfully. No associated notifications found."}), 200

# @app.route('/delete_job_post/<int:job_id>', methods=['POST'])
# def delete_job_post(job_id):
#     # data=request.json
#     # job_id=data['job_id']
#     job_post = JobPost.query.get(job_id)
#     if job_post:
#         JobPost.query.filter_by(id=job_id).delete()
#         db.session.commit()
#         return jsonify({"message": "Job Post Deleted Successfully"}), 200
#     else:
#         return jsonify({"error": "Job Post not found"}), 404

@app.route('/download_jd/<int:job_id>')
def download_jd(job_id):
    jobpost = JobPost.query.get(job_id)
    if jobpost is None or jobpost.jd_pdf is None:
        return redirect(url_for('dashboard'))

    jd_file = io.BytesIO(jobpost.jd_pdf)
    is_pdf = jd_file.getvalue().startswith(b"%PDF")
    if is_pdf : 
        jd_filename = f"{jobpost.client}_jd.pdf"  # Set the filename as desired
        jd_path = os.path.join(app.config['UPLOAD_FOLDER'], jd_filename)
        with open(jd_path, 'wb') as file:
            file.write(jobpost.jd_pdf)

        # Send the saved resume file for download
        return send_file(jd_path, as_attachment=True)
    else:
        jd_filename = f"{jobpost.client}_jd.docx"  # Set the filename as desired
        jd_path = os.path.join(app.config['UPLOAD_FOLDER'], jd_filename)
        with open(jd_path, 'wb') as file:
            file.write(jobpost.jd_pdf)

        # Send the saved resume file for download
        return send_file(jd_path, as_attachment=True)


import base64
import io
from flask import send_file



@app.route('/view_jd/<int:job_id>', methods=['GET'])
def view_jd(job_id):
    # Retrieve the resume data from the database using SQLAlchemy
    jobpost = JobPost.query.filter_by(id=job_id).first()
    if not jobpost:
        return 'Job post not found', 404  # Return 404 Not Found status

    if jobpost.jd_pdf is None:
        return 'jd_pdf not found for this job', 404

    try:
        # If the resume data is a base64 encoded string, decode it
        if isinstance(jobpost.jd_pdf, str):
            jd_pdf_binary = base64.b64decode(jobpost.jd_pdf)
        elif isinstance(jobpost.jd_pdf, bytes):
            jd_pdf_binary = jobpost.jd_pdf
        else:
            return 'Invalid jd_pdf format', 400

        # Determine the mimetype based on the file content
        is_pdf = jd_pdf_binary.startswith(b"%PDF")
        mimetype = 'application/pdf' if is_pdf else 'application/msword'

        # Send the file as a response
        return send_file(
            io.BytesIO(jd_pdf_binary),
            mimetype=mimetype,
            as_attachment=False
        )
    except Exception as e:
        return f'Error processing resume: {str(e)}', 500




from flask import Flask, request, jsonify
from datetime import datetime, timedelta
# from sqlalchemy import func
from sqlalchemy import func, text,case
from sqlalchemy.sql import text 
from sqlalchemy import text
import pandas as pd  # Import pandas for date_range

from io import BytesIO
import base64

# from sqlalchemy.orm import sessionmaker
# from sqlalchemy import create_engine

# import plotly.express as px
# import plotly.io as pio

# import plotly.express as px
# import plotly.io as pio
# from flask import Flask, request, jsonify
# from sqlalchemy import func, extract
# from datetime import datetime
# import io
# import base64


############## Here i am converting the username to name ####################################


# @app.route('/analyze_recruitment', methods=['POST', 'GET'])
# def analyze_recruitment():
#     data = request.json

#     if not data:
#         return jsonify({'error': 'No JSON data provided'})

#     recruiter_usernames = data.get('recruiter_names', [])

#     if not recruiter_usernames:
#         return jsonify({'error': 'Please select any Recruiter'})

#     # Fetch recruiter names from the User table
#     recruiter_names = db.session.query(User.username, User.name).filter(User.username.in_(recruiter_usernames)).all()

#     if not recruiter_names:
#         return jsonify({'error': 'No matching recruiters found'})

#     recruiter_names_dict = {username: name for username, name in recruiter_names}
#     print("recruiter_names_dict : ",recruiter_names_dict)

#     try:
#         from_date_str = data.get('from_date')
#         to_date_str = data.get('to_date')
#         from_date = datetime.strptime(from_date_str, "%d-%m-%Y")
#         to_date = datetime.strptime(to_date_str, "%d-%m-%Y")
#     except ValueError:
#         return jsonify({'status': 'error', 'message': 'Invalid date format. Please use DD-MM-YYYY format.'})

#     from_date = from_date.strftime("%Y-%m-%d")
#     to_date = to_date.strftime("%Y-%m-%d")

#     recruiter_data = {}
#     total_candidate_count = 0
#     total_selected_candidates = 0
#     total_rejected_candidates_count = 0
#     total_process_candidates = 0

#     for recruiter_username in recruiter_usernames:
#         recruiter_name = recruiter_names_dict.get(recruiter_username, recruiter_username)
#         print("recruiter_name :",recruiter_name)
        
#         candidates_query = db.session.query(Candidate).filter(
#             Candidate.recruiter == recruiter_name,
#             Candidate.date_created >= from_date,
#             Candidate.date_created <= to_date
#         )

#         candidates = candidates_query.all()
#         recruiter_candidate_count = candidates_query.count()
#         print("recruiter_candidate_count : ",recruiter_candidate_count)
#         total_candidate_count += recruiter_candidate_count
#         print("total_candidate_count :",total_candidate_count)

#         if recruiter_candidate_count > 0:   
#             selected_candidates_count = candidates_query.filter(Candidate.status == 'ON-BOARDED').count()
#             in_process_candidates_count = candidates_query.filter(Candidate.status.notin_([
#                 'SCREENING', 'SCREENING SELECTED','L1 - SCHEDULE', 'L1 - FEEDBACK', 'L1 - SELECTED', 'L1 - CANDIDATE RESCHEDULE','L3 - CANDIDATE RESCHEDULE', 
#                 'L3 - CANDIDATE RESCHEDULE','L1 - PANEL RESCHEDULE', 'L2 - PANEL RESCHEDULE','L3 - PANEL RESCHEDULE','L2 - SCHEDULE', 'L2 - FEEDBACK','L3 - FEEDBACK' ,'L2 - SELECTED', 'L3 - SELECTED' ,'L3 - SCHEDULE','HR - ROUND', 'MANAGERIAL ROUND',
#                 'NEGOTIATION', 'SELECTED','None'])).count()
#             rejected_candidates_count = candidates_query.filter(Candidate.status.in_([
#                 'SCREEN REJECTED', 'L1 - REJECTED', 'L2 - REJECTED', 'L3 - REJECTED', 'OFFER - DECLINED', 
#                 'OFFER-REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO - SHOW','NO SHOW','CANDIDATE HOLD'])).count()
#         else:
#             selected_candidates_count = 0      
#             in_process_candidates_count = 0
#             rejected_candidates_count = 0
#         print("rejected_candidates_count : ",rejected_candidates_count)
#         total_selected_candidates += selected_candidates_count
#         total_process_candidates += in_process_candidates_count
#         total_rejected_candidates_count += rejected_candidates_count

#         in_process_candidates = recruiter_candidate_count - (selected_candidates_count + rejected_candidates_count)
        
#         total_process_candidates_count = total_candidate_count - (total_selected_candidates + total_rejected_candidates_count)

#         # role_industry_location_analysis_result = get_role_industry_location_analysis(
#         #     recruiter_username, from_date, to_date)

#         conversion_rate = get_conversion_rate(candidates_query)
#         # analysis_result = get_role_industry_location_analysis(recruiter_username, from_date, to_date)
#         analysis_result = get_role_industry_location_analysis()
#         # time_to_close_analysis = get_time_to_close_analysis(recruiter_usernames)
#         historical_performance_analysis = calculate_historical_performance_analysis(recruiter_usernames, from_date, to_date)
#         client_closure_rates, highest_closure_client, lowest_closure_client, _, _ = get_client_closure_rates(candidates_query)

#         percentage_of_selected = (selected_candidates_count / recruiter_candidate_count) * 100 if recruiter_candidate_count > 0 else 0.0

#         recruiter_data[recruiter_name] = {
#             'submission_counts_daily': get_submission_counts(candidates_query, from_date, to_date, 'daily'),
#             'submission_counts_weekly': get_submission_counts(candidates_query, from_date, to_date, 'weekly'),
#             'submission_counts_monthly': get_submission_counts(candidates_query, from_date, to_date, 'monthly'),
#             'submission_counts_yearly': get_submission_counts(candidates_query, from_date, to_date, 'yearly'),
#             'selected_candidates_count': selected_candidates_count,
#             'rejected_candidates_count': rejected_candidates_count,
#             'in_process_candidates_count': in_process_candidates,
#             'conversion_rate': conversion_rate,
#             'client_closure_rates': client_closure_rates,
#             'highest_closure_client': highest_closure_client,
#             'lowest_closure_client': lowest_closure_client,
#             'candidate_count': recruiter_candidate_count,
#             'percentage_of_selected': percentage_of_selected,
#             'candidates': [{
#                 'candidate_name': candidate.name,
#                 'job_id': candidate.job_id,
#                 'client': candidate.client,
#                 'recruiter': candidate.recruiter,
#                 'date_created': candidate.date_created.strftime('%Y-%m-%d'),
#                 'time_created': candidate.date_created.strftime('%H:%M:%S'),
#                 'profile': candidate.profile,
#                 'last_working_date': candidate.last_working_date.strftime('%Y-%m-%d') if candidate.last_working_date else None,
#                 'status': candidate.status
#             } for candidate in candidates]
#         }

#     ranked_recruiters = sorted(recruiter_data.items(), key=lambda x: (x[1]['selected_candidates_count'] / x[1]['candidate_count'] if x[1]['candidate_count'] > 0 else 0), reverse=True)

#     rank = 0

#     for recruiter, data in ranked_recruiters:
#         if data['candidate_count'] > 0:
#             if data['selected_candidates_count'] > 0:
#                 rank += 1
#                 recruiter_data[recruiter]['ranking'] = rank
#             else:
#                 recruiter_data[recruiter]['ranking'] = 0
#         else:
#             recruiter_data[recruiter]['ranking'] = 0

#     response_data = {
#         'status': 'success',
#         # 'time_to_close_analysis':time_to_close_analysis,
#         'historical_performance_analysis':historical_performance_analysis,
#         'Job_Type_Analysis': analysis_result,
#         'recruiter_data': recruiter_data,
#         'total_candidate_count': total_candidate_count,
#         'total_selected_candidates': total_selected_candidates,
#         'total_rejected_candidates_count': total_rejected_candidates_count,
#         # 'total_process_candidates_count': total_process_candidates,
#         'total_process_candidates_count': total_process_candidates_count,
#         'from_date_str': from_date_str,
#         'to_date_str': to_date_str,
#         'message': 'Ranking calculations completed successfully',
#         'bar_graph_data': generate_bar_graph_data(recruiter_data)
#     }

#     return jsonify(response_data)


import plotly.express as px
import plotly.io as pio
from flask import Flask, request, jsonify
from sqlalchemy import func, extract
from datetime import datetime
import io
import base64

@app.route('/analyze_recruitment', methods=['POST','GET'])
def analyze_recruitment():
    data = request.json

    if not data:
        return jsonify({'error': 'No JSON data provided'})

    recruiter_usernames = data.get('recruiter_names', [])

    if not recruiter_usernames:
        return jsonify({'error': 'Please select any Recruiter'})

    try:
        from_date_str = data.get('from_date')
        to_date_str = data.get('to_date')
        from_date = datetime.strptime(from_date_str, "%d-%m-%Y")
        to_date = datetime.strptime(to_date_str, "%d-%m-%Y")
    except ValueError:
        return jsonify({'status': 'error', 'message': 'Invalid date format. Please use DD-MM-YYYY format.'})

    from_date = from_date.strftime("%Y-%m-%d")
    to_date = to_date.strftime("%Y-%m-%d")

    recruiter_data = {}
    total_candidate_count = 0
    total_selected_candidates = 0
    total_rejected_candidates_count = 0
    total_process_candidates = 0

    for recruiter_username in recruiter_usernames:
        candidates_query = db.session.query(Candidate).filter(
            Candidate.recruiter == recruiter_username,
            Candidate.date_created >= from_date,
            Candidate.date_created <= to_date
        )

        candidates = candidates_query.all()
        recruiter_candidate_count = candidates_query.count()
        total_candidate_count += recruiter_candidate_count

        if recruiter_candidate_count > 0:
            selected_candidates_count = candidates_query.filter(Candidate.status == 'ON-BOARDED').count()
            in_process_candidates_count = candidates_query.filter(Candidate.status.notin_([
                'SCREENING', 'SCREENING SELECTED','L1-SCHEDULE', 'L1-FEEDBACK', 'L1-SELECTED', 'L1-CANDIDATE RESCHEDULE','L3-CANDIDATE RESCHEDULE', 
                'L3-CANDIDATE RESCHEDULE','L1-PANEL RESCHEDULE', 'L2-PANEL RESCHEDULE','L3-PANEL RESCHEDULE','L2-SCHEDULE', 'L2-FEEDBACK','L3-FEEDBACK' ,'L2-SELECTED', 'L3-SELECTED' ,'L3-SCHEDULE','HR-ROUND', 'MANAGERIAL ROUND',
                'NEGOTIATION', 'SELECTED'])).count()
            rejected_candidates_count = candidates_query.filter(Candidate.status.in_([
                'SCREEN REJECTED', 'L1-REJECTED', 'L2-REJECTED', 'L3-REJECTED', 'OFFER-DECLINED', 
                'OFFER-REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO-SHOW','NO SHOW'])).count()
        else:
            selected_candidates_count = 0
            in_process_candidates_count = 0
            rejected_candidates_count = 0

        total_selected_candidates += selected_candidates_count
        total_process_candidates += in_process_candidates_count
        total_rejected_candidates_count += rejected_candidates_count

        in_process_candidates = recruiter_candidate_count - (selected_candidates_count + rejected_candidates_count)
        
        total_process_candidates_count = total_candidate_count - (total_selected_candidates + total_rejected_candidates_count)

        # role_industry_location_analysis_result = get_role_industry_location_analysis(
        #     recruiter_username, from_date, to_date)

        conversion_rate = get_conversion_rate(candidates_query)
        # analysis_result = get_role_industry_location_analysis(recruiter_username, from_date, to_date)
        analysis_result = get_role_industry_location_analysis()
        # time_to_close_analysis = get_time_to_close_analysis(recruiter_usernames)
        historical_performance_analysis = calculate_historical_performance_analysis(recruiter_usernames, from_date, to_date)
        client_closure_rates, highest_closure_client, lowest_closure_client, _, _ = get_client_closure_rates(candidates_query)

        percentage_of_selected = (selected_candidates_count / recruiter_candidate_count) * 100 if recruiter_candidate_count > 0 else 0.0

        recruiter_data[recruiter_username] = {
            'submission_counts_daily': get_submission_counts(candidates_query, from_date, to_date, 'daily'),
            'submission_counts_weekly': get_submission_counts(candidates_query, from_date, to_date, 'weekly'),
            'submission_counts_monthly': get_submission_counts(candidates_query, from_date, to_date, 'monthly'),
            'submission_counts_yearly': get_submission_counts(candidates_query, from_date, to_date, 'yearly'),
            'selected_candidates_count': selected_candidates_count,
            'rejected_candidates_count': rejected_candidates_count,
            'in_process_candidates_count': in_process_candidates,
            'conversion_rate': conversion_rate,
            'client_closure_rates': client_closure_rates,
            'highest_closure_client': highest_closure_client,
            'lowest_closure_client': lowest_closure_client,
            'candidate_count': recruiter_candidate_count,
            'percentage_of_selected': percentage_of_selected,
            'candidates': [{
                'candidate_name': candidate.name,
                'job_id': candidate.job_id,
                'client': candidate.client,
                'recruiter': candidate.recruiter,
                'date_created': candidate.date_created.strftime('%Y-%m-%d'),
                'time_created': candidate.date_created.strftime('%H:%M:%S'),
                'profile': candidate.profile,
                'last_working_date': candidate.last_working_date.strftime('%Y-%m-%d') if candidate.last_working_date else None,
                'status': candidate.status
            } for candidate in candidates]
        }

    ranked_recruiters = sorted(recruiter_data.items(), key=lambda x: (x[1]['selected_candidates_count'] / x[1]['candidate_count'] if x[1]['candidate_count'] > 0 else 0), reverse=True)

    rank = 0

    for recruiter, data in ranked_recruiters:
        if data['candidate_count'] > 0:
            if data['selected_candidates_count'] > 0:
                rank += 1
                recruiter_data[recruiter]['ranking'] = rank
            else:
                recruiter_data[recruiter]['ranking'] = 0
        else:
            recruiter_data[recruiter]['ranking'] = 0

    response_data = {
        'status': 'success',
        # 'time_to_close_analysis':time_to_close_analysis,
        'historical_performance_analysis':historical_performance_analysis,
        'Job_Type_Analysis': analysis_result,
        'recruiter_data': recruiter_data,
        'total_candidate_count': total_candidate_count,
        'total_selected_candidates': total_selected_candidates,
        'total_rejected_candidates_count': total_rejected_candidates_count,
        # 'total_process_candidates_count': total_process_candidates,
        'total_process_candidates_count': total_process_candidates_count,
        'from_date_str': from_date_str,
        'to_date_str': to_date_str,
        'message': 'Ranking calculations completed successfully',
        'bar_graph_data': generate_bar_graph_data(recruiter_data)
    }

    return jsonify(response_data)
    
def calculate_historical_performance_analysis(recruiter_usernames, from_date_str, to_date_str):
    try:
        from_date = datetime.strptime(from_date_str, "%Y-%m-%d")
        to_date = datetime.strptime(to_date_str, "%Y-%m-%d")
    except ValueError:
        return {'status': 'error', 'message': 'Invalid date format. Please use YYYY-MM-DD format.'}

    result = {}

    for recruiter_name in recruiter_usernames:
        # Initialize data structures
        monthly_data = {}
        total_months = (to_date.year - from_date.year) * 12 + (to_date.month - from_date.month) + 1

        for i in range(total_months):
            # Calculate start and end of current month
            start_date = from_date + relativedelta(months=i)
            end_date = start_date + relativedelta(months=1) - relativedelta(days=1)

            # Query candidates for the recruiter within the current month (mockup query)
            candidates = db.session.query(Candidate).filter(
                Candidate.recruiter == recruiter_name,
                Candidate.date_created >= start_date,
                Candidate.date_created <= end_date,
                Candidate.status.in_(['SCREENING', 'ON-BOARDED', 'SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW'])
            ).all()

            candidates_data = []
            total_screening_candidates = 0
            total_days_to_close = 0
            count_of_onboarded_positions = 0
            total_candidates = len(candidates)
            unsuccessful_closures = 0

            for candidate in candidates:
                days_to_close = None  # Initialize days_to_close

                if candidate.status == 'SCREENING':
                    total_screening_candidates += 1
                elif candidate.status == 'ON-BOARDED':
                    count_of_onboarded_positions += 1

                    # Calculate days to close (mockup calculation)
                    if candidate.date_created and candidate.data_updated_date:
                        days_to_close = (candidate.data_updated_date - candidate.date_created).days
                        total_days_to_close += days_to_close

                    # Prepare candidate data (mockup preparation)
                    candidate_data = {
                        'candidate_name': candidate.name,
                        'job_id': candidate.job_id,  # Assuming job_id is a regular column
                        'client': candidate.client,
                        'recruiter': candidate.recruiter,
                        'date_created': candidate.date_created.strftime('%Y-%m-%d') if candidate.date_created else None,
                        'date_updated': candidate.data_updated_date.strftime('%Y-%m-%d') if candidate.data_updated_date else None,
                        'days_to_close': days_to_close if candidate.status == 'ON-BOARDED' else None,
                        'profile': candidate.profile,
                        'status': candidate.status
                    }
                    candidates_data.append(candidate_data)
                elif candidate.status in ['SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW']:
                    unsuccessful_closures += 1

            # Calculate average days to close (mockup calculation)
            average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

            # Calculate percentage of onboarded candidates (mockup calculation)
            percentage_onboarded = (count_of_onboarded_positions / total_candidates) * 100 if total_candidates > 0 else 0

            # Store monthly data
            monthly_data[start_date.strftime('%Y-%m')] = {
                'candidates': candidates_data,
                'total_days_to_close': total_days_to_close,
                'count_of_screening_candidates': total_screening_candidates,
                'count_of_onboarded_positions': count_of_onboarded_positions,
                'unsuccessful_closures': unsuccessful_closures,
                'average_days_to_close': average_days_to_close,
                'percentage_onboarded': percentage_onboarded,
                'total_candidates_count': total_candidates
            }

        # Evaluate trend in closure rates (mockup evaluation)
        first_month_data = next(iter(monthly_data.values()), None)
        last_month_data = next(iter(list(monthly_data.values())[::-1]), None)  # Convert dict_values to list and then reverse

        if first_month_data and last_month_data:
            initial_onboarded_percentage = first_month_data.get('percentage_onboarded', 0)
            final_onboarded_percentage = last_month_data.get('percentage_onboarded', 0)

            if final_onboarded_percentage > initial_onboarded_percentage:
                trend = 'improving'
            elif final_onboarded_percentage < initial_onboarded_percentage:
                trend = 'declining'
            else:
                trend = 'stable'

            # Prepare line graph data
            line_graph_data = prepare_line_graph_data(recruiter_name, monthly_data)

            result[recruiter_name] = {
                'line_graph_data': line_graph_data,
                'monthly_data': monthly_data,
                'overall_summary': {
                    'total_months_analyzed': total_months,
                    'trend_in_closure_rates': trend
                }
            }
        else:
            result[recruiter_name] = {
                'monthly_data': monthly_data,
                'overall_summary': {
                    'total_months_analyzed': total_months,
                    'trend_in_closure_rates': 'insufficient data'  # Or handle as appropriate
                }
            }

    return result

# Function to prepare line graph data
def prepare_line_graph_data(recruiter_name, monthly_data):
    line_graph_data = {}

    for month, data in monthly_data.items():
        line_graph_data[month] = {
            'Total Days to Close': data['total_days_to_close'],
            'Screening Candidates': data['count_of_screening_candidates'],
            'Onboarded Positions': data['count_of_onboarded_positions'],
            'Unsuccessful Closures': data['unsuccessful_closures'],
            'Average Days to Close': data['average_days_to_close'],
            'Percentage Onboarded': data['percentage_onboarded'],
            'Total Candidates Count': data['total_candidates_count']
        }

    return line_graph_data
    
def calculate_historical_performance_analysis(recruiter_usernames, from_date_str, to_date_str):
    try:
        from_date = datetime.strptime(from_date_str, "%Y-%m-%d")
        to_date = datetime.strptime(to_date_str, "%Y-%m-%d")
    except ValueError:
        return {'status': 'error', 'message': 'Invalid date format. Please use YYYY-MM-DD format.'}

    result = {}

    # Fetch the recruiter names from the User table
    recruiters = db.session.query(User.username, User.name).filter(User.username.in_(recruiter_usernames)).all()
    username_to_name_map = {recruiter.username: recruiter.name for recruiter in recruiters}

    for recruiter_username in recruiter_usernames:
        recruiter_name = username_to_name_map.get(recruiter_username, recruiter_username)
        # Initialize data structures
        monthly_data = {}
        total_months = (to_date.year - from_date.year) * 12 + (to_date.month - from_date.month) + 1

        for i in range(total_months):
            # Calculate start and end of current month
            start_date = from_date + relativedelta(months=i)
            end_date = start_date + relativedelta(months=1) - relativedelta(days=1)

            # Query candidates for the recruiter within the current month (mockup query)
            candidates = db.session.query(Candidate).filter(
                Candidate.recruiter == recruiter_name,
                Candidate.date_created >= start_date,
                Candidate.date_created <= end_date,
                Candidate.status.in_(['SCREENING', 'ON-BOARDED', 'SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW'])
            ).all()

            candidates_data = []
            total_screening_candidates = 0
            total_days_to_close = 0
            count_of_onboarded_positions = 0
            total_candidates = len(candidates)
            unsuccessful_closures = 0

            for candidate in candidates:
                days_to_close = None  # Initialize days_to_close

                if candidate.status == 'SCREENING':
                    total_screening_candidates += 1
                elif candidate.status == 'ON-BOARDED':
                    count_of_onboarded_positions += 1

                    # Calculate days to close (mockup calculation)
                    if candidate.date_created and candidate.data_updated_date:
                        days_to_close = (candidate.data_updated_date - candidate.date_created).days
                        total_days_to_close += days_to_close

                    # Prepare candidate data (mockup preparation)
                    candidate_data = {
                        'candidate_name': candidate.name,
                        'job_id': candidate.job_id,  # Assuming job_id is a regular column
                        'client': candidate.client,
                        'recruiter': candidate.recruiter,
                        'date_created': candidate.date_created.strftime('%Y-%m-%d') if candidate.date_created else None,
                        'date_updated': candidate.data_updated_date.strftime('%Y-%m-%d') if candidate.data_updated_date else None,
                        'days_to_close': days_to_close if candidate.status == 'ON-BOARDED' else None,
                        'profile': candidate.profile,
                        'status': candidate.status
                    }
                    candidates_data.append(candidate_data)
                elif candidate.status in ['SCREEN REJECTED', 'L1 REJECTED', 'L2 REJECTED', 'L3 REJECTED', 'OFFER DECLINED/REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO SHOW']:
                    unsuccessful_closures += 1

            # Calculate average days to close (mockup calculation)
            average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

            # Calculate percentage of onboarded candidates (mockup calculation)
            percentage_onboarded = (count_of_onboarded_positions / total_candidates) * 100 if total_candidates > 0 else 0

            # Store monthly data
            monthly_data[start_date.strftime('%Y-%m')] = {
                'candidates': candidates_data,
                'total_days_to_close': total_days_to_close,
                'count_of_screening_candidates': total_screening_candidates,
                'count_of_onboarded_positions': count_of_onboarded_positions,
                'unsuccessful_closures': unsuccessful_closures,
                'average_days_to_close': average_days_to_close,
                'percentage_onboarded': percentage_onboarded,
                'total_candidates_count': total_candidates
            }

        # Evaluate trend in closure rates (mockup evaluation)
        first_month_data = next(iter(monthly_data.values()), None)
        last_month_data = next(iter(list(monthly_data.values())[::-1]), None)  # Convert dict_values to list and then reverse

        if first_month_data and last_month_data:
            initial_onboarded_percentage = first_month_data.get('percentage_onboarded', 0)
            final_onboarded_percentage = last_month_data.get('percentage_onboarded', 0)

            if final_onboarded_percentage > initial_onboarded_percentage:
                trend = 'improving'
            elif final_onboarded_percentage < initial_onboarded_percentage:
                trend = 'declining'
            else:
                trend = 'stable'

            # Prepare line graph data
            line_graph_data = prepare_line_graph_data(recruiter_name, monthly_data)

            result[recruiter_name] = {
                'line_graph_data': line_graph_data,
                'monthly_data': monthly_data,
                'overall_summary': {
                    'total_months_analyzed': total_months,
                    'trend_in_closure_rates': trend
                }
            }
        else:
            result[recruiter_name] = {
                'monthly_data': monthly_data,
                'overall_summary': {
                    'total_months_analyzed': total_months,
                    'trend_in_closure_rates': 'insufficient data'  # Or handle as appropriate
                }
            }

    return result

# Function to prepare line graph data
def prepare_line_graph_data(recruiter_name, monthly_data):
    line_graph_data = {}

    for month, data in monthly_data.items():
        line_graph_data[month] = {
            'Total Days to Close': data['total_days_to_close'],
            'Screening Candidates': data['count_of_screening_candidates'],
            'Onboarded Positions': data['count_of_onboarded_positions'],
            'Unsuccessful Closures': data['unsuccessful_closures'],
            'Average Days to Close': data['average_days_to_close'],
            'Percentage Onboarded': data['percentage_onboarded'],
            'Total Candidates Count': data['total_candidates_count']
        }

    return line_graph_data




def get_submission_counts(candidates_query, from_date, to_date, interval):
    if interval == 'daily':
        grouped_query = candidates_query.filter(
            Candidate.date_created >= from_date,
            Candidate.date_created <= to_date
        ).group_by(func.DATE(Candidate.date_created)).with_entities(
            func.DATE(Candidate.date_created).label('date_part'),
            func.count().label('count')
        )
    elif interval == 'weekly':
        grouped_query = candidates_query.filter(
            Candidate.date_created >= from_date,
            Candidate.date_created <= to_date
        ).group_by(func.DATE(Candidate.date_created)).with_entities(
            func.DATE(Candidate.date_created).label('date_part'),
            func.count().label('count')
        )
    elif interval == 'monthly':
        grouped_query = candidates_query.filter(
            Candidate.date_created >= from_date,
            Candidate.date_created <= to_date
        ).group_by(func.TO_CHAR(Candidate.date_created, 'YYYY-MM')).with_entities(
            func.TO_CHAR(Candidate.date_created, 'YYYY-MM').label('date_part'),
            func.count().label('count')
        )
    elif interval == 'yearly':
        grouped_query = candidates_query.filter(
            Candidate.date_created >= from_date,
            Candidate.date_created <= to_date
        ).group_by(func.TO_CHAR(Candidate.date_created, 'YYYY')).with_entities(
            func.TO_CHAR(Candidate.date_created, 'YYYY').label('date_part'),
            func.count().label('count')
        )
    else:
        return []

    submission_counts = grouped_query.all()
    return [{'date_part': str(item.date_part), 'count': item.count} for item in submission_counts]

def get_role_industry_location_analysis():  # mugilan
    # Get all distinct roles from JobPost table
    roles = db.session.query(
        JobPost.role,
        JobPost.location,
        JobPost.job_type,
        JobPost.client
    ).distinct().all()

    # Query to get count of all candidates by role, location, job type, and client
    role_industry_location_analysis = db.session.query(
        JobPost.role,
        JobPost.location,
        JobPost.job_type,
        JobPost.client,
        func.count(Candidate.id).label('total_count')
    ).join(
        Candidate,
        Candidate.job_id == JobPost.id  # Assuming job_id links Candidate to JobPost
    ).filter(
        JobPost.role.isnot(None)  # Ensure JobPost.role is not None (to filter out non-linked candidates)
    ).group_by(
        JobPost.role,
        JobPost.location,
        JobPost.job_type,
        JobPost.client
    ).all()

    # Query to get the count of on-boarded candidates by role, location, job type, and client
    on_boarded_candidates_analysis = db.session.query(
        JobPost.role,
        JobPost.location,
        JobPost.job_type,
        JobPost.client,
        func.count(Candidate.id).label('on_boarded_count')
    ).join(
        Candidate,
        Candidate.job_id == JobPost.id
    ).filter(
        Candidate.status == 'ON-BOARDED'
    ).group_by(
        JobPost.role,
        JobPost.location,
        JobPost.job_type,
        JobPost.client
    ).all()

    # Combine results
    result = []
    on_boarded_dict = {(item.role, item.location, item.job_type, item.client): item.on_boarded_count for item in on_boarded_candidates_analysis}
    total_dict = {(item.role, item.location, item.job_type, item.client): item.total_count for item in role_industry_location_analysis}

    for role, location, job_type, client in roles:
        key = (role, location, job_type, client)
        total_count = total_dict.get(key, 0)
        on_boarded_count = on_boarded_dict.get(key, 0)
        on_boarded_percentage = (on_boarded_count / total_count) * 100 if total_count > 0 else 0

        result.append({
            'role': role,
            'location': location,
            'job_type': job_type,
            'client': client,
            'total_count': total_count,
            'on_boarded_count': on_boarded_count,
            'on_boarded_percentage': on_boarded_percentage
        })

    return result



def get_conversion_rate(query):
    total_submissions = query.count()
    if total_submissions > 0:
        successful_closures = query.filter(Candidate.status == 'ON-BOARDED').count()
        conversion_rate = successful_closures / total_submissions
    else:
        conversion_rate = 0.0

    return conversion_rate


def get_client_closure_rates(query): # Avadhut
    client_closure_counts = query.filter(Candidate.status.in_(['SELECTED', 'ON-BOARDED'])).group_by(Candidate.client).with_entities(
        Candidate.client,
        func.count().label('count')
    ).all()

    client_closure_rates = [{'client': item.client, 'count': item.count} for item in client_closure_counts]

    highest_closure_client = max(client_closure_rates, key=lambda x: x['count']) if client_closure_rates else None
    lowest_closure_client = min(client_closure_rates, key=lambda x: x['count']) if client_closure_rates else None

    highest_closure_candidates = []
    lowest_closure_candidates = []

    if highest_closure_client:
        highest_closure_candidates = query.filter(
            Candidate.client == highest_closure_client['client'],
            Candidate.status.in_(['SELECTED', 'ON-BOARDED'])
        ).all()
        highest_closure_candidates = [{
            'candidate_name': candidate.name,
            'job_id': candidate.job_id,
            'client': candidate.client,
            'recruiter': candidate.recruiter,
            'date_created': candidate.date_created.strftime('%Y-%m-%d'),
            'time_created': candidate.date_created.strftime('%H:%M:%S'),
            'profile': candidate.profile,
            'last_working_date': candidate.last_working_date.strftime('%Y-%m-%d') if candidate.last_working_date else None,
            'status': candidate.status
        } for candidate in highest_closure_candidates]

    if lowest_closure_client:
        lowest_closure_candidates = query.filter(
            Candidate.client == lowest_closure_client['client'],
            Candidate.status.in_(['SELECTED', 'ON-BOARDED'])
        ).all()
        lowest_closure_candidates = [{
            'candidate_name': candidate.name,
            'job_id': candidate.job_id,
            'client': candidate.client,
            'recruiter': candidate.recruiter,
            'date_created': candidate.date_created.strftime('%Y-%m-%d'),
            'time_created': candidate.date_created.strftime('%H:%M:%S'),
            'profile': candidate.profile,
            'last_working_date': candidate.last_working_date.strftime('%Y-%m-%d') if candidate.last_working_date else None,
            'status': candidate.status
        } for candidate in lowest_closure_candidates]

    return client_closure_rates, highest_closure_client, lowest_closure_client, highest_closure_candidates, lowest_closure_candidates


def generate_bar_graph_data(recruiter_data):
    bar_graph_data = {
        'recruiters': [],
        'selected_candidates_counts': []
    }

    for recruiter, data in recruiter_data.items():
        bar_graph_data['recruiters'].append(recruiter)
        bar_graph_data['selected_candidates_counts'].append(data['selected_candidates_count'])

    return bar_graph_data



@app.route('/time_to_close_position_for_recruiter', methods=['POST'])
def get_time_to_close_analysis():
    data = request.json
    recruiter_usernames = data.get('recruiter_names', [])
    result = {}

    # Fetch the recruiter names from the User table
    recruiters = db.session.query(User.username, User.name).filter(User.username.in_(recruiter_usernames)).all()
    username_to_name_map = {recruiter.username: recruiter.name for recruiter in recruiters}

    for recruiter_username in recruiter_usernames:
        recruiter_name = username_to_name_map.get(recruiter_username, recruiter_username)

        # Query candidates for the recruiter
        candidates = db.session.query(Candidate).filter(
            Candidate.recruiter == recruiter_name,
            Candidate.status.in_([
                'SCREENING', 'SCREENING SELECTED','L1-SCHEDULE', 'L1-FEEDBACK', 'L1-SELECTED', 'L1-CANDIDATE RESCHEDULE','L3-CANDIDATE RESCHEDULE', 
                'L3-CANDIDATE RESCHEDULE','L1-PANEL RESCHEDULE', 'L2-PANEL RESCHEDULE','L3-PANEL RESCHEDULE','L2-SCHEDULE', 'L2-FEEDBACK','L3-FEEDBACK' ,'L2-SELECTED', 'L3-SELECTED' ,'L3-SCHEDULE','HR-ROUND', 'MANAGERIAL ROUND',
                'NEGOTIATION', 'SELECTED','SCREEN REJECTED', 'L1-REJECTED', 'L2-REJECTED', 'L3-REJECTED', 'OFFER-DECLINED', 
                'OFFER-REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO-SHOW','NO SHOW','ON-BOARDED','None','SCREENING'])
        ).all()

        candidates_data = []
        total_screening_candidates = 0
        total_days_to_close = 0
        count_of_onboarded_positions = 0
        total_candidates = len(candidates)
        unsuccessful_closures = 0

        for candidate in candidates:
            days_to_close = None  # Initialize days_to_close to None
            if candidate.status == 'SCREENING':
                total_screening_candidates += 1
            elif candidate.status == 'ON-BOARDED':
                count_of_onboarded_positions += 1

                # Calculate days to close
                if candidate.date_created and candidate.data_updated_date:
                    days_to_close = (candidate.data_updated_date - candidate.date_created).days
                    total_days_to_close += days_to_close

                # Prepare candidate data
                candidate_data = {
                    'candidate_name': candidate.name,
                    'job_id': candidate.job_id,  # Assuming job_id is a regular column
                    'client': candidate.client,
                    'recruiter': recruiter_name,
                    'date_created': candidate.date_created.strftime('%Y-%m-%d') if candidate.date_created else None,
                    'date_updated': candidate.data_updated_date.strftime('%Y-%m-%d') if candidate.data_updated_date else None,
                    'days_to_close': days_to_close,  # This will be None if not ON-BOARDED
                    'profile': candidate.profile,
                    'status': candidate.status
                }
                candidates_data.append(candidate_data)
            elif candidate.status in [
                'SCREEN REJECTED', 'L1-REJECTED', 'L2-REJECTED', 'L3-REJECTED', 'OFFER-DECLINED', 
                'OFFER-REJECTED', 'DUPLICATE', 'HOLD', 'DROP', 'CANDIDATE NO-SHOW','NO SHOW','None','NEGOTIATION']:
                unsuccessful_closures += 1

        # Calculate average days to close
        average_days_to_close = (total_days_to_close / count_of_onboarded_positions) if count_of_onboarded_positions > 0 else 0

        candidates_in_progress = total_candidates - (count_of_onboarded_positions + total_screening_candidates + unsuccessful_closures)
        
        # Calculate percentage of onboarded candidates
        percentage_onboarded = (count_of_onboarded_positions / total_candidates) * 100 if total_candidates > 0 else 0
            
        # Append summary and candidates data for the recruiter
        recruiter_data = {
            'recruiter_name': recruiter_name,
            'candidates': candidates_data,
            'total_days_to_close': total_days_to_close,
            'count_of_screening_candidates': total_screening_candidates,
            'count_of_onboarded_positions': count_of_onboarded_positions,
            'unsuccessful_closures': unsuccessful_closures,
            'candidates_in_progress': candidates_in_progress,
            'average_days_to_close': average_days_to_close,
            'percentage_onboarded': percentage_onboarded,
            'total_candidates_count': total_candidates
        }

        result[recruiter_name] = recruiter_data

    # Sort recruiters by average days to close (ascending order)
    ranked_recruiters = sorted(result.values(), key=lambda x: x['average_days_to_close'])

    # Assign rankings
    rank = 1
    for data in ranked_recruiters:
        if data['count_of_onboarded_positions'] > 0:
            data['ranking'] = rank
            rank += 1
        else:
            data['ranking'] = 0

    return jsonify(ranked_recruiters)



import itertools

@app.route('/generate_excel', methods=['POST'])
def generate_excel():
    data = request.json

    if not data:
        return jsonify({'error': 'No JSON data provided'})

    user_id = data.get('user_id')
    from_date_str = data.get('from_date')
    to_date_str = data.get('to_date')
    recruiter_names = data.get('recruiter_names', [])

    if not recruiter_names:
        return jsonify({'error': 'Please select any Recruiter'})

    try:
        from_date = datetime.strptime(from_date_str, "%d-%m-%Y")
        to_date = datetime.strptime(to_date_str, "%d-%m-%Y")
    except ValueError:
        return jsonify({'error': 'Invalid date format. Please use DD-MM-YYYY format.'})

    session = Session()
    
    # Generate all recruiter-date combinations within the specified range
    all_recruiter_date_combinations = list(itertools.product(recruiter_names, pd.date_range(from_date, to_date, freq='D')))

    # Sort combinations by date (earliest to latest)
    all_recruiter_date_combinations.sort(key=lambda x: x[1])

    # Fetch all candidates within the specified date range
    candidates_query = session.query(Candidate.recruiter, Candidate.date_created, func.count(Candidate.id).label('count')).filter(
        Candidate.recruiter.in_(recruiter_names),
        Candidate.date_created >= from_date,
        Candidate.date_created <= to_date
    ).group_by(Candidate.recruiter, Candidate.date_created).all()

    # Convert candidate data to a dictionary for easy lookup
    candidate_data_dict = {(row.recruiter, row.date_created.strftime("%d-%m-%Y")): row.count for row in candidates_query}

    # Prepare data for the report
    data = []
    for recruiter, date in all_recruiter_date_combinations:
        count = candidate_data_dict.get((recruiter, date.strftime("%d-%m-%Y")), 0)
        data.append({
            "recruiter": recruiter,
            "date_created": date.strftime("%d-%m-%Y"),
            "count": count
        })

    # Convert data to DataFrame
    data_df = pd.DataFrame(data)

    # Creating a pivot table for the report
    pivot_table = data_df.pivot_table(index='recruiter', columns='date_created', values='count', aggfunc='sum',
                                      fill_value=0, margins=True, margins_name='Grand Total')

    # Extract 'Grand Total' column
    grand_total = pivot_table['Grand Total']

    # Drop 'Grand Total' column from the pivot table
    pivot_table.drop(columns=['Grand Total'], inplace=True, errors='ignore')

    # Ensure columns are sorted in chronological order
    pivot_table.columns = pd.to_datetime(pivot_table.columns, format="%d-%m-%Y")
    pivot_table = pivot_table.sort_index(axis=1)

    # Convert columns back to string format for JSON serialization
    pivot_table.columns = pivot_table.columns.strftime("%d-%m-%Y")

    # Add 'Grand Total' column back to the pivot table
    pivot_table['Grand Total'] = grand_total

    # Convert pivot table to JSON for response
    styled_pivot_table_json = pivot_table.to_json()

    return jsonify({
        'user_id': user_id,
        'from_date_str': from_date_str,
        'to_date_str': to_date_str,
        'pivot_table': styled_pivot_table_json
    })


#############  Here coverting the username to name ########################

# import itertools

# @app.route('/generate_excel', methods=['POST'])
# def generate_excel():
#     data = request.json

#     if not data:
#         return jsonify({'error': 'No JSON data provided'})

#     user_id = data.get('user_id')
#     from_date_str = data.get('from_date')
#     to_date_str = data.get('to_date')
#     recruiter_usernames = data.get('recruiter_names', [])

#     if not recruiter_usernames:
#         return jsonify({'error': 'Please select any Recruiter'})

#     try:
#         from_date = datetime.strptime(from_date_str, "%d-%m-%Y")
#         to_date = datetime.strptime(to_date_str, "%d-%m-%Y")
#     except ValueError:
#         return jsonify({'error': 'Invalid date format. Please use DD-MM-YYYY format.'})

#     session = Session()

#     # Fetch the recruiter names based on the recruiter usernames
#     recruiter_names_map = session.query(User.username, User.name).filter(User.username.in_(recruiter_usernames)).all()
#     recruiter_names_dict = {username: name for username, name in recruiter_names_map}

#     # Replace usernames with names in the recruiter_names list
#     recruiter_names = [recruiter_names_dict.get(username, username) for username in recruiter_usernames]

#     # Generate all recruiter-date combinations within the specified range
#     all_recruiter_date_combinations = list(itertools.product(recruiter_names, pd.date_range(from_date, to_date, freq='D')))

#     # Sort combinations by date (earliest to latest)
#     all_recruiter_date_combinations.sort(key=lambda x: x[1])

#     # Fetch all candidates within the specified date range using the recruiter names
#     candidates_query = session.query(Candidate.recruiter, Candidate.date_created, func.count(Candidate.id).label('count')).filter(
#         Candidate.recruiter.in_(recruiter_names),  # Use the list of recruiter names here
#         Candidate.date_created >= from_date,
#         Candidate.date_created <= to_date
#     ).group_by(Candidate.recruiter, Candidate.date_created).all()

#     # Convert candidate data to a dictionary for easy lookup
#     candidate_data_dict = {(row.recruiter, row.date_created.strftime("%d-%m-%Y")): row.count for row in candidates_query}

#     # Prepare data for the report
#     data = []
#     for recruiter, date in all_recruiter_date_combinations:
#         count = candidate_data_dict.get((recruiter, date.strftime("%d-%m-%Y")), 0)
#         data.append({
#             "recruiter": recruiter,
#             "date_created": date.strftime("%d-%m-%Y"),
#             "count": count
#         })

#     # Convert data to DataFrame
#     data_df = pd.DataFrame(data)

#     # Creating a pivot table for the report
#     pivot_table = data_df.pivot_table(index='recruiter', columns='date_created', values='count', aggfunc='sum',
#                                       fill_value=0, margins=True, margins_name='Grand Total')

#     # Extract 'Grand Total' column
#     grand_total = pivot_table['Grand Total']

#     # Drop 'Grand Total' column from the pivot table
#     pivot_table.drop(columns=['Grand Total'], inplace=True, errors='ignore')

#     # Ensure columns are sorted in chronological order
#     pivot_table.columns = pd.to_datetime(pivot_table.columns, format="%d-%m-%Y")
#     pivot_table = pivot_table.sort_index(axis=1)

#     # Convert columns back to string format for JSON serialization
#     pivot_table.columns = pivot_table.columns.strftime("%d-%m-%Y")

#     # Add 'Grand Total' column back to the pivot table
#     pivot_table['Grand Total'] = grand_total

#     # Convert pivot table to JSON for response
#     styled_pivot_table_json = pivot_table.to_json()

#     return jsonify({
#         'user_id': user_id,
#         'from_date_str': from_date_str,
#         'to_date_str': to_date_str,
#         'pivot_table': styled_pivot_table_json
#     })




# def re_send_notification(recruiter_email, job_id):
#     msg = Message('Job Update Notification', sender=config.sender_email, recipients=[recruiter_email])
#     msg.body = f'Hello,\n\nThe job post with ID {job_id} has been updated.\n\nPlease check your dashboard for more details.'
#     mail.send(msg)


def job_removed_send_notification(recruiter_email, new_recruiter_name, job_data, job_id):
    html_body = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                color: #333;
                line-height: 1.6;
                background-color: #f4f4f4;
                margin: 0;
                padding: 0;
            }}
            .container {{
                padding: 20px;
                margin: 20px auto;
                max-width: 600px;
                background-color: #ffffff;
                border: 1px solid #ddd;
                border-radius: 8px;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                text-align: center;
                font-size: 20px;
                border-radius: 8px 8px 0 0;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin-top: 10px;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #4CAF50;
                color: white;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            p {{
                margin: 10px 0;
            }}
            .footer {{
                margin-top: 20px;
                font-size: 12px;
                color: #777;
                text-align: center;
                border-top: 1px solid #ddd;
                padding-top: 10px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                Job Removal Notification
            </div>
            <p>Dear {new_recruiter_name},</p>
            <p>The job post with ID <b>{job_id}</b> has been removed from your login.</p>
            <p>Please find the details below:</p>
            <table>
                <tr>
                    <th style="width: 20%;">Job ID</th>
                    <th style="width: 30%;">Client</th>
                    <th style="width: 30%;">Role/Profile</th>
                    <th style="width: 20%;">Location</th>
                </tr>
                {job_data}
            </table>
            <p>Please check your dashboard for more details.</p>
            <p>Regards,</p>
            <p><b>Makonis Talent Track Pro Team</b></p>
        </div>
    </body>
    </html>
    """

    msg = Message(
        f'Job Removal Notification: Job ID {job_id}',
        sender=config.sender_email,
        recipients=[recruiter_email]
    )
    msg.html = html_body
    mail.send(msg)



def job_updated_send_notification(recruiter_email, new_recruiter_name, job_data, job_id):
    html_body = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                color: #333;
                line-height: 1.6;
                background-color: #f4f4f4;
                margin: 0;
                padding: 0;
            }}
            .container {{
                padding: 20px;
                margin: 20px auto;
                max-width: 600px;
                background-color: #ffffff;
                border: 1px solid #ddd;
                border-radius: 8px;
                box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
            }}
            .header {{
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                text-align: center;
                font-size: 20px;
                border-radius: 8px 8px 0 0;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin-top: 10px;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #4CAF50;
                color: white;
            }}
            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            p {{
                margin: 10px 0;
            }}
            .footer {{
                margin-top: 20px;
                font-size: 12px;
                color: #777;
                text-align: center;
                border-top: 1px solid #ddd;
                padding-top: 10px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                Job Update Notification
            </div>
            <p>Dear {new_recruiter_name},</p>
            <p>The job post with ID <b>{job_id}</b> has been updated.</p>
            <p>Please find the details below:</p>
            <table>
                <tr>
                    <th style="width: 20%;">Job ID</th>
                    <th style="width: 30%;">Client</th>
                    <th style="width: 30%;">Role/Profile</th>
                    <th style="width: 20%;">Location</th>
                </tr>
                {job_data}
            </table>
            <p>Please check your dashboard for more details.</p>
            <p>Regards,</p>
            <p><b>Makonis Talent Track Pro Team</b></p>
        </div>
    </body>
    </html>
    """

    msg = Message(
        f'Job Update Notification: Job ID {job_id}',
        sender=config.sender_email,
        recipients=[recruiter_email]
    )
    msg.html = html_body
    mail.send(msg)


@app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
def edit_job_post(job_post_id):
    try:
        data = request.json
        user_id = data.get('user_id')
        
        # Retrieve the user
        user = User.query.filter_by(id=user_id).first()
        
        if user and user.user_type == 'management':
            # Retrieve the job post to be edited
            job_post = JobPost.query.get(job_post_id)
            
            if job_post:
                old_recruiter_usernames = set(job_post.recruiter.split(', ')) if isinstance(job_post.recruiter, str) else set()

                # Check if any field except 'recruiter' is updated
                fields_updated = set(data.keys()) - {'recruiter'}
                if fields_updated:
                    # Update job post fields
                    job_post.client = data.get('client', job_post.client)
                    job_post.experience_min = data.get('experience_min', job_post.experience_min)
                    job_post.experience_max = data.get('experience_max', job_post.experience_max)
                    job_post.budget_min = data.get('budget_min', job_post.budget_min)
                    job_post.budget_max = data.get('budget_max', job_post.budget_max)
                    job_post.location = data.get('location', job_post.location)
                    job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
                    job_post.notice_period = data.get('notice_period', job_post.notice_period)
                    job_post.role = data.get('role', job_post.role)
                    job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
                    job_post.mode = data.get('mode', job_post.mode)
                    job_post.job_status = data.get('job_status', job_post.job_status)
                    job_post.skills = data.get('skills', job_post.skills)
                    job_post.no_of_positions = data.get('no_of_positions', job_post.no_of_positions)
                    
                    recruiters = data.get('recruiter', job_post.recruiter)
                    recruiters = set(recruiters if isinstance(recruiters, list) else [recruiters])
                    job_post.recruiter = ', '.join(list(recruiters))
                    
                    job_type = data.get('Job_Type')
                    job_post.job_type = job_type
                    if job_type == 'Contract':
                        job_post.contract_in_months = data.get('Job_Type_details')
                    
                    # Handle jd_pdf field
                    jd_pdf = data.get('jd_pdf')
                    if jd_pdf is not None:
                        jd_binary = base64.b64decode(jd_pdf)
                        job_post.jd_pdf = jd_binary
                        job_post.jd_pdf_present = True

                    # Update job post in the database
                    db.session.commit()
                    
                    # Create notification records for each recruiter
                    if job_post.recruiter:
                        for recruiter in recruiters:
                            notification = Notification.query.filter_by(job_post_id=job_post_id, recruiter_name=recruiter).first()
                            if notification:
                                notification.num_notification += 1
                            else:
                                new_notification = Notification(job_post_id=job_post_id, recruiter_name=recruiter)
                                db.session.add(new_notification)
                                new_notification.num_notification = 1
                    
                    # Update candidate details
                    candidates = Candidate.query.filter_by(job_id=job_post_id).all()
                    for candidate in candidates:
                        candidate.client = job_post.client
                        candidate.profile = job_post.role

                    db.session.commit()
                    
                    return jsonify({
                        'status': 'success',
                        "message": "Job post details updated successfully",
                        'job_post_id': job_post_id,
                        'old_recruiter_usernames': list(old_recruiter_usernames),
                        'new_recruiter_usernames': list(recruiters)
                    }), 200
                else:
                    return jsonify({'status': 'success', "message": "No fields updated other than recruiter"})
            else:
                return jsonify({'status': 'error', "message": "Job post not found"})
        else:
            return jsonify({'status': 'error', "message": "Unauthorized"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/send_edit_notifications', methods=['POST'])
def send_edit_notifications():
    try:
        data = request.json
        job_post_id = data.get('job_post_id')
        old_recruiter_usernames = set(data.get('old_recruiter_usernames', []))
        new_recruiter_usernames = set(data.get('new_recruiter_usernames', []))

        if not job_post_id:
            return jsonify({'status': 'error', 'message': 'job_post_id is required'}), 400

        job_post = JobPost.query.filter_by(id=job_post_id).first()
        if not job_post:
            return jsonify({'status': 'error', 'message': 'Job post not found'}), 404

        job_data = f"<tr><td>{job_post.id}</td><td>{job_post.client}</td><td>{job_post.role}</td><td>{job_post.location}</td></tr>"

        unchanged_recruiters = set(old_recruiter_usernames) & set(new_recruiter_usernames)
        removed_recruiters = set(old_recruiter_usernames) - set(new_recruiter_usernames)
        added_recruiters = set(new_recruiter_usernames) - set(old_recruiter_usernames)

        all_recruiters = User.query.filter(
            User.name.in_(set(new_recruiter_usernames).union(set(old_recruiter_usernames))),
            User.user_type == 'recruiter',
            User.is_active == True,
            User.is_verified == True
        ).all()

        all_recruiter_emails = {recruiter.name: recruiter.email for recruiter in all_recruiters}

        for recruiter_name in unchanged_recruiters:
            email = all_recruiter_emails.get(recruiter_name)
            if email:
                job_updated_send_notification(recruiter_email=email, new_recruiter_name=recruiter_name, job_data=job_data, job_id=job_post.id)

        for recruiter_name in removed_recruiters:
            email = all_recruiter_emails.get(recruiter_name)
            if email:
                job_removed_send_notification(recruiter_email=email, new_recruiter_name=recruiter_name, job_data=job_data, job_id=job_post.id)

        for recruiter_name in added_recruiters:
            email = all_recruiter_emails.get(recruiter_name)
            if email:
                post_job_send_notification(recruiter_email=email, new_recruiter_name=recruiter_name, job_data=job_data)
        
        return jsonify({'status': 'success', 'message': 'Notifications sent successfully'}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500



# def job_updated_send_notification(recruiter_email, new_recruiter_name, job_data, job_id):
#     html_body = f"""
#     <html>
#     <head>
#         <style>
#             body {{
#                 font-family: Arial, sans-serif;
#                 color: #333;
#                 line-height: 1.6;
#                 background-color: #f4f4f4;
#                 margin: 0;
#                 padding: 0;
#             }}
#             .container {{
#                 padding: 20px;
#                 margin: 20px auto;
#                 max-width: 600px;
#                 background-color: #ffffff;
#                 border: 1px solid #ddd;
#                 border-radius: 8px;
#                 box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
#             }}
#             .header {{
#                 background-color: #4CAF50;
#                 color: white;
#                 padding: 10px;
#                 text-align: center;
#                 font-size: 20px;
#                 border-radius: 8px 8px 0 0;
#             }}
#             table {{
#                 border-collapse: collapse;
#                 width: 100%;
#                 margin-top: 10px;
#             }}
#             th, td {{
#                 border: 1px solid #ddd;
#                 padding: 8px;
#                 text-align: left;
#             }}
#             th {{
#                 background-color: #4CAF50;
#                 color: white;
#             }}
#             tr:nth-child(even) {{
#                 background-color: #f9f9f9;
#             }}
#             p {{
#                 margin: 10px 0;
#             }}
#             .footer {{
#                 margin-top: 20px;
#                 font-size: 12px;
#                 color: #777;
#                 text-align: center;
#                 border-top: 1px solid #ddd;
#                 padding-top: 10px;
#             }}
#         </style>
#     </head>
#     <body>
#         <div class="container">
#             <div class="header">
#                 Job Update Notification
#             </div>
#             <p>Dear {new_recruiter_name},</p>
#             <p>The job post with ID <b>{job_id}</b> has been updated.</p>
#             <p>Please find the details below:</p>
#             <table>
#                 <tr>
#                     <th style="width: 20%;">Job ID</th>
#                     <th style="width: 30%;">Client</th>
#                     <th style="width: 30%;">Role/Profile</th>
#                     <th style="width: 20%;">Location</th>
#                 </tr>
#                 {job_data}
#             </table>
#             <p>Please check your dashboard for more details.</p>
#             <p>Regards,</p>
#             <p><b>Makonis Talent Track Pro Team</b></p>
#         </div>
#     </body>
#     </html>
#     """

#     msg = Message(
#         f'Job Update Notification: Job ID {job_id}',
#         sender=config.sender_email,
#         recipients=[recruiter_email]
#     )
#     msg.html = html_body
#     mail.send(msg)


# @app.route('/edit_job_post/<int:job_post_id>', methods=['POST'])
# def edit_job_post(job_post_id):
#     try:
#         data = request.json
#         user_id = data.get('user_id')
        
#         # Retrieve the user
#         user = User.query.filter_by(id=user_id).first()
        
#         if user and user.user_type == 'management':
#             # Retrieve the job post to be edited
#             job_post = JobPost.query.get(job_post_id)
            
#             if job_post:
#                 old_recruiter_usernames = set(job_post.recruiter.split(', ')) if isinstance(job_post.recruiter, str) else set()

#                 # Check if any field except 'recruiter' is updated
#                 fields_updated = set(data.keys()) - {'recruiter'}
#                 if fields_updated:
#                     # Update job post fields
#                     job_post.client = data.get('client', job_post.client)
#                     job_post.experience_min = data.get('experience_min', job_post.experience_min)
#                     job_post.experience_max = data.get('experience_max', job_post.experience_max)
#                     job_post.budget_min = data.get('budget_min', job_post.budget_min)
#                     job_post.budget_max = data.get('budget_max', job_post.budget_max)
#                     job_post.location = data.get('location', job_post.location)
#                     job_post.shift_timings = data.get('shift_timings', job_post.shift_timings)
#                     job_post.notice_period = data.get('notice_period', job_post.notice_period)
#                     job_post.role = data.get('role', job_post.role)
#                     job_post.detailed_jd = data.get('detailed_jd', job_post.detailed_jd)
#                     job_post.mode = data.get('mode', job_post.mode)
#                     job_post.job_status = data.get('job_status', job_post.job_status)
#                     job_post.skills = data.get('skills', job_post.skills)
#                     job_post.no_of_positions = data.get('no_of_positions', job_post.no_of_positions)
                    
#                     recruiters = data.get('recruiter', job_post.recruiter)
#                     recruiters = set(recruiters if isinstance(recruiters, list) else [recruiters])
#                     job_post.recruiter = ', '.join(list(recruiters))
                    
#                     job_type = data.get('Job_Type')
#                     job_post.job_type = job_type
#                     if job_type == 'Contract':
#                         job_post.contract_in_months = data.get('Job_Type_details')
                    
#                     # Handle jd_pdf field
#                     jd_pdf = data.get('jd_pdf')
#                     if jd_pdf is not None:
#                         jd_binary = base64.b64decode(jd_pdf)
#                         job_post.jd_pdf = jd_binary
#                         job_post.jd_pdf_present = True

#                     # Update job post in the database
#                     db.session.commit()
                    
#                     # Create notification records for each recruiter
#                     if job_post.recruiter:
#                         for recruiter in recruiters:
#                             notification = Notification.query.filter_by(job_post_id=job_post_id, recruiter_name=recruiter).first()
#                             if notification:
#                                 notification.num_notification += 1
#                             else:
#                                 new_notification = Notification(job_post_id=job_post_id, recruiter_name=recruiter)
#                                 db.session.add(new_notification)
#                                 new_notification.num_notification = 1
                    
#                     # Update candidate details
#                     candidates = Candidate.query.filter_by(job_id=job_post_id).all()
#                     for candidate in candidates:
#                         candidate.client = job_post.client
#                         candidate.profile = job_post.role

#                     db.session.commit()
                    
#                     return jsonify({
#                         'status': 'success',
#                         "message": "Job post details updated successfully",
#                         'job_post_id': job_post_id,
#                         'old_recruiter_usernames': list(old_recruiter_usernames),
#                         'new_recruiter_usernames': list(recruiters)
#                     }), 200
#                 else:
#                     return jsonify({'status': 'success', "message": "No fields updated other than recruiter"})
#             else:
#                 return jsonify({'status': 'error', "message": "Job post not found"})
#         else:
#             return jsonify({'status': 'error', "message": "Unauthorized"})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500



@app.route('/jobs_notification/<int:user_id>', methods=['GET'])
def get_jobs_notification(user_id):
    # Retrieve the user
    user = User.query.filter_by(id=user_id).first()
    
    # Check if the user exists and has the right permissions
    if user and user.user_type == 'recruiter':
        recruiter_name = user.name
        
        # Retrieve the notifications for the recruiter where num_notification >= 1
        notifications = Notification.query.filter_by(recruiter_name=recruiter_name).filter(Notification.num_notification >= 1).all()
        
        # Format the notifications as a list of dictionaries
        notifications_list = [
            {
                # 'id': notification.id,
                'job_post_id': notification.job_post_id,
                'recruiter_name': notification.recruiter_name,
                'notification_status': notification.notification_status,
                'num_notification': notification.num_notification
            } for notification in notifications
        ]
        
        return jsonify(notifications_list), 200
    else:
        return jsonify({'error': 'User not found or does not have the right permissions'}), 404


@app.route('/checked_jobs_notification/<int:user_id>', methods=['POST'])
def checked_jobs_notification(user_id):
    data = request.json
    checked_notification_status = data.get('checked_notification_status')

    user = User.query.filter_by(id=user_id).first()
    
    # Check if the user exists and has the right permissions
    if user and user.user_type == 'recruiter':
        recruiter_name = user.name
        
        # Retrieve the notifications for the recruiter
        notifications = Notification.query.filter_by(recruiter_name=recruiter_name).all()
        
        if checked_notification_status:
            # Update the num_notification to 0 for each notification
            for notification in notifications:
                notification.notification_status = checked_notification_status
                notification.num_notification = 0
                db.session.commit()
                
            # Delete notifications where num_notification is 0
            notifications_to_delete = Notification.query.filter_by(recruiter_name=recruiter_name, num_notification=0).all()
            for notification in notifications_to_delete:
                db.session.delete(notification)
                db.session.commit()
        
        # Format the notifications as a list of dictionaries
        notifications_list = [
            {
                'id': notification.id,
                'job_post_id': notification.job_post_id,
                'recruiter_name': notification.recruiter_name,
                'notification_status': notification.notification_status,
                'num_notification': notification.num_notification
            } for notification in notifications
        ]
        
        return jsonify(notifications_list), 200
    else:
        return jsonify({'error': 'User not found or does not have the right permissions'}), 404



    
@app.route('/get_candidate_data')
def get_candidate_data():
    candidates = Candidate.query.all()
    candidate_data = []
    for candidate in candidates:
        candidate_data.append({
            'id': candidate.id,
            'name': candidate.name,
            'email': candidate.email,
            'client': candidate.client,
            'current_company':candidate.current_company,
            'position': candidate.position,
            'profile': candidate.profile,
            'current_job_location':candidate.current_job_location,
            'preferred_job_location':candidate.preferred_job_location,
            'skills':candidate.skills,
            'status':candidate.status,
        })
    return jsonify(candidate_data)


@app.route('/send_email', methods=['POST'])
def send_email():
    recipient_email = request.form.get('recipient_email')

    if not recipient_email:
        flash('Recipient email is required.', 'error')
        return redirect(url_for('careers'))

    # Create a link to the page you want to send
    page_link = 'http://127.0.0.1:5001/careers'  # Replace with the actual link

    # Create the email content with a hyperlink
    email_content = f"Click the link below to view active job posts: <a href='{page_link}'>{page_link}</a>"

    # Create an email message
    message = Message('Active Job Posts', sender=config.sender_email, recipients=[recipient_email])
    message.html = email_content

    # Send the email
    mail.send(message)

    flash('Email sent successfully!', 'success')
    return redirect(url_for('careers'))

#new
@app.route('/careers', methods=['GET'])
def careers():
    user_type = session.get('user_type', None)
    is_logged_in = 'user_id' in session
    candidate_message = request.args.get('candidate_message')
    print(candidate_message)

    # Query the database to retrieve active job posts and sort them by date_created in descending order
    active_jobs = JobPost.query.filter_by(job_status='Active').order_by(JobPost.date_created.desc()).all()

    return render_template('careers.html', jobs=active_jobs, user_type=user_type, is_logged_in=is_logged_in,candidate_message=candidate_message)

#new
@app.route('/apply_careers', methods=['GET', 'POST'])
def apply_careers():
    user_id = session.get('user_id')
    if not user_id:
        # User is not authenticated, you can redirect them to a login page or take appropriate action
        return redirect(url_for('career_login'))
    user = Career_user.query.get(user_id)
    if request.method == 'GET':
        job_id = request.args.get('job_id')
        client = request.args.get('client')
        profile = request.args.get('profile')
        name = user.name
        email = user.email

        if job_id:
            matching_job_post = JobPost.query.filter(and_(JobPost.id == job_id, JobPost.job_status == 'Hold')).first()
            if matching_job_post:
                return render_template('job_on_hold.html')
        
        job_post = JobPost.query.get(job_id)
        experience_min = job_post.experience_min

        job_ids = db.session.query(JobPost.id).filter(JobPost.client == client, JobPost.job_status == 'Active').all()
        job_roles = db.session.query(JobPost.role).filter(JobPost.client == client).all()

        ids = [job_id[0] for job_id in job_ids]
        roles = [job_role[0] for job_role in job_roles]

        candidate_data = None
        if 'candidate_data' in request.args:
            candidate_data = ast.literal_eval(request.args['candidate_data'])

        return render_template('apply_careers.html', candidate_data=candidate_data, job_id=job_id,
                               client=client, profile=profile, ids=ids, roles=roles,
                               name=name, email=email,experience_min=experience_min)

    if request.method == 'POST':
        try:
            job_id = request.form['job_id']
            name = request.form['name']
            mobile = request.form['mobile']
            email = request.form['email']
            client = request.form['client']
            profile = request.form['profile']
            skills = request.form['skills']

            # Ensure client and job_id are integers
            job_id = int(job_id)

            # Check if the job post is active
            matching_job_post = JobPost.query.filter(and_(JobPost.id == job_id, JobPost.job_status == 'Active')).first()
            if not matching_job_post:
                return render_template('job_on_hold.html')

            # Handle other form fields...
            current_company = request.form['current_company']
            position = request.form['position']
            current_job_location = request.form['current_job_location']
            preferred_job_location = request.form['preferred_job_location']
            qualifications = request.form['qualifications']
            experience = request.form['experience']
            exp_months = request.form['exp_months']
            experience = experience + '.' + exp_months
            relevant_experience = request.form['relevant_experience']
            relevant_exp_months = request.form['relevant_exp_months']
            relevant_experience = relevant_experience + '.' + relevant_exp_months
            currency_type_current = request.form['currency_type_current']
            currency_type_except = request.form['currency_type_except']
            current_ctc = currency_type_current + " " + request.form['current_ctc']
            expected_ctc = currency_type_except + " " + request.form['expected_ctc']
            linkedin = request.form['linkedin']

            # Handle file upload
            filename = None
            resume_binary = None
            if 'resume' in request.files:
                resume_file = request.files['resume']
                if resume_file and allowed_file(resume_file.filename):
                    # Convert the resume file to binary data
                    resume_binary = resume_file.read()
                    filename = secure_filename(resume_file.filename)
                else:
                    return render_template('apply_careers.html', error_message='Invalid file extension')

            notice_period = request.form['notice_period']
            last_working_date = None
            buyout = False
            period_of_notice = None

            if notice_period == 'yes':
                last_working_date = request.form['last_working_date']
                buyout = 'buyout' in request.form
            elif notice_period == 'no':
                period_of_notice = request.form['months']
                buyout = 'buyout' in request.form
            elif notice_period == 'completed':
                last_working_date = request.form['last_working_date']

            holding_offer = request.form['holding_offer']

            if holding_offer == 'yes':
                total = request.form['total']
                if total == '':
                    total = 0
                else:
                    total = int(request.form['total'])
                package_in_lpa = request.form['package_in_lpa']
                if package_in_lpa == '':
                    package_in_lpa = 0
                else:
                    package_in_lpa = float(request.form['package_in_lpa'])
            else:
                total = None
                package_in_lpa = None

            reason_for_job_change = request.form.get('reason_for_job_change')
            remarks = request.form.get('remarks')

            reference = request.form['reference']
            reference_name = None
            reference_position = None
            reference_information = None

            if reference == 'yes':
                reference_name = request.form['reference_name']
                reference_position = request.form['reference_position']
                reference_information = request.form['reference_information']
            elif reference == 'no':
                reference_name = None
                reference_position = None
                reference_information = None

            existing_candidate = Candidate.query.filter(
                and_(Candidate.profile == profile, Candidate.client == client, Candidate.email == email,
                     Candidate.mobile == mobile)).first()
            if existing_candidate:
                return render_template('candidate_exists.html',
                                       error_message='Candidate with the same profile and client already exists')

            # Create a new Candidate object
            new_candidate = Candidate(
                job_id=job_id,
                name=name,
                mobile=mobile,
                email=email,
                client=client,
                current_company=current_company,
                position=position,
                profile=profile,
                resume=resume_binary,
                current_job_location=current_job_location,
                preferred_job_location=preferred_job_location,
                qualifications=qualifications,
                experience=experience,
                relevant_experience=relevant_experience,
                current_ctc=current_ctc,
                expected_ctc=expected_ctc,
                notice_period=notice_period,
                last_working_date=last_working_date if notice_period == 'yes' or notice_period == 'completed' else None,
                buyout=buyout,
                holding_offer=holding_offer,
                total=total,
                package_in_lpa=package_in_lpa,
                linkedin_url=linkedin,
                reason_for_job_change=reason_for_job_change,
                status='None',
                remarks=remarks,
                skills=skills,
                period_of_notice=period_of_notice,
                reference=reference,
                reference_name=reference_name,
                reference_position=reference_position,
                reference_information=reference_information
            )

            new_candidate.date_created = date.today()
            new_candidate.time_created = datetime.now().time()

            # Commit the new candidate to the database
            db.session.add(new_candidate)
            db.session.commit()

            try:
                msg = Message('Successful Submission of Your Job Application', sender=config.sender_email, recipients=[email])
                msg.body = f"Dear { name },\n Congratulations! Your job application has been successfully submitted for the position at {client} for the role of {profile}. We appreciate your interest in joining our team.\n\n  Our dedicated recruiter will review your application, and you can expect to hear from us within the next 24 hours.\n\nBest wishes for your application process!\n\n Regards, \n\nTeam\nMakonis Talent Track Pro\nrecruiterpro@makonissoft.com\n"
                mail.send(msg)
            except Exception as e:
                # Handle email sending errors, log the error
                return render_template('error.html', error_message=f"Error sending thank-you email: {str(e)}")

            return redirect(url_for('careers', candidate_message='Candidate Added Successfully'))

        except Exception as e:
            # Handle any exceptions here (e.g., log the error, return an error page)
            return render_template('error.html', error_message=str(e))

    return redirect(url_for('careers'))


#new
# User Login
@app.route('/career_login', methods=['GET', 'POST'])
def career_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = Career_user.query.filter_by(username=username, password=password).first()

        if user:
            # Store the user's session or token
            session['user_id'] = user.id
            return redirect(url_for('careers'))

    return render_template('career_login.html')

#new
@app.route('/career_logout')
def career_logout():
    # Clear the user's session
    session.pop('user_id', None)
    return redirect(url_for('careers'))

#new
@app.route('/career_register', methods=['GET', 'POST'])
def career_register():
    if request.method == 'POST':
        username = request.form['username']
        name = request.form['name']
        email = request.form['email']
        password = request.form['password']

        # Create a new user and add it to the database
        new_user = Career_user(username=username, name=name, email=email, password=password)
        db.session.add(new_user)
        db.session.commit()

        return redirect(url_for('career_login'))

    return render_template('career_registration.html')

#new
@app.route('/career_dashboard')
def career_dashboard():
    edit_candidate_message = request.args.get('edit_candidate_message')
    if 'user_id' in session and 'user_type' in session:
        page_no = request.args.get('page_no')
        candidate_message = request.args.get('candidate_message')
        signup_message = request.args.get('signup_message')
        job_message = request.args.get('job_message')
        update_candidate_message = request.args.get('update_candidate_message')
        delete_message = request.args.get("delete_message")

        user_id = session['user_id']
        user_type = session['user_type']
        user_name = session['user_name']

        if user_type == 'management':
            users = User.query.all()
            candidates = Candidate.query.filter((Candidate.reference.is_not(None))).all()
            candidates = sorted(candidates, key=lambda candidate: candidate.id)
            JobsPosted = JobPost.query.all()
            # data = json.dumps(candidates, sort_keys=False)
            return render_template('career_dashboard.html', users=users, user_type=user_type, user_name=user_name,
                                   candidates=candidates, update_candidate_message=update_candidate_message,
                                   candidate_message=candidate_message, delete_message=delete_message,
                                   JobsPosted=JobsPosted, signup_message=signup_message, job_message=job_message,
                                   page_no=page_no, edit_candidate_message=edit_candidate_message)
        elif user_type == 'recruiter':
            recruiter = User.query.filter_by(id=user_id, user_type='recruiter').first()
            recruiter_name = User.query.get(user_id).name
            if recruiter:
                candidates = Candidate.query.filter(and_(Candidate.recruiter == recruiter.name,
                                                         Candidate.reference.is_not(None))).all()
                candidates = sorted(candidates, key=lambda candidate: candidate.id)
                career_count_notification_no = Career_notification.query.filter(Career_notification.notification_status == 'false',
                                                                  Career_notification.recruiter_name == user_name).count()
                career_notifications = Career_notification.query.filter(
                    Career_notification.recruiter_name.contains(recruiter_name)).all()

                for career_notification in career_notifications:
                    if career_notification.notification_status == False:
                        career_notification.notification_status = True
                        db.session.commit()
                return render_template('career_dashboard.html', user=recruiter, user_type=user_type, user_name=user_name,
                                       candidates=candidates, candidate_message=candidate_message,
                                       update_candidate_message=update_candidate_message,
                                       career_count_notification_no=career_count_notification_no,
                                       edit_candidate_message=edit_candidate_message, page_no=page_no)
        else:
            user = User.query.filter_by(id=user_id).first()
            if user:
                candidates = Candidate.query.filter_by(recruiter=user.name).all()  # Filter candidates by user's name
                return render_template('career_dashboard.html', user=user, user_type=user_type, candidates=candidates)

    return redirect(url_for('index'))

#new
@app.route('/website_candidate_assign', methods=['GET', 'POST'])
def website_candidate_assign():
    assignment_message = request.args.get('assignment_message')
    if 'user_id' in session and 'user_type' in session and session['user_type'] == 'management':
        user_name = session['user_name']
        recruiters = User.query.filter_by(user_type='recruiter').all()

        if request.method == 'POST':
            assign_recruiter_id = request.form.get('assign_recruiter_id')
            selected_candidate_ids = request.form.getlist('selected_candidate_ids')

            if assign_recruiter_id and selected_candidate_ids:
                assigned_recruiter = User.query.get(assign_recruiter_id)
                if assigned_recruiter:
                    # Fetch selected candidates by their IDs
                    candidates = Candidate.query.filter(
                        Candidate.id.in_(selected_candidate_ids),
                        Candidate.recruiter.is_(None),
                        Candidate.management.is_(None)
                    ).all()

                    for candidate in candidates:
                        # Assign the selected recruiter to the candidate
                        candidate.recruiter = assigned_recruiter.name
                        # Send an email to the assigned recruiter
                        send_career_email(assigned_recruiter.email, 'Alert! New Candidate Assignment ',
                                          f'Dear {assigned_recruiter.name}\n\n,A new candidate application has been assigned to you. Please access your dashboard to view the details.\n\nCandidate Name: {candidate.name}\n\nClient: {candidate.client}\n\nRole: {candidate.profile}\n\nAssigned by Manager: {user_name}\n\nFeel free to reach out if you have any questions during the recruitment process.\n\nRegards,\n\nTeam\nMakonis Talent Track Pro\nrecruiterpro@makonissoft.com')

                    db.session.commit()

                    # Create notifications for the assigned recruiter
                    notifications = []
                    for candidate in candidates:
                        notification = Career_notification(
                            recruiter_name=assigned_recruiter.name,
                            notification_status=False  # You may set this to True for unread notifications
                        )
                        notifications.append(notification)

                    db.session.add_all(notifications)
                    db.session.commit()

                    return redirect(
                        url_for('website_candidate_assign', assignment_message='Candidates Assigned Successfully'))

        candidates = Candidate.query.filter(
            Candidate.recruiter.is_(None),
            Candidate.management.is_(None)
        ).all()

        candidate_count = Candidate.query.filter(
            Candidate.recruiter.is_(None),
            Candidate.management.is_(None)
        ).count()

        return render_template(
            'website_candidate_assign.html',
            recruiters=recruiters,
            candidates=candidates,
            assignment_message=assignment_message,
            user_name=user_name,
            candidate_count=candidate_count
        )

    return redirect(url_for('index'))

#new
def send_career_email(to, subject, message):
    msg = Message(subject, sender=config.sender_email, recipients=[to])
    msg.body = message
    mail.send(msg)

####################################################################################################################################

import base64
import io
import re
from flask import Flask, request, jsonify
import fitz  # PyMuPDF
from docx import Document

ALLOWED_EXTENSIONS = {'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text(file):
    """
    Extract text from PDF or DOCX files.
    
    Parameters:
        file (BytesIO): File-like object.
    
    Returns:
        str: Extracted text.
    """
    try:
        file.seek(0)
        header = file.read(4)
        file.seek(0)
        if header.startswith(b'%PDF'):
            return extract_text_from_pdf(file)
        elif header.startswith(b'PK\x03\x04'):
            return extract_text_from_docx(file)
        else:
            return ""  # Unsupported file format
    except Exception as e:
        print(f"Error determining file type: {e}")
        return ""

def extract_text_from_pdf(file):
    """
    Extract text from a PDF file.
    
    Parameters:
        file (BytesIO): PDF file-like object.
    
    Returns:
        str: Extracted text.
    """
    text = ""
    try:
        with fitz.open(stream=file, filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
    return text

def extract_text_from_docx(file):
    """
    Extract text from a DOCX file.
   
    Parameters:
        file (BytesIO): DOCX file-like object.
   
    Returns:
        str: Extracted text.
    """
    text = ""
    try:
        text = docx2txt.process(file)
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
    return text


# def extract_text_from_docx(file):
#     """
#     Extract text from a DOCX file.
    
#     Parameters:
#         file (BytesIO): DOCX file-like object.
    
#     Returns:
#         str: Extracted text.
#     """
#     text = ""
#     try:
#         doc = Document(file)
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + '\n'
#     except Exception as e:
#         print(f"Error extracting text from DOCX: {e}")
#     return text

def extract_skills_from_resume(text, skills_list):
    found_skills = [skill for skill in skills_list if skill.lower() in text.lower()]
    return found_skills

def extract_email(text):
    email_regex = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    email_matches = re.findall(email_regex, text)
    return email_matches[-1].rstrip('.,') if email_matches else "No email found"
    
def extract_phone_number(text):
    phone_regex = r'\b\d{10}\b'
    phone_matches = re.findall(phone_regex, text)
    return phone_matches[-1] if phone_matches else "No phone number found"

# def extract_phone_number(text):
#     phone_regex = r'\+?\d[\d -]{8,12}\d'
#     phone_matches = re.findall(phone_regex, text)
#     return phone_matches[-1] if phone_matches else "No phone number found"

def extract_name(text):
    """
    Extract the name from the first few lines of the resume text.
    
    Parameters:
        text (str): Resume text.
    
    Returns:
        str: Extracted name.
    """
    lines = text.split('\n')
    name_words = []  # List to store the words of the name
    
    # Regular expressions to identify lines that are likely contact details
    phone_pattern = re.compile(r'\b(\+?\d[\d\-\.\s]+)?\d{10}\b')
    email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
    
    for line in lines[:5]:  # Look at the first five lines where the name is likely to appear
        # Skip lines that are likely to be contact details
        if phone_pattern.search(line) or email_pattern.search(line):
            continue
        
        # Remove common salutations and titles
        cleaned_line = re.sub(r'\b(Mr\.|Mrs\.|Ms\.|Miss|Dr\.|Sir|Madam)\b', '', line, flags=re.IGNORECASE).strip()
        
        # Extract names with up to three words
        words = cleaned_line.split()
        name_words.extend(words)  # Add words from the current line to the list
        
        if len(name_words) <= 2:
            continue  # Continue accumulating words if we have less than or equal to three words
        else:
            # Stop accumulating if we exceed three words and return the concatenated name
            return ' '.join(word.capitalize() for word in name_words[:3]).rstrip('.,')
    
    # Return the concatenated name if found within the first five lines
    if name_words:
        return ' '.join(word.capitalize() for word in name_words[:3]).rstrip('.,')
    
    return "No name found"



@app.route('/parse_resume', methods=['POST'])
def parse_resume():
    if 'resume' not in request.json:
        return jsonify({'status':'error',"message": "No resume data provided"})
    
    data = request.json
    resume_data = data['resume']
    
    try:
        decoded_resume = base64.b64decode(resume_data)
    except Exception as e:
        return jsonify({'status':'error',"message": "Invalid resume data"})
    
    resume_file = io.BytesIO(decoded_resume)
    resume_text = extract_text(resume_file)
    
    if not resume_text:
        return jsonify({'status':'error',"message": "No text found in the resume data"})

    it_skills = [ 
        'Data Analysis', 'Machine Learning', 'Communication', 'Project Management',
        'Deep Learning', 'SQL', 'Tableau', 'C++', 'C', 'Front End Development', 'JAVA', 
        'Java Full Stack', 'React JS', 'Node JS','Programming (Python, Java, C++)',
        'Data Analysis and Visualization','Artificial Intelligence','Programming',
        'Database Management (SQL)','Web Development (HTML, CSS, JavaScript)',
        'Machine Learning and Artificial Intelligence','Network Administration',
        'Software Development and Testing','Embedded Systems','CAD and 3D Modeling',
        'HTML5', 'CSS3', 'Jquery', 'Bootstrap', 'XML', 'JSON', 'ABAP', 'SAPUI5',
        'Agile Methodology', 'Frontend Development', 'Jira', 'Odata', 'BTP', 'Fiori Launchpad', 
        'Python', 'JavaScript', 'HTML', 'CSS','React', 'Node.js', 'Django', 'Git', 'AWS',
        'Linux','DevOps','Linear Regression','Logistic Regression','Decision Tree',
        'SVM (Support Vector Machine)','Ensembles','Random Forest','Clustering',
        'PCA (Principal Component Analysis)','K-means','Recommendation System',
        'Market Basket Analysis','CNN','RNN','LSTM','Natural Language Processing',
        'NLTK','LGBM','XGBoost','Transformers','Siamese network','BTYD (Buy Till You Die)',
        'ML Ops Tools: Azure Synapse','Azure ML','Azure Databricks','ML flow','Airflow',
        'Kubernetes','Dockers','Data Streaming – Kafka','Flask','LT Spice','Wireshark',
        'Ansys Lumerical','Zemax OpticStudio','Xilinx Vivado','Google Collab','MATLAB'
    ]
    
    non_it_skills = [
        'Communication Skills', 'Teamwork', 'Problem Solving', 'Time Management', 'Leadership',
        'Creativity', 'Adaptability', 'Critical Thinking', 'Analytical Skills', 'Attention to Detail',
        'Customer Service', 'Interpersonal Skills', 'Negotiation Skills', 'Project Management', 
        'Presentation Skills', 'Research Skills', 'Organizational Skills', 'Multitasking',
        'Decision Making', 'Emotional Intelligence', 'Conflict Resolution', 'Networking', 
        'Strategic Planning', 'Public Speaking', 'Writing Skills', 'Sales Skills', 'Marketing', 
        'Finance', 'Human Resources', 'Training and Development', 'Event Planning', 'Language Proficiency',
        'Problem-Solving', 'Sales', 'Marketing', 'Financial Analysis', 'Customer Relationship Management (CRM)', 
        'Quality Management', 'Supply Chain Management', 'Logistics', 'Health and Safety', 'Public Relations', 
        'Social Media Management', 'Content Creation', 'Graphic Design', 'Video Editing', 'Photography', 
        'Data Entry', 'Administrative Support', 'Customer Support'
    ]

    extracted_it_skills = extract_skills_from_resume(resume_text, it_skills)
    extracted_nonit_skills = extract_skills_from_resume(resume_text, non_it_skills)
    non_it_skills_final = list(set(extracted_nonit_skills) - set(extracted_it_skills))

    skills_it = ", ".join(extracted_it_skills) if extracted_it_skills else "No skills found"
    skills_non_it = ", ".join(non_it_skills_final) if non_it_skills_final else "No skills found"

    email_text = extract_email(resume_text)
    phone_text = extract_phone_number(resume_text)
    name_text = extract_name(resume_text)

    return jsonify({
        'status':'success',
        'message':'resume parsed successfully',
        "name": name_text,
        "mail": email_text,
        "phone": phone_text,
        "skill1": skills_it,
        "skill2": skills_non_it
    })

if __name__ == '__main__':
    app.run(debug=True, host="0.0.0.0",port=5000)
